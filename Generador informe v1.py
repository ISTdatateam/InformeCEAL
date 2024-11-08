import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io

# Supongamos que df_resultados ya está cargado en un DataFrame.
# Aquí simulamos la carga desde un archivo CSV, pero puedes cargarlo como corresponda en tu entorno.
# df_resultados = pd.read_csv("ruta_al_archivo.csv")

# Datos de ejemplo de df_resultados
data = {
    "Rut empresa": ["76.027.979-K"],
    "Razón social": ["Summit Motors S.A."],
    "Nombre fantasía": ["Summit Motors"],
    "Nombre Centro de Trabajo (CT)": ["Antofagasta Ónix"],
    "CUV": [123456],
    "Folio": [1],
    "Organismo Administrador": ["OAL XYZ"],
    "CIIU": ["G456"],
    "Fecha de activación cuestionario": ["08-07-2024"],
    "Fecha de cierre cuestionario": ["07-08-2024"],
    "Opción de evaluación": ["Completa"],
    "¿Es espejo?": ["No"],
    "Universo de trabajadores de evaluación": [30],
    "Participación (%)": [79],
    "Estado": ["Medio"],
    "Razón aplicación cuestionario": ["Cumplimiento de norma"],
    "Nombre Archivo": ["informe_antofagasta_onix.docx"]
}
df_resultados = pd.DataFrame(data)


# Función para generar el encabezado del informe usando df_resultados
def generar_encabezado(doc, datos):
    doc.add_heading("INFORME DE IMPLEMENTACIÓN", level=1)
    doc.add_paragraph("PROTOCOLO DE VIGILANCIA DE RIESGOS PSICOSOCIALES EN EL TRABAJO")
    doc.add_paragraph(f"Razón Social: {datos['Razón social']}")
    doc.add_paragraph(f"RUT: {datos['Rut empresa']}")
    doc.add_paragraph(f"Nombre del centro de trabajo o aglomeración: {datos['Nombre Centro de Trabajo (CT)']}")
    doc.add_paragraph(f"CUV: {datos['CUV']}")
    doc.add_paragraph(f"Organismo Administrador: {datos['Organismo Administrador']}")
    doc.add_paragraph(f"CIIU: {datos['CIIU']}")
    doc.add_paragraph(f"Fecha de activación del cuestionario: {datos['Fecha de activación cuestionario']}")
    doc.add_paragraph(f"Fecha de cierre del cuestionario: {datos['Fecha de cierre cuestionario']}")
    doc.add_paragraph(f"Universo de trabajadores de evaluación: {datos['Universo de trabajadores de evaluación']}")
    doc.add_paragraph(f"Participación: {datos['Participación (%)']}%")
    doc.add_paragraph(f"Estado de Riesgo: {datos['Estado']}")
    doc.add_paragraph(f"Razón aplicación cuestionario: {datos['Razón aplicación cuestionario']}")


# Función para generar el listado de medidas con datos dummy
def generar_listado_medidas(doc):
    doc.add_heading("LISTADO DE MEDIDAS", level=2)
    medidas = [
        {
            "medida": "Revisar y definir puestos de trabajo clave para el desarrollo de la operación.",
            "dimensiones": "Carga de trabajo",
            "origen_riesgo": "Debido a la falta de puesto de trabajo clave, los vendedores realizan labores extras.",
            "alcance": "Ventas",
            "plazo": "Corto plazo (180 días)",
            "responsable": "Gerencia de área"
        }
    ]
    for medida in medidas:
        doc.add_paragraph(f"MEDIDA: {medida['medida']}")
        doc.add_paragraph(f"Dimensión(es): {medida['dimensiones']}")
        doc.add_paragraph(f"Origen del Riesgo: {medida['origen_riesgo']}")
        doc.add_paragraph(f"Alcance (GES): {medida['alcance']}")
        doc.add_paragraph(f"Plazo: {medida['plazo']}")
        doc.add_paragraph(f"Responsable: {medida['responsable']}")
        doc.add_paragraph("\n")  # Espacio entre medidas


# Función para agregar un gráfico dummy
def agregar_grafico(doc):
    plt.figure(figsize=(6, 4))
    plt.bar(["Dimensión A", "Dimensión B", "Dimensión C"], [70, 85, 60])
    plt.title("Resultados de Dimensiones")
    plt.xlabel("Dimensiones")
    plt.ylabel("Puntaje")

    # Guardar el gráfico en un buffer de memoria
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)

    # Insertar el gráfico en el documento
    doc.add_paragraph("Gráfico de Resultados:")
    doc.add_picture(img_stream, width=Inches(5))
    img_stream.close()
    plt.close()


# Función principal para generar el informe para cada fila de df_resultados
def generar_informe(df):
    for index, row in df.iterrows():
        doc = Document()

        # Generar contenido del informe
        generar_encabezado(doc, row)
        doc.add_heading("RESULTADOS CUESTIONARIO CEAL-SM SUSESO", level=2)
        doc.add_paragraph("Datos del cuestionario con resultados dummy.")
        generar_listado_medidas(doc)
        agregar_grafico(doc)

        # Guardar el informe con el nombre de archivo especificado en el DataFrame
        nombre_archivo = row["Nombre Archivo"]
        doc.save(nombre_archivo)
        print(f"Informe generado: {nombre_archivo}")


# Generar informes
generar_informe(df_resultados)
