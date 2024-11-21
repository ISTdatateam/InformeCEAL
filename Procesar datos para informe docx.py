import os
import pandas as pd
import re
import logging


# Configuración de logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

'''# Definir la ruta de la carpeta
# Definir la ruta de la carpeta
folder_path = r'H:\Mi unidad\Informes PBI\Paneles CEAL\SMU\Buen trato\Archivos'
output_path = r'H:\Mi unidad\Informes PBI\Paneles CEAL\SMU\Buen trato\salida_test.xlsx'
resultados_path = r'H:\Mi unidad\Informes PBI\Paneles CEAL\SMU\resultados.xlsx'
output_archivos = r'H:\Mi unidad\Informes PBI\Paneles CEAL\SMU\Buen trato'
'''

folder_path = r'H:\Mi unidad\SM-CEAL\Reporteria masiva\tablas'
output_path = r'H:\Mi unidad\SM-CEAL\Reporteria masiva\salida_test.xlsx'
resultados_path = r'H:\Mi unidad\SM-CEAL\Reporteria masiva\database.xlsx'
output_archivos = r'H:\Mi unidad\SM-CEAL\Reporteria masiva'






# Obtener una lista de todos los archivos .xlsx en la carpeta
file_list = [f for f in os.listdir(folder_path) if f.endswith('.xlsx')]

# Listas para almacenar los datos de cada hoja
data_frames_base_completa = []
resumen_info = []
ct_values = {}
cuv_values = {}

ceal = [
    {"Coddim": "CT", "Dimensión": "Carga de trabajo", "Codpreg": "QD1",
     "Pregunta": "¿Su carga de trabajo se distribuye de manera desigual de modo que se le acumula el trabajo?"},
    {"Coddim": "CT", "Dimensión": "Carga de trabajo", "Codpreg": "QD2",
     "Pregunta": "¿Con qué frecuencia le falta tiempo para completar sus tareas?"},
    {"Coddim": "CT", "Dimensión": "Carga de trabajo", "Codpreg": "QD3",
     "Pregunta": "¿Se retrasa en la entrega de su trabajo?"},
    {"Coddim": "EM", "Dimensión": "Exigencias emocionales", "Codpreg": "ED1",
     "Pregunta": "Su trabajo, ¿le coloca en situaciones emocionalmente perturbadoras?"},
    {"Coddim": "EM", "Dimensión": "Exigencias emocionales", "Codpreg": "ED2",
     "Pregunta": "Como parte de su trabajo, ¿tiene que lidiar con los problemas personales de usuarios o clientes?"},
    {"Coddim": "EM", "Dimensión": "Exigencias emocionales", "Codpreg": "HE2",
     "Pregunta": "Su trabajo, ¿le exige esconder sus emociones?"},
    {"Coddim": "DP", "Dimensión": "Desarrollo profesional", "Codpreg": "DP2",
     "Pregunta": "¿Tiene la posibilidad de adquirir nuevos conocimientos a través de su trabajo?"},
    {"Coddim": "DP", "Dimensión": "Desarrollo profesional", "Codpreg": "DP3",
     "Pregunta": "En su trabajo, ¿puede utilizar sus habilidades o experiencia?"},
    {"Coddim": "DP", "Dimensión": "Desarrollo profesional", "Codpreg": "DP4",
     "Pregunta": "Su trabajo, ¿le da la oportunidad de desarrollar sus habilidades?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "PR2",
     "Pregunta": "¿Recibe toda la información que necesita para hacer bien su trabajo?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "RE1",
     "Pregunta": "Su trabajo, ¿es reconocido y valorado por sus superiores?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "RE2",
     "Pregunta": "En su trabajo, ¿es respetado por sus superiores?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "RE3",
     "Pregunta": "En su trabajo, ¿es tratado de forma justa?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "MW1",
     "Pregunta": "Su trabajo, ¿tiene sentido para usted?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "CL1",
     "Pregunta": "Su trabajo, ¿tiene objetivos claros?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "CL2",
     "Pregunta": "En su trabajo, ¿sabe exactamente qué tareas son de su responsabilidad?"},
    {"Coddim": "RC", "Dimensión": "Reconocimiento y claridad de rol", "Codpreg": "CL3",
     "Pregunta": "¿Sabe exactamente lo que se espera de usted en el trabajo?"},
    {"Coddim": "CR", "Dimensión": "Conflicto de rol", "Codpreg": "CO2",
     "Pregunta": "En su trabajo, ¿se le exigen cosas contradictorias?"},
    {"Coddim": "CR", "Dimensión": "Conflicto de rol", "Codpreg": "CO3",
     "Pregunta": "¿Tiene que hacer tareas que usted cree que deberían hacerse de otra manera?"},
    {"Coddim": "CR", "Dimensión": "Conflicto de rol", "Codpreg": "IT1",
     "Pregunta": "¿Tiene que realizar tareas que le parecen innecesarias?"},
    {"Coddim": "QL", "Dimensión": "Calidad del liderazgo", "Codpreg": "QL3",
     "Pregunta": "Su superior inmediato, ¿planifica bien el trabajo?"},
    {"Coddim": "QL", "Dimensión": "Calidad del liderazgo", "Codpreg": "QL4",
     "Pregunta": "Su superior inmediato, ¿resuelve bien los conflictos?"},
    {"Coddim": "QL", "Dimensión": "Calidad del liderazgo", "Codpreg": "SS1",
     "Pregunta": "Si usted lo necesita, ¿con qué frecuencia su superior inmediato está dispuesto a escuchar sus problemas?"},
    {"Coddim": "QL", "Dimensión": "Calidad del liderazgo", "Codpreg": "SS2",
     "Pregunta": "Si usted lo necesita, ¿con qué frecuencia obtiene ayuda y apoyo de su superior inmediato?"},
    {"Coddim": "CM", "Dimensión": "Compañerismo", "Codpreg": "SC1",
     "Pregunta": "De ser necesario, ¿con qué frecuencia obtiene ayuda y apoyo de sus compañeros(as) de trabajo?"},
    {"Coddim": "CM", "Dimensión": "Compañerismo", "Codpreg": "SC2",
     "Pregunta": "De ser necesario, ¿con qué frecuencia sus compañeros(as) de trabajo están dispuestos(as) a escuchar problemas?"},
    {"Coddim": "CM", "Dimensión": "Compañerismo", "Codpreg": "SW1",
     "Pregunta": "¿Hay un buen ambiente entre usted y sus compañeros(as) de trabajo?"},
    {"Coddim": "CM", "Dimensión": "Compañerismo", "Codpreg": "SW3",
     "Pregunta": "En su trabajo, ¿usted siente que forma parte de un equipo?"},
    {"Coddim": "IT", "Dimensión": "Inseguridad en las condiciones de trabajo", "Codpreg": "IW1",
     "Pregunta": "¿Está preocupado(a) de que le cambien sus tareas laborales en contra de su voluntad?"},
    {"Coddim": "IT", "Dimensión": "Inseguridad en las condiciones de trabajo", "Codpreg": "IW2",
     "Pregunta": "¿Está preocupado(a) por si le trasladan a otro lugar de trabajo, obra, funciones, unidad, departamento o sección en contra de su voluntad?"},
    {"Coddim": "IT", "Dimensión": "Inseguridad en las condiciones de trabajo", "Codpreg": "IW3",
     "Pregunta": "¿Está preocupado(a) de que le cambien el horario (turnos, días de la semana, hora de entrada y salida) en contra de su voluntad?"},
    {"Coddim": "TV", "Dimensión": "Equilibrio trabajo y vida privada", "Codpreg": "WF2",
     "Pregunta": "¿Siente que su trabajo le consume demasiada ENERGÍA teniendo un efecto negativo en su vida privada?"},
    {"Coddim": "TV", "Dimensión": "Equilibrio trabajo y vida privada", "Codpreg": "WF3",
     "Pregunta": "¿Siente que su trabajo le consume demasiado TIEMPO teniendo un efecto negativo en su vida privada?"},
    {"Coddim": "TV", "Dimensión": "Equilibrio trabajo y vida privada", "Codpreg": "WF5",
     "Pregunta": "Las exigencias de su trabajo, ¿interfieren con su vida privada y familiar?"},
    {"Coddim": "CJ", "Dimensión": "Confianza y justicia organizacional", "Codpreg": "TE1",
     "Pregunta": "En general, ¿los trabajadores(as) en su organización confían entre sí?"},
    {"Coddim": "CJ", "Dimensión": "Confianza y justicia organizacional", "Codpreg": "TM1",
     "Pregunta": "¿Los gerentes o directivos confían en que los trabajadores(as) hacen bien su trabajo?"},
    {"Coddim": "CJ", "Dimensión": "Confianza y justicia organizacional", "Codpreg": "TM2",
     "Pregunta": "¿Los trabajadores(as) confían en la información que proviene de los gerentes, directivos o empleadores?"},
    {"Coddim": "CJ", "Dimensión": "Confianza y justicia organizacional", "Codpreg": "TM4",
     "Pregunta": "¿Los trabajadores(as) pueden expresar sus opiniones y sentimientos?"},
    {"Coddim": "CJ", "Dimensión": "Confianza y justicia organizacional", "Codpreg": "JU1",
     "Pregunta": "En su trabajo, ¿los conflictos se resuelven de manera justa?"},
    {"Coddim": "CJ", "Dimensión": "Confianza y justicia organizacional", "Codpreg": "JU2",
     "Pregunta": "¿Se valora a los trabajadores(as) cuando han hecho un buen trabajo?"},
    {"Coddim": "CJ", "Dimensión": "Confianza y justicia organizacional", "Codpreg": "JU4",
     "Pregunta": "¿Se distribuye el trabajo de manera justa?"},
    {"Coddim": "VU", "Dimensión": "Vulnerabilidad", "Codpreg": "VU1",
     "Pregunta": "¿Tiene miedo a pedir mejores condiciones de trabajo?"},
    {"Coddim": "VU", "Dimensión": "Vulnerabilidad", "Codpreg": "VU2",
     "Pregunta": "¿Se siente indefenso(a) ante el trato injusto de sus superiores?"},
    {"Coddim": "VU", "Dimensión": "Vulnerabilidad", "Codpreg": "VU3",
     "Pregunta": "¿Tiene miedo de que lo(la) despidan si no hace lo que le piden?"},
    {"Coddim": "VU", "Dimensión": "Vulnerabilidad", "Codpreg": "VU4",
     "Pregunta": "¿Considera que sus superiores lo(la) tratan de forma discriminatoria o injusta?"},
    {"Coddim": "VU", "Dimensión": "Vulnerabilidad", "Codpreg": "VU5",
     "Pregunta": "¿Considera que lo(la) tratan de forma autoritaria o violenta?"},
    {"Coddim": "VU", "Dimensión": "Vulnerabilidad", "Codpreg": "VU6",
     "Pregunta": "¿Lo(la) hacen sentir que usted puede ser fácilmente reemplazado(a)?"},
    {"Coddim": "VA", "Dimensión": "Violencia y acoso", "Codpreg": "CQ1",
     "Pregunta": "En su trabajo, durante los últimos 12 meses, ¿ha estado involucrado(a) en disputas o conflictos?"},
    {"Coddim": "VA", "Dimensión": "Violencia y acoso", "Codpreg": "UT1",
     "Pregunta": "En su trabajo, durante los últimos 12 meses, ¿ha estado expuesto(a) a bromas desagradables?"},
    {"Coddim": "VA", "Dimensión": "Violencia y acoso", "Codpreg": "HSM1",
     "Pregunta": "En los últimos 12 meses, ¿ha estado expuesto(a) a acoso relacionado al trabajo por correo electrónico, mensajes de texto y/o en las redes sociales (por ejemplo, Facebook, Instagram, Twitter)?"},
    {"Coddim": "VA", "Dimensión": "Violencia y acoso", "Codpreg": "SH1",
     "Pregunta": "En su trabajo, durante los últimos 12 meses, ¿ha estado expuesta(o) a acoso sexual?"},
    {"Coddim": "VA", "Dimensión": "Violencia y acoso", "Codpreg": "PV1",
     "Pregunta": "En su trabajo, en los últimos 12 meses, ¿ha estado expuesta(o) a violencia física?"},
    {"Coddim": "VA", "Dimensión": "Violencia y acoso", "Codpreg": "AL",
     "Pregunta": "En su trabajo, en los últimos 12 meses, ¿ha estado expuesto(a) a bullying o acoso?"},
    {"Coddim": "VA", "Dimensión": "Violencia y acoso", "Codpreg": "HO",
     "Pregunta": "¿Con qué frecuencia se siente intimidado(a), colocado(a) en ridículo o injustamente criticado(a), frente a otros por sus compañeros(as) de trabajo o su superior?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ1",
     "Pregunta": "¿Ha podido concentrarse bien en lo que hace?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ2",
     "Pregunta": "¿Sus preocupaciones le han hecho perder mucho sueño?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ3",
     "Pregunta": "¿Ha sentido que está jugando un papel útil en la vida?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ4",
     "Pregunta": "¿Se ha sentido capaz de tomar decisiones?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ5",
     "Pregunta": "¿Se ha sentido constantemente agobiado(a) y en tensión?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ6",
     "Pregunta": "¿Ha sentido que no puede superar sus dificultades?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ7",
     "Pregunta": "¿Ha sido capaz de disfrutar sus actividades normales de cada día?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ8",
     "Pregunta": "¿Ha sido capaz de hacer frente a sus problemas?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ9",
     "Pregunta": "¿Se ha sentido poco feliz y deprimido(a)?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ10",
     "Pregunta": "¿Ha perdido confianza en sí mismo?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ11",
     "Pregunta": "¿Ha pensado que usted es una persona que no vale para nada?"},
    {"Coddim": "GHQ", "Dimensión": "Cuestionario de salud general", "Codpreg": "GHQ12",
     "Pregunta": "¿Se siente razonablemente feliz considerando todas las circunstancias?"}
]

df_ceal = pd.DataFrame(ceal)

# Arreglo con la información de la tabla
risk_intervals = [
    {"Dimensión": "Carga de trabajo", "Nivel de riesgo bajo": (0, 1), "Nivel de riesgo medio": (2, 4),
     "Nivel de riesgo alto": (5, 12)},
    {"Dimensión": "Exigencias emocionales", "Nivel de riesgo bajo": (0, 1), "Nivel de riesgo medio": (2, 5),
     "Nivel de riesgo alto": (6, 12)},
    {"Dimensión": "Desarrollo profesional", "Nivel de riesgo bajo": (0, 1), "Nivel de riesgo medio": (2, 5),
     "Nivel de riesgo alto": (6, 12)},
    {"Dimensión": "Reconocimiento y claridad de rol", "Nivel de riesgo bajo": (0, 4), "Nivel de riesgo medio": (5, 9),
     "Nivel de riesgo alto": (10, 32)},
    {"Dimensión": "Conflicto de rol", "Nivel de riesgo bajo": (0, 2), "Nivel de riesgo medio": (3, 5),
     "Nivel de riesgo alto": (6, 12)},
    {"Dimensión": "Calidad del liderazgo", "Nivel de riesgo bajo": (0, 2), "Nivel de riesgo medio": (3, 7),
     "Nivel de riesgo alto": (8, 16)},
    {"Dimensión": "Compañerismo", "Nivel de riesgo bajo": (0, 0), "Nivel de riesgo medio": (1, 4),
     "Nivel de riesgo alto": (5, 16)},
    {"Dimensión": "Inseguridad en las condiciones de trabajo", "Nivel de riesgo bajo": (0, 2),
     "Nivel de riesgo medio": (3, 5), "Nivel de riesgo alto": (6, 12)},
    {"Dimensión": "Equilibrio trabajo y vida privada", "Nivel de riesgo bajo": (0, 2), "Nivel de riesgo medio": (3, 5),
     "Nivel de riesgo alto": (6, 12)},
    {"Dimensión": "Confianza y justicia organizacional", "Nivel de riesgo bajo": (0, 7),
     "Nivel de riesgo medio": (8, 12), "Nivel de riesgo alto": (13, 28)},
    {"Dimensión": "Violencia y acoso", "Nivel de riesgo bajo": (0, 0), "Nivel de riesgo medio": (1, 14),
     "Nivel de riesgo alto": (15, 28)},
    {"Dimensión": "Vulnerabilidad", "Nivel de riesgo bajo": (1, 6), "Nivel de riesgo medio": (7, 11),
     "Nivel de riesgo alto": (12, 24)}
]

# Crear el DataFrame a partir del arreglo
df_risk_intervals = pd.DataFrame(risk_intervals)


def rename_duplicate_columns(df):
    # Renombrar la columna 5 como 'CdT' si se llama 'TE1'
    if df.columns[2] == 'TE1':
        df.columns.values[2] = 'CdT'
    # Renombrar la columna 62 como 'TE1' si no se llama 'TE1'
    if df.columns[59] != 'TE1':
        df.columns.values[59] = 'TE1'

    col_count = df.columns.value_counts()
    duplicate_cols = col_count[col_count > 1].index.tolist()

    for col in duplicate_cols:
        duplicate_indices = [i for i, x in enumerate(df.columns) if x == col]
        for index, i in enumerate(duplicate_indices):
            if index == 0 and col == 'TE1':
                continue
            df.columns.values[i] = f"{col}_{index}" if index > 0 else col
    return df


for file_name in file_list:
    file_path = os.path.join(folder_path, file_name)

    # Leer la hoja "BaseCompleta"
    try:
        df_base_completa = pd.read_excel(file_path, sheet_name='BaseCompleta', header=1, usecols='C:CO')
        df_base_completa = rename_duplicate_columns(df_base_completa)
        df_base_completa.rename(columns={'DD1': 'Genero', 'DD2': 'Edad'}, inplace=True)
        df_base_completa['Genero'] = df_base_completa['Genero'].replace(
            {1: 'Hombre', 2: 'Mujer', 3: 'NcOtro', 4: 'NcOtro'})
        #       ct_values[file_name] = df_base_completa['CdT'].iloc[0]

        # Extraer RUT y CUV del nombre del archivo
        rut_match = re.search(r'\d{8}-[\dkK]', file_name)
        cuv_match = re.search(r'(\d+)(?=\.xlsx)', file_name)
        cdt_match = re.search(r'\d{8}-[\dkK]-(.*?)-\d+\.xlsx', file_name)

        # Imprimir resultados de las coincidencias de regex
        print(f"Procesando archivo: {file_name}")
        print(f"{rut_match}")
        print(f"{cuv_match}")
        print(f"{cdt_match}")

        if rut_match and cuv_match and cdt_match:
            rut = rut_match.group(0)
            cuv = cuv_match.group(1)
            cdtm = cdt_match.group(1)
            print(f"Archivo: {file_name} - RUT: {rut} - CUV: {cuv} - CDT: {cdtm}")
            df_base_completa['RUT_empleador'] = rut
            df_base_completa['CUV'] = cuv
            df_base_completa['CDT_glosa'] = cdtm
        else:
            print(f"Error extrayendo RUT o CUV del archivo: {file_name}")

        ct_values[file_name] = df_base_completa['CUV'].iloc[0]
        data_frames_base_completa.append(df_base_completa)

    except Exception as e:
        print(f"Error procesando el archivo {file_name}: {e}")

# Combinar todos los DataFrames de la hoja "BaseCompleta" en uno solo
combined_df_base_completa = pd.concat(data_frames_base_completa, ignore_index=True)


# Función para comparar y modificar los campos
def compare_and_concat(row):
    cdt = row['CdT'].strip()
    cdt_glosa = row['CDT_glosa'].strip()

    if cdt != cdt_glosa:
        return f"{cdt_glosa} - {cdt}"
    return cdt


# Aplicar la función a cada fila del DataFrame y eliminar la columna 'CDT_glosa'
combined_df_base_completa['CdT'] = combined_df_base_completa.apply(compare_and_concat, axis=1)
combined_df_base_completa = combined_df_base_completa.drop(columns=['CDT_glosa'])

# Crear una nueva columna para los rangos de edad
bins = [18, 25, 36, 49, float('inf')]
labels = ['18 a 25', '26 a 36', '37 a 49', '50 o más']
combined_df_base_completa['Rango Edad'] = pd.cut(combined_df_base_completa['Edad'], bins=bins, labels=labels,
                                                 right=False)


# Definir la función para asignar puntaje
def calcular_puntaje(row):
    if row['Dimensión'] == 'Salud mental (GHQ)':
        return None
    if row['Nivel'] == 'Bajo' and row['Porcentaje'] >= 50:
        return -2
    elif row['Nivel'] == 'Medio' and row['Porcentaje'] >= 50:
        return 1
    elif row['Nivel'] == 'Alto' and row['Porcentaje'] >= 50:
        return 2
    else:
        return None


# Calcular el Riesgo basado en el Puntaje
def calcular_riesgo(puntaje):
    if puntaje <= 1:
        return 'Riesgo bajo'
    elif 2 <= puntaje <= 12:
        return 'Riesgo medio'
    else:
        return 'Riesgo alto'


# Crear un diccionario donde las claves sean los Coddim y los valores sean listas de Codpreg
coddim_to_codpreg = df_ceal.groupby('Coddim')['Codpreg'].apply(list).to_dict()

# Crear las columnas de subtotales en df_base_completa
for coddim, codpreg_list in coddim_to_codpreg.items():
    # Sumar las columnas correspondientes a cada Codpreg y crear una nueva columna con el subtotal
    combined_df_base_completa[coddim] = combined_df_base_completa[codpreg_list].sum(axis=1)

# Crear un diccionario que mapee los Coddim a los nombres de la dimensión
coddim_to_dimension = df_ceal[['Coddim', 'Dimensión']].drop_duplicates().set_index('Coddim')['Dimensión'].to_dict()


# Función para determinar el nivel de riesgo basado en el puntaje
def determinar_nivel_riesgo(dim, puntaje):
    fila = df_risk_intervals[df_risk_intervals['Dimensión'] == dim]
    if fila.empty:
        return 'dimensión no encontrada'
    fila = fila.iloc[0]
    if fila['Nivel de riesgo bajo'][0] <= puntaje <= fila['Nivel de riesgo bajo'][1]:
        return 'Bajo'
    elif fila['Nivel de riesgo medio'][0] <= puntaje <= fila['Nivel de riesgo medio'][1]:
        return 'Medio'
    elif fila['Nivel de riesgo alto'][0] <= puntaje <= fila['Nivel de riesgo alto'][1]:
        return 'Alto'
    else:
        return 'fuera de rango'


# Aplicar la función para determinar el nivel de riesgo a cada dimensión
for coddim, dimension in coddim_to_dimension.items():
    column_name = f'{coddim}_riesgo'
    if dimension in df_risk_intervals['Dimensión'].values:
        combined_df_base_completa[column_name] = combined_df_base_completa[coddim].apply(
            lambda x: determinar_nivel_riesgo(dimension, x))
    else:
        print(f"La dimensión '{dimension}' no está presente en df_risk_intervals")

print("listo")

combined_df_base_complet2 = combined_df_base_completa


def calcular_porcentaje_respuestas_nuevas_columnas(df, nuevas_columnas):
    resultados = []
    total_respuestas = df.shape[0]
    mapping01 = {0: "No expuesto", 1: "Expuesto"}

    for columna in nuevas_columnas:
        for valor in [0, 1]:
            conteo = df[columna].value_counts().get(valor, 0)
            porcentaje = round((conteo / total_respuestas) * 100, 2)
            resultados.append({
                'CUV': df['CUV'].unique()[0],
                'CdT': df['CdT'].unique()[0],
                'Codpreg': columna,
                'Valor': valor,
                'Exposición': mapping01[valor],
                'Porcentaje': porcentaje,
                'Respuestas': conteo
            })
    return pd.DataFrame(resultados)


# Crear nuevas columnas basadas en la condición de ser diferente de 0
columnas_originales = ['CQ1', 'UT1', 'HSM1', 'SH1', 'PV1', 'AL', 'HO']
nuevas_columnas = []
for columna in columnas_originales:
    nueva_columna = columna + '_01'
    combined_df_base_complet2[nueva_columna] = combined_df_base_complet2[columna].apply(lambda x: 1 if x != 0 else 0)
    nuevas_columnas.append(nueva_columna)

# Crear una nueva columna que cuente cuántos 1 hay en las nuevas columnas
combined_df_base_complet2['Expo_total'] = combined_df_base_complet2[nuevas_columnas].sum(axis=1)
combined_df_base_complet2['Expo_total_01'] = combined_df_base_complet2['Expo_total'].apply(lambda x: 1 if x != 0 else 0)
nuevas_columnas.append('Expo_total_01')

df_exposicionviolencia = combined_df_base_complet2.groupby('CUV')[nuevas_columnas].agg(['mean'])
df_exposicionviolencia.columns = ['_'.join(col).strip() for col in df_exposicionviolencia.columns.values]
df_exposicionviolencia = df_exposicionviolencia.reset_index()


# Calcular porcentajes para las nuevas columnas
df_porcentajes = calcular_porcentaje_respuestas_nuevas_columnas(combined_df_base_complet2, nuevas_columnas)

# df_porcentajes
df_resultados_porcentaje_nuevas = pd.concat(
    [calcular_porcentaje_respuestas_nuevas_columnas(grupo, nuevas_columnas)
     for _, grupo in combined_df_base_complet2.groupby('CUV')],
    ignore_index=True
)

violencia = [
    {"Codpreg": "CQ1_01", "Temática": "Disputas o conflictos"},
    {"Codpreg": "UT1_01", "Temática": "Bromas desagradables"},
    {"Codpreg": "HSM1_01", "Temática": "Acoso virtual"},
    {"Codpreg": "SH1_01", "Temática": "Acoso sexual"},
    {"Codpreg": "PV1_01", "Temática": "Violencia física"},
    {"Codpreg": "AL_01", "Temática": "Bullying o acoso"},
    {"Codpreg": "HO_01", "Temática": "Humillaciones"},
    {"Codpreg": "Expo_total_01", "Temática": "Exposicion a violencia"}
]
# Crear el DataFrame a partir del arreglo
df_violencia = pd.DataFrame(violencia)

# Realizar un merge con el DataFrame de resultados de porcentajes
df_resultados_porcentaje_nuevas = df_resultados_porcentaje_nuevas.merge(df_violencia, on='Codpreg', how='left')

combined_df_base_complet3 = combined_df_base_complet2


def calcular_porcentaje_respuestas_nuevas_columnasp(df, nuevas_columnasn, nuevas_columnasd):
    resultados2 = []
    mapping01 = {0: "No Proteccion", 1: "Proteccion"}

    for coln, cold in zip(nuevas_columnasn, nuevas_columnasd):
        for valor in [0, 1]:
            conteon = df[coln].value_counts().get(valor, 0)
            conteod = df[cold].value_counts().get(1, 0)
            porcentaje = round((conteon / conteod) * 100, 2) if conteod != 0 else 0

            resultados2.append({
                'CUV': df['CUV'].unique()[0],
                'CdT': df['CdT'].unique()[0],
                'Codpreg': coln,
                'Exposición': mapping01[valor],  # Assuming you want to use "No proteccion" for Exposición
                'Valor': valor,
                'Porcentaje': porcentaje,
                'Respuestan': conteon,
                'Respuestad': conteod
            })

    return pd.DataFrame(resultados2)


# Definir las columnas originales
columnas_originalesp = ['SS1', 'SS2', 'SC1', 'SC2', 'SW1', 'SW3']
nuevas_columnasn = []
nuevas_columnasd = []

# Crear nuevas columnas 'n' y 'd' basadas en las reglas especificadas
for columna in columnas_originalesp:
    nueva_columnan = columna + '_n'
    nueva_columnad = columna + '_d'

    # Aplicar reglas para nueva_columnan
    combined_df_base_complet3[nueva_columnan] = combined_df_base_complet3[columna].apply(
        lambda x: 1 if x in [0, 1] else (np.nan if x == 5 else 0))
    nuevas_columnasn.append(nueva_columnan)

    # Aplicar reglas para nueva_columnad
    combined_df_base_complet3[nueva_columnad] = combined_df_base_complet3[columna].apply(
        lambda x: 1 if x in [0, 1, 2, 3, 4] else 0)
    nuevas_columnasd.append(nueva_columnad)

# Calcular porcentajes para las nuevas columnas
df_resultados_porcentaje_nuevas2 = pd.concat(
    [calcular_porcentaje_respuestas_nuevas_columnasp(grupo, nuevas_columnasn, nuevas_columnasd)
     for _, grupo in combined_df_base_complet3.groupby('CUV')],
    ignore_index=True
)

protectores = [
    {"Codpreg": "SS1_n", "Temática": "superior1"},
    {"Codpreg": "SS2_n", "Temática": "superior2"},
    {"Codpreg": "SC1_n", "Temática": "compañeros1"},
    {"Codpreg": "SC2_n", "Temática": "compañeros2"},
    {"Codpreg": "SW1_n", "Temática": "oficina1"},
    {"Codpreg": "SW3_n", "Temática": "oficina2"}
]

# Crear el DataFrame a partir del arreglo
df_protectores = pd.DataFrame(protectores)

# Realizar un merge con el DataFrame de resultados de porcentajes
df_resultados_porcentaje_nuevas2 = df_resultados_porcentaje_nuevas2.merge(df_protectores, on='Codpreg', how='left')

# Verificar el resultado
#print(df_resultados_porcentaje_nuevas2)


def calcular_porcentaje_respuestas(df, coddim_to_dimension, df_risk_intervals):
    resultados = []
    nivel_mapping = {'Bajo': 1, 'Medio': 2, 'Alto': 3}

    for coddim, dimension in coddim_to_dimension.items():
        if dimension in df_risk_intervals['Dimensión'].values:
            columna_riesgo = f'{coddim}_riesgo'
            total_respuestas = df.shape[0]
            for nivel in ['Bajo', 'Medio', 'Alto']:
                conteo = df[columna_riesgo].value_counts().get(nivel, 0)
                porcentaje = round((conteo / total_respuestas) * 100, 2)
                resultados.append({
                    'CUV': df['CUV'].unique()[0],
                    'CdT': df['CdT'].unique()[0],
                    'Dimensión': dimension,
                    'Nivel': nivel,
                    'Nivel_n': nivel_mapping[nivel],
                    'Porcentaje': porcentaje,
                    'Respuestas': conteo
                })
    return pd.DataFrame(resultados)


# Aplicar la función a cada grupo de CdT y concatenar los resultados
df_resultados_porcentaje = pd.concat([calcular_porcentaje_respuestas(grupo, coddim_to_dimension, df_risk_intervals)
                                      for _, grupo in combined_df_base_complet3.groupby('CUV')], ignore_index=True)

df_resultados_porcentaje['Puntaje'] = df_resultados_porcentaje.apply(calcular_puntaje, axis=1)

# Crear un DataFrame con los puntajes máximos por 'CdT' y 'Dimensión'
df_puntajes_max = df_resultados_porcentaje.dropna(subset=['Puntaje'])
df_puntajes_max = (df_puntajes_max.groupby(['CUV', 'Dimensión'])
                   .apply(lambda x: x.loc[x['Puntaje'].idxmax()])
                   .reset_index(drop=True))

# Hacer un merge con el DataFrame original utilizando 'CdT', 'Dimensión', 'Nivel'
df_resultados_porcentaje = df_resultados_porcentaje.merge(
    df_puntajes_max[['CUV', 'CdT', 'Dimensión', 'Nivel', 'Nivel_n', 'Porcentaje', 'Respuestas', 'Puntaje']],
    on=['CUV', 'Dimensión', 'Nivel'],
    suffixes=('', '_max'),
    how='left'
)

# Actualizar la columna 'Puntaje' con los valores de 'Puntaje_max' y eliminar 'Puntaje_max'
df_resultados_porcentaje['Puntaje'] = df_resultados_porcentaje['Puntaje_max']
df_resultados_porcentaje.drop(columns=['CdT_max', 'Puntaje_max', 'Nivel_n_max', 'Porcentaje_max', 'Respuestas_max'],
                              inplace=True)

# Asegurarse de que las columnas 'CUV' y 'CdT' sean del tipo string
df_resultados_porcentaje['CUV'] = df_resultados_porcentaje['CUV'].astype(str)
df_resultados_porcentaje['CdT'] = df_resultados_porcentaje['CdT'].astype(str)

# Agrupar por archivo y CdT para calcular Puntaje total y número de evaluaciones
summary_df = df_resultados_porcentaje.groupby(['CUV', 'CdT']).agg(
    Puntaje=('Puntaje', 'sum')
).reset_index()

summary_df['Riesgo'] = summary_df['Puntaje'].apply(calcular_riesgo)

# Reordenar las columnas en el orden deseado
summary_df = summary_df[['CUV', 'CdT', 'Puntaje', 'Riesgo']]

print("Listo")




############

def obtener_porcentaje_niveles(df, coddim_to_dimension, df_intervalos_riesgo):
    resultados = []
    nivel_mapping = {'Bajo': 1, 'Medio': 2, 'Alto': 3}

    for coddim, dimension in coddim_to_dimension.items():
        if dimension in df_intervalos_riesgo['Dimensión'].values:
            columna_riesgo = f'{coddim}_riesgo'
            total_respuestas = df.shape[0]
            for nivel in ['Bajo', 'Medio', 'Alto']:
                conteo = df[columna_riesgo].value_counts().get(nivel, 0)
                porcentaje = round((conteo / total_respuestas) * 100, 2)
                resultados.append({
                    'CUV': df['CUV'].unique()[0],
                    'CdT': df['CdT'].unique()[0],
                    'TE3': df['TE3'].unique()[0],  # Agregar TE3
                    'Dimensión': dimension,
                    'Nivel': nivel,
                    'Nivel_n': nivel_mapping[nivel],
                    'Porcentaje': porcentaje,
                    'Respuestas': conteo
                })
    return pd.DataFrame(resultados)


# Aplicar la función a cada grupo de CUV y TE3, y concatenar los resultados
df_porcentajes_niveles = pd.concat([
    obtener_porcentaje_niveles(grupo, coddim_to_dimension, df_risk_intervals)
    for _, grupo in combined_df_base_complet3.groupby(['CUV', 'TE3'])
], ignore_index=True)

df_porcentajes_niveles['Puntaje'] = df_porcentajes_niveles.apply(calcular_puntaje, axis=1)

# Crear un DataFrame con los puntajes máximos por 'CUV', 'TE3' y 'Dimensión'
df_max_puntajes = df_porcentajes_niveles.dropna(subset=['Puntaje'])
df_max_puntajes = (df_max_puntajes.groupby(['CUV', 'TE3', 'Dimensión'])
                   .apply(lambda x: x.loc[x['Puntaje'].idxmax()])
                   .reset_index(drop=True))

# Hacer un merge con el DataFrame original utilizando 'CUV', 'TE3', 'Dimensión' y 'Nivel'
df_porcentajes_niveles = df_porcentajes_niveles.merge(
    df_max_puntajes[['CUV', 'TE3', 'CdT', 'Dimensión', 'Nivel', 'Nivel_n', 'Porcentaje', 'Respuestas', 'Puntaje']],
    on=['CUV', 'TE3', 'Dimensión', 'Nivel'],
    suffixes=('', '_max'),
    how='left'
)

# Actualizar la columna 'Puntaje' con los valores de 'Puntaje_max' y eliminar 'Puntaje_max'
df_porcentajes_niveles['Puntaje'] = df_porcentajes_niveles['Puntaje_max']
df_porcentajes_niveles.drop(columns=['CdT_max', 'Puntaje_max', 'Nivel_n_max', 'Porcentaje_max', 'Respuestas_max'],
                            inplace=True)

# Asegurarse de que las columnas 'CUV', 'CdT' y 'TE3' sean del tipo string
df_porcentajes_niveles['CUV'] = df_porcentajes_niveles['CUV'].astype(str)
df_porcentajes_niveles['CdT'] = df_porcentajes_niveles['CdT'].astype(str)
df_porcentajes_niveles['TE3'] = df_porcentajes_niveles['TE3'].astype(str)

# Agrupar por archivo, CdT y TE3 para calcular Puntaje total y número de evaluaciones
df_resumen = df_porcentajes_niveles.groupby(['CUV', 'CdT', 'TE3']).agg(
    Puntaje=('Puntaje', 'sum')
).reset_index()

df_resumen['Riesgo'] = df_resumen['Puntaje'].apply(calcular_riesgo)

# Reordenar las columnas en el orden deseado
df_resumen = df_resumen[['CUV', 'CdT', 'TE3', 'Puntaje', 'Riesgo']]

#df_res_dimTE3 = df_porcentajes_niveles[df_porcentajes_niveles['Puntaje'] >= 1]
df_res_dimTE3 = df_porcentajes_niveles

############


columns_to_keep = ['CUV', 'CdT', 'QD1', 'QD2', 'QD3', 'ED1', 'ED2', 'HE2', 'DP2', 'DP3', 'DP4', 'MW1', 'PR2', 'RE1',
                   'RE2', 'RE3', 'CL1', 'CL2', 'CL3', 'CO2', 'CO3',
                   'IT1', 'QL3', 'QL4', 'SS1', 'SS2', 'SC2', 'SC1', 'SW1', 'SW3', 'IW1',
                   'IW2', 'IW3', 'WF2', 'WF3', 'WF5', 'TE1', 'TM1', 'TM2', 'TM4', 'JU1',
                   'JU2', 'JU4', 'VU1', 'VU2', 'VU3', 'VU4', 'VU5', 'VU6', 'CQ1', 'UT1',
                   'HSM1', 'SH1', 'PV1', 'AL', 'HO', 'GHQ1', 'GHQ2', 'GHQ3', 'GHQ4', 'GHQ5',
                   'GHQ6', 'GHQ7', 'GHQ8', 'GHQ9', 'GHQ10', 'GHQ11', 'GHQ12']

# Crear un nuevo DataFrame con solo las columnas seleccionadas
new_df = combined_df_base_complet3[columns_to_keep]

# Obtener las columnas a analizar
columns = new_df.columns[2:]

# Resultado final
result = []

# Recorrer cada grupo de CUV y CdT
for (CUV, CdT), group in new_df.groupby(['CUV', 'CdT']):
    # Recorrer cada columna
    for col in columns:
        # Contar la frecuencia de cada valor en la columna
        freq = group[col].value_counts().sort_index()
        # Añadir los resultados a la lista final
        for value, count in freq.items():
            result.append([CUV, CdT, col, value, count])

# Convertir la lista de resultados a un DataFrame
result_df = pd.DataFrame(result, columns=['CUV', 'CdT', 'Codpreg', 'valor', 'frec'])

# Reorganizar el DataFrame para que el formato sea correcto
result_df = result_df[['CUV', 'CdT', 'Codpreg', 'valor', 'frec']]

# Agregar columnas adicionales desde df_ceal
result_df = result_df.merge(df_ceal, on='Codpreg', how='left')

# Agregar columnas adicionales desde df_ceal
df_resultados_porcentaje = df_resultados_porcentaje.merge(df_ceal[['Dimensión', 'Coddim']], on='Dimensión', how='left')
df_resultados_porcentaje = df_resultados_porcentaje.drop_duplicates(subset=None)

# Leer 'resultados.xlsx'
df_resultados = pd.read_excel(resultados_path, sheet_name='Datos', usecols=['CUV', 'Folio'])
df_res_com = pd.read_excel(resultados_path, sheet_name='Datos')


# Convertir 'CUV' a int64 en df_resultados
df_resultados['CUV'] = df_resultados['CUV'].astype('int64')

# Lista de DataFrames y sus nombres
dataframes = [combined_df_base_complet3, result_df, df_resultados_porcentaje, df_resultados_porcentaje_nuevas, df_resultados_porcentaje_nuevas2, df_exposicionviolencia,  summary_df]
nombres = ['combined_df_base_complet3', 'result_df', 'df_resultados_porcentaje', 'df_resultados_porcentaje_nuevas', 'df_resultados_porcentaje_nuevas2', 'df_exposicionviolencia','summary_df']

# Función para convertir 'CUV' a int64 y manejar errores
def convert_cuv_to_int64(df, nombre_df):
    if 'CUV' in df.columns:
        df['CUV'] = pd.to_numeric(df['CUV'], errors='coerce')
        if df['CUV'].isnull().any():
            print(f"Valores no numéricos encontrados en 'CUV' de {nombre_df}.")
            df = df.dropna(subset=['CUV'])
        df['CUV'] = df['CUV'].astype('int64')
    return df

# Aplicar la función a cada DataFrame
for df, nombre in zip(dataframes, nombres):
    df = convert_cuv_to_int64(df, nombre)

# Realizar el merge
def merge_if_cuv_exists(df):
    if 'CUV' in df.columns:
        return df.merge(df_resultados, on='CUV', how='left')
    else:
        return df

# Aplicar el merge a cada DataFrame
combined_df_base_complet3 = merge_if_cuv_exists(combined_df_base_complet3)
result_df = merge_if_cuv_exists(result_df)
df_resultados_porcentaje = merge_if_cuv_exists(df_resultados_porcentaje)
df_resultados_porcentaje_nuevas = merge_if_cuv_exists(df_resultados_porcentaje_nuevas)
df_resultados_porcentaje_nuevas2 = merge_if_cuv_exists(df_resultados_porcentaje_nuevas2)
df_exposicionviolencia= merge_if_cuv_exists(df_exposicionviolencia)
summary_df = merge_if_cuv_exists(summary_df)


df_resultados_porcentaje['Descripción'] = df_resultados_porcentaje.apply(
    lambda row: f"{row['Dimensión']} ({row['Porcentaje']}% Riesgo {row['Nivel']}, {row['Respuestas']} personas)"
    if row['Puntaje'] in [1, 2] else "",
    axis=1
)

df_porcentajes_niveles['Descripción'] = df_porcentajes_niveles.apply(
    lambda row: f"{row['Dimensión']} ({row['Porcentaje']}% Riesgo {row['Nivel']}, {row['Respuestas']} personas)"
    if row['Puntaje'] in [1, 2] else "",
    axis=1
)

# Calcular la columna 'Factor' como la multiplicación de 'valor' por 'frec'
result_df['Factor'] = result_df['valor'] * result_df['frec']
# Agrupar por 'CUV', 'Dimensión', y 'Pregunta' y sumar los valores de 'Factor' en cada grupo
result_df = result_df.groupby(['CUV', 'Dimensión', 'Pregunta'], as_index=False)['Factor'].sum()
# Para cada combinación de 'CUV' y 'Dimensión', seleccionar la pregunta con el mayor 'Factor'
top_glosas = result_df.sort_values(by='Factor', ascending=False).groupby(['CUV', 'Dimensión']).head(2)


#print(top_glosas)


# Escribir en 'combined_output.xlsx'
with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
    combined_df_base_complet3.to_excel(writer, sheet_name='basecompleta', index=False)
    result_df.to_excel(writer, sheet_name='recuentopreguntas', index=False)
    top_glosas.to_excel(writer, sheet_name='top_glosas', index=False)
    df_resultados_porcentaje.to_excel(writer, sheet_name='resultado', index=False)
    df_resultados_porcentaje_nuevas.to_excel(writer, sheet_name='violencia', index=False)
    df_resultados_porcentaje_nuevas2.to_excel(writer, sheet_name='protectores', index=False)
    df_exposicionviolencia.to_excel(writer, sheet_name='expoviolencia', index=False)
    summary_df.to_excel(writer, sheet_name='Summary', index=False)
    df_porcentajes_niveles.to_excel(writer, sheet_name='df_porcentajes_niveles', index=False)
    df_res_dimTE3.to_excel(writer, sheet_name='df_res_dimTE3', index=False)
    df_resumen.to_excel(writer, sheet_name='df_resumen', index=False)


    ###########

    # Cargar el archivo 'Recomendaciones2.xlsx' en el DataFrame (cambia 'hoja1' al nombre correcto de la hoja si es necesario)
    df_rec = pd.read_excel('Recomendaciones2.xlsx', sheet_name='Hoja1')

    # Crear un DataFrame vacío para almacenar el resultado
    df_reco = pd.DataFrame(columns=['Dimensión', 'Rubro', 'Recomendación'])

    # Recorrer las filas del DataFrame original y extraer las recomendaciones por cada rubro
    for index, row in df_rec.iterrows():
        dimension = row['Dimensión']  # Tomar el valor de la Dimensión en la fila actual
        for i in range(0, len(row) - 1, 2):
            rubro_col = f'Rubro.{i // 2}' if i // 2 > 0 else 'Rubro'
            recomendacion_col = f'Recomendación.{i // 2}' if i // 2 > 0 else 'Recomendación'

            if rubro_col in row and recomendacion_col in row:
                rubro = row[rubro_col]
                recomendacion = row[recomendacion_col]

                # Evitar agregar filas con valores vacíos
                if pd.notna(rubro) and pd.notna(recomendacion):
                    df_reco = pd.concat(
                        [df_reco, pd.DataFrame([[dimension, rubro, recomendacion]], columns=df_reco.columns)],
                        ignore_index=True)


    with pd.ExcelWriter('df_reco.xlsx', engine='xlsxwriter') as writer:
        df_reco.to_excel(writer, sheet_name='df_reco', index=False)

    df_ciiu = pd.read_excel('ciiu.xlsx', sheet_name='ciiu')


    #df_concatenado = df_recomendaciones.groupby('Dimensión')['Recomendación'].apply(lambda x: '\n'.join(x)).reset_index()
    #with pd.ExcelWriter('reco_concatenado.xlsx', engine='xlsxwriter') as writer:
    #    df_concatenado.to_excel(writer, sheet_name='recomendaciones', index=False)

    # Realizar el merge entre `df_ciiu` y `df_recomendaciones` usando la columna `Sección` y `Rubro`
    df_recomendaciones = pd.merge(df_ciiu, df_reco, left_on='Sección', right_on='Rubro')

    print (df_recomendaciones)
    ###############

    import unicodedata

    def normalizar_texto(texto):
        if isinstance(texto, str):
            texto = texto.strip().lower()
            texto = ''.join(
                c for c in unicodedata.normalize('NFD', texto)
                if unicodedata.category(c) != 'Mn'
            )
            return texto
        else:
            return ''

    def obtener_dimm_por_dimension(nombre_dimension):
        nombre_dimension_normalizado = normalizar_texto(nombre_dimension)
        df_dimensiones = dataframes['dimensiones'].copy()  # Ahora accede al DataFrame desde el diccionario
        df_dimensiones['dimension_normalizada'] = df_dimensiones['dimension'].apply(normalizar_texto)
        resultado = df_dimensiones[df_dimensiones['dimension_normalizada'] == nombre_dimension_normalizado]
        if not resultado.empty:
            return resultado.iloc[0]['dimm']
        else:
            print(f"No se encontró el código 'dimm' para la dimensión '{nombre_dimension}'.")
            return None


    ############


    # Asegurarse de que la columna 'CUV' en todos los DataFrames sea de tipo str
    df_res_com['CUV'] = df_res_com['CUV'].astype(str)
    summary_df['CUV'] = summary_df['CUV'].astype(str)
    df_resultados_porcentaje['CUV'] = df_resultados_porcentaje['CUV'].astype(str)
    df_porcentajes_niveles['CUV'] = df_porcentajes_niveles['CUV'].astype(str)
    df_res_dimTE3['CUV'] = df_res_dimTE3['CUV'].astype(str)
    df_resumen['CUV'] = df_resumen['CUV'].astype(str)
    top_glosas['CUV'] = top_glosas['CUV'].astype(str)
    df_recomendaciones['ciiu'] =df_recomendaciones['ciiu'].astype(str)


    # Convertir las fechas al tipo datetime y luego al formato deseado
    df_res_com['Fecha Inicio'] = pd.to_datetime(df_res_com['Fecha Inicio'], errors='coerce').dt.strftime('%d-%m-%Y')
    df_res_com['Fecha Fin'] = pd.to_datetime(df_res_com['Fecha Fin'], errors='coerce').dt.strftime('%d-%m-%Y')


    #########

    import matplotlib.pyplot as plt
    import numpy as np

    # Función para guardar el gráfico principal como imagen temporal
    def guardar_grafico_principal(df, CUV):
        df_filtrado = df[df['CUV'] == CUV]
        if df_filtrado.empty:
            print(f"No se encontraron datos para el CUV {CUV}.")
            return None

        # Pivoteo de los datos para organizar por Dimensión y Nivel, y luego invertir el orden de las filas
        df_pivot = df_filtrado.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[::-1]
        fig, ax = plt.subplots(figsize=(12, 8))

        niveles = ["Bajo", "Medio", "Alto"]
        colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
        posiciones = np.arange(len(df_pivot.index))
        ancho_barra = 0.2

        for i, nivel in enumerate(niveles):
            valores = df_pivot[nivel]
            ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                    label=f"Porcentaje de trabajadores en riesgo {nivel.lower()} (%)", color=colores[nivel])

        ax.axvline(50, color="blue", linestyle="--", linewidth=1)
        ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}", pad=50)
        ax.set_xlabel("Porcentaje")
        ax.set_ylabel("Dimensiones")
        ax.set_xlim(0, 100)
        ax.set_yticks(posiciones + ancho_barra)
        ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

        fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.93), ncol=3)
        plt.subplots_adjust(left=0.3, top=0.85)

        img_path = f"grafico_resultado_CUV_{CUV}.png"
        plt.savefig(img_path, format="png", bbox_inches="tight")
        plt.close()

        return img_path


    #############

    # Filtrar solo los registros con puntaje 1 o 2
    df_filtrado_porcuv = df_res_dimTE3[df_res_dimTE3['Puntaje'].isin([1, 2])]

    # Agrupar por 'CUV' y 'Dimensión' y combinar los valores únicos de 'TE3' en una lista
    resultado_porcuv = df_filtrado_porcuv.groupby(['CUV', 'Dimensión'])['TE3'].unique().reset_index()

    # Convertir las listas de valores de 'TE3' a una cadena separada por "; "
    resultado_porcuv['GES'] = resultado_porcuv['TE3'].apply(lambda x: '; '.join(x))

    # Reemplazar caracteres especiales en toda la columna 'GES'
    resultado_porcuv['GES'] = resultado_porcuv['GES'].str.replace(' | ', '_', regex=False)
    resultado_porcuv['GES'] = resultado_porcuv['GES'].str.replace(':', '_', regex=False)
    resultado_porcuv['GES'] = resultado_porcuv['GES'].str.replace('?', '_', regex=False)

    # Eliminar la columna temporal 'TE3'
    resultado_porcuv = resultado_porcuv[['CUV', 'Dimensión', 'GES']]

    # Mostrar el resultado
    #print(resultado_porcuv)

    ###############


    import matplotlib.pyplot as plt
    import numpy as np

    def guardar_graficos_por_te3(df, CUV):
        """
        Guarda gráficos para cada valor de TE3 dentro del CUV especificado.

        Parámetros:
        df (pd.DataFrame): DataFrame que contiene los datos de TE3, Dimensión, Nivel y Porcentaje.
        CUV (str): El identificador del CUV para filtrar los datos.

        Retorna:
        list: Lista de tuplas con la ruta de cada imagen generada y el valor de TE3 correspondiente.
        """
        df_cuv = df[df['CUV'] == CUV]
        if df_cuv.empty:
            print(f"No se encontraron datos para el CUV {CUV}.")
            return []

        if 'TE3' not in df_cuv.columns:
            print(f"La columna 'TE3' no existe en el DataFrame para CUV {CUV}.")
            return []

        te3_values = df_cuv['TE3'].unique()

        # Verificar si la cardinalidad de TE3 es 1
        if len(te3_values) == 1:
            print(f"Solo hay un valor de TE3 ({te3_values[0]}) para el CUV {CUV}. No se generarán gráficos.")
            return []

        img_paths_te3 = []

        for te3 in te3_values:
            df_te3 = df_cuv[df_cuv['TE3'] == te3]
            df_pivot = df_te3.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[::-1]

            fig, ax = plt.subplots(figsize=(12, 8))

            niveles = ["Bajo", "Medio", "Alto"]
            colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
            posiciones = np.arange(len(df_pivot.index))
            ancho_barra = 0.2

            for i, nivel in enumerate(niveles):
                valores = df_pivot[nivel]
                ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                        label=f"Porcentaje de trabajadores en riesgo {nivel.lower()} (%)", color=colores[nivel])

            ax.axvline(50, color="blue", linestyle="--", linewidth=1)
            ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}, TE3 {te3}", pad=50)
            ax.set_xlabel("Porcentaje")
            ax.set_ylabel("Dimensiones")
            ax.set_xlim(0, 100)
            ax.set_yticks(posiciones + ancho_barra)
            ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

            fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.93), ncol=3)
            plt.subplots_adjust(left=0.3, top=0.85)

            img_path_te3 = f"grafico_resultado_CUV_{CUV}_TE3_{te3}.png".replace('|', '_').replace(':', '_').replace('?', '_')
            plt.savefig(img_path_te3, format="png", bbox_inches="tight")
            plt.close()
            img_paths_te3.append((img_path_te3, te3))  # Guardar como tupla

        return img_paths_te3


    #########
    '''
    # Insertar recomendaciones para cada dimensión en riesgo
                dimensiones_en_riesgo = dimensiones_riesgo_alto + dimensiones_riesgo_medio + dimensiones_riesgo_bajo
                if dimensiones_en_riesgo:
                    doc.add_heading('Recomendaciones por Dimensión en Riesgo', level=3)
                    for dimension in dimensiones_en_riesgo:
                        # Determinar el nivel de riesgo de la dimensión
                        if dimension in dimensiones_riesgo_alto:
                            nivel_riesgo = "Alto"
                        elif dimension in dimensiones_riesgo_medio:
                            nivel_riesgo = "Medio"
                        elif dimension in dimensiones_riesgo_bajo:
                            nivel_riesgo = "Bajo"
                        else:
                            nivel_riesgo = "Desconocido"  # Solo por precaución
    
                        # Agregar el nombre de la dimensión y el nivel de riesgo
                        doc.add_paragraph(f"Dimensión {dimension} (Riesgo {nivel_riesgo})", style="Heading 4")
    
                        # Agregar las recomendaciones para esta dimensión
                        recomendaciones = df_recomendaciones[df_recomendaciones['Dimensión'] == dimension][
                            'Recomendación'].tolist()
                        for rec in recomendaciones:
                            doc.add_paragraph(f"- {rec}")
    
    
    '''





    #############

    from docx.shared import Cm
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches

    #print("ESTAS SON LAS TOP GLOSAS")
    #print(top_glosas.columns)
    #print("FIN TOP GLOSAS")

    def agregar_tabla_ges_por_dimension(doc, df, cuv, df_recomendaciones, df_resultados_porcentaje,df_porcentajes_niveles, top_glosas, datos):
        """
        Agrega una tabla de dimensiones y GES para un CUV específico en el documento de Word.

        Parámetros:
        doc (Document): El objeto del documento de Word.
        df (pd.DataFrame): DataFrame con los datos de dimensiones y GES filtrados.
        cuv (str): El CUV específico para el que se generará la tabla.
        df_recomendaciones (pd.DataFrame): DataFrame con las recomendaciones por dimensión.
        """

        df_revision = df[(df['CUV'] == cuv)]
        if len(df_revision['TE3'].unique()) < 2 :
            noges = 1
            print(f"Solo hay un GES para el CUV {CUV}. No se generarán recomendaciones por GES.")
        else:
            noges = 0



        # Filtrar el DataFrame para el CUV específico y obtener solo las columnas necesarias
        df_cuv = df[(df['CUV'] == cuv) & (df['Puntaje'].isin([1, 2]))]

        # Agrupar por 'Dimensión' y combinar los valores únicos de 'TE3' en una lista separada por "; "
        resultado = df_cuv.groupby('Dimensión')['TE3'].unique().reset_index()
        resultado['GES'] = resultado['TE3'].apply(lambda x: '; '.join(x))
        resultado['GES'] = resultado['GES'].str.replace('|', '_', regex=False).str.replace(':', '_',
                                                                                           regex=False).str.replace('?',
                                                                                                                    '_',
                                                                                                                    regex=False)
        # Asegúrate de que los valores en 'CIIU' son numéricos y trata con NaN o valores no numéricos
        datos['CIIU'] = datos['CIIU'].apply(lambda x: x.split('_')[-1] if isinstance(x, str) else x)

        # Filtra el DataFrame para el CUV específico y convierte el valor a entero si es numérico
        ciiu_valor = datos.loc[datos['CUV'] == cuv, 'CIIU'].copy()

        # Verifica que se ha encontrado el valor y que es un número antes de aplicar la lógica
        if len(ciiu_valor) > 0:
            ciiu_unico = ciiu_valor.iloc[0]

            # Si el valor es un número, aplica la extracción de decenas de miles
            if isinstance(ciiu_unico, str) and ciiu_unico.isdigit():
                ciiu = int(ciiu_unico[:2]) if len(ciiu_unico) > 5 else int(ciiu_unico[:1])
                print("El valor del ciuu para buscar es: " + str(ciiu))
            else:
                print("El valor de CIIU no es numérico.")
        else:
            print("CUV no encontrado en la tabla de datos.")

        # Filtrar por el valor de `ciiu` y `Dimensión` deseados
        #df_resultado = df_merged[(df_merged['ciiu'] == ciiu_valor) & (df_merged['Dimensión'] == dimension_valor)]

        # Mostrar el resultado
        #print(df_resultado[['ciiu', 'Sección', 'Dimensión', 'Recomendación']])



        # Crear la tabla en el documento
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        column_widths = [Inches(0.5),Inches(0.5),Inches(0.5),Inches(7),Inches(0.5),Inches(0.5)]
        # Configurar el ancho de cada columna
        for col_idx, width in enumerate(column_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width

        # Agregar encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dimensión en riesgo'
        hdr_cells[1].text = 'Preguntas clave'
        hdr_cells[2].text = 'Explicación'
        hdr_cells[3].text = 'Medidas propuestas'
        hdr_cells[4].text = 'Fecha monitoreo'
        hdr_cells[5].text = 'Responsable seguimiento'

        # Rellenar la tabla con los datos de 'Dimensión' y 'GES'

        for _, row in resultado.iterrows():
            dim = row['Dimensión']
            ges = row['GES']

            # Obtener las recomendaciones para esta dimensión
            recomendaciones = df_recomendaciones[
                (df_recomendaciones['Dimensión'] == dim) &
                (df_recomendaciones['ciiu'] == str(ciiu))
                ]['Recomendación'].tolist()
            medidas_propuestas = '\n'.join([f"- {rec}" for rec in recomendaciones])

            # Obtener la descripción relacionada desde df_resultados_porcentaje
            descripcion = df_resultados_porcentaje[(df_resultados_porcentaje['Dimensión'] == dim) &
                                                   (df_resultados_porcentaje['CUV'] == cuv)]['Descripción'].values
            descripcion = [desc for desc in descripcion if str(desc).strip() != '']

            descripcion2 = [
                f"{desc} en {ges}"
                for desc in df_porcentajes_niveles[
                    (df_porcentajes_niveles['Dimensión'] == dim) &
                    (df_porcentajes_niveles['CUV'] == cuv) &
                    (df_porcentajes_niveles['TE3'] == ges) &
                    (df_porcentajes_niveles['Descripción'].str.strip() != '')  # Excluir vacíos
                    ]['Descripción'].values
            ]

            if len(descripcion) > 0 and len(descripcion[0]) > 0:
                descripcion_text = descripcion[0] + " para todo el centro de trabajo\n"  # Usa el primer elemento si no es una cadena vacía
            elif len(descripcion) > 1 and len(descripcion[1]) > 0:
                descripcion_text = descripcion[1] + " para todo el centro de trabajo\n" # Usa el segundo elemento si no es una cadena vacía
            elif len(descripcion) > 2 and len(descripcion[2]) > 0:
                descripcion_text = descripcion[2] + " para todo el centro de trabajo\n" # Usa el segundo elemento si no es una cadena vacía
            else:
                descripcion_text = ""

            descripcion2_text = '\n'.join(descripcion2)
            descripcion2_text = str(descripcion2_text).replace("[", "").replace("]", "").replace("'", "")

            te3_values = df_cuv['TE3'].unique()
            # Verificar si la cardinalidad de TE3 es 1
            if noges == 1:
                descripcion2_text = ""
                print(f"Solo hay un GES para el CUV {CUV}. No se generarán recomendaciones por GES.")
            else:
                descripcion2_text = descripcion2_text



            # Depuración: Verificar si hay filas en top_glosas que coincidan con los valores de dim y cuv
            filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dim) & (top_glosas['CUV'] == cuv)]
            #print(f"Dimensión: {dim}, CUV: {cuv}, Filas coincidentes en top_glosas: {len(filtro_glosas)}")

            preguntas = [
                desc for desc in filtro_glosas['Pregunta'].values
            ]

            # Convertir la lista de preguntas a texto
            preguntas_text = '\n'.join(preguntas)

            # Verificar si preguntas se imprime correctamente
            #print("Preguntas seleccionadas:", preguntas_text)

            # Rellenar las celdas de la tabla
            row_cells = table.add_row().cells
            row_cells[0].text = str(str(descripcion_text).lstrip() + "\n"+ str(descripcion2_text).lstrip()).lstrip()
            row_cells[1].text = preguntas_text.lstrip()  # Asignar preguntas_text en lugar de preguntas
            row_cells[2].text = ''  # Espacio para 'Explicación'
            row_cells[3].text = medidas_propuestas  # Medidas propuestas

        # Ajustar el tamaño de fuente de las celdas (opcional)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Ajusta el tamaño de la fuente
    ######


    def establecer_orientacion_apaisada(doc):
        """
        Configura el documento en orientación apaisada (horizontal).
        """
        # Acceder al elemento de configuración de la sección (Sección 1, que es la sección predeterminada)
        section = doc.sections[0]

        # Intercambiar los valores de ancho y alto para hacer la página horizontal
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Configurar márgenes (opcional)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


    ########


    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches


    def generar_contenido_word(datos, estado_riesgo, img_path_principal, img_paths_te3, df_resumen, df_porcentajes_niveles, df_resultados_porcentaje, nombre_archivo):
        """
        Genera el contenido del informe en un documento Word, incluyendo la información general,
        gráficos y recomendaciones de intervenciones.

        Parámetros:
        datos (pd.Series): Serie con los datos de un CUV específico.
        estado_riesgo (str): Nivel de riesgo del CUV.
        img_path_principal (str): Ruta de la imagen principal del gráfico.
        img_paths_te3 (list): Lista de tuplas con rutas de imágenes y valores de TE3.
        df_resumen (pd.DataFrame): DataFrame con datos resumidos de riesgo.
        df_res_dimTE3 (pd.DataFrame): DataFrame con datos de dimensiones por TE3.
        nombre_archivo (str): Ruta donde se guardará el documento Word generado.
        """
        # Crear un nuevo documento
        doc = Document()
        establecer_orientacion_apaisada(doc)
        #section = doc.sections[0]
        #section.page_height = Inches(11)  # 11 pulgadas de alto para Carta
        #section.page_width = Inches(8.5)  # 8.5 pulgadas de ancho para Carta

        # Establecer Calibri como fuente predeterminada para el estilo 'Normal'
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(9)  # Tamaño de fuente opcional; ajusta según prefieras

        # Crear un nuevo estilo llamado 'destacado' con Calibri y tamaño de fuente 12
        destacado = doc.styles.add_style('destacado', 1)  # 1 para párrafos
        destacado_font = destacado.font
        destacado_font.name = 'Calibri'
        destacado_font.size = Pt(12)  # Tamaño de la fuente en puntos

        # Configurar el idioma del documento en español
        lang = doc.styles['Normal'].element
        lang.set(qn('w:lang'), 'es-ES')

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        # Agregar imagen del logo (ajusta la ruta de la imagen a tu ubicación)
        doc.add_picture('ist.jpg', width=Inches(2))  # Ajusta el tamaño según sea necesario
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinear a la derecha
        doc.add_paragraph()

        # Título principal
        titulo = doc.add_heading('INFORME TÉCNICO', level=1)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Subtítulo
        subtitulo = doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2)
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Subtítulo
        subtitulo = doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2)
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
        doc.add_paragraph()


        # Información general
        p = doc.add_paragraph()
        p.add_run('Razón Social: ').bold = True
        p.add_run(f"{datos['Nombre Empresa']}\n")
        p.add_run('RUT: ').bold = True
        p.add_run(f"{datos['RUT Empresa Lugar Geográfico']}\n")
        p.add_run('Nombre del centro de trabajo: ').bold = True
        p.add_run(f"{datos['Nombre Centro de Trabajo']}\n")
        p.add_run('CUV: ').bold = True
        p.add_run(f"{datos['CUV']}\n")
        #p.add_run('Organismo Administrador: ').bold = True
        #p.add_run(f"{datos['Organismo Administrador']}\n")
        p.add_run('CIIU: ').bold = True
        p.add_run(f"{datos['CIIU CT'].split('_')[-1]}\n")
        p.add_run('Fecha de activación del cuestionario: ').bold = True
        p.add_run(f"{datos['Fecha Inicio']}\n")
        p.add_run('Fecha de cierre del cuestionario: ').bold = True
        p.add_run(f"{datos['Fecha Fin']}\n")
        p.add_run('Universo de trabajadores de evaluación: ').bold = True
        p.add_run(f"{datos['Nº Trabajadores CT']}\n")
        #p.add_run('Participación: ').bold = True
        #p.add_run(f"{datos['Participación (%)']}%\n")
        #p.add_run('Razón aplicación cuestionario: ').bold = True
        #p.add_run(f"{datos['Razón aplicación cuestionario']}\n")
        p.paragraph_format.left_indent = Cm(15)

        # Salto de página
        doc.add_page_break()
        doc.add_picture('ist.jpg', width=Inches(1))  # Ajusta el tamaño según sea necesario
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Alinear a la derecha


        # Título de sección
        doc.add_heading('RESULTADOS GENERALES CEAL-SM SUSESO', level=2)

        # Información de riesgo general
        p = doc.add_paragraph()
        p.add_run('Nivel de riesgo: ').bold = True
        p.add_run(f"{estado_riesgo}\n")
        p.style = destacado


        # Insertar imagen principal
        if img_path_principal:
            doc.add_picture(img_path_principal, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Obtener dimensiones en riesgo
        dimensiones_riesgo_altog = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == datos['CUV']) & (
                        df_resultados_porcentaje['Puntaje'] == 2)
            ]['Dimensión'].tolist()

        dimensiones_riesgo_mediog = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == datos['CUV']) & (
                        df_resultados_porcentaje['Puntaje'] == 1)
            ]['Dimensión'].tolist()

        dimensiones_riesgo_bajog = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == datos['CUV'])  & (
                        df_resultados_porcentaje['Puntaje'] == -2)
            ]['Dimensión'].tolist()

        # Dimensiones en riesgo
        p = doc.add_paragraph()
        p.add_run('Dimensiones en riesgo alto: ').bold = True
        p.add_run(f"{', '.join(dimensiones_riesgo_altog) if dimensiones_riesgo_altog else 'Ninguna'}\n")
        p.add_run('Dimensiones en riesgo medio: ').bold = True
        p.add_run(f"{', '.join(dimensiones_riesgo_mediog) if dimensiones_riesgo_mediog else 'Ninguna'}\n")
        p.add_run('Dimensiones en riesgo bajo: ').bold = True
        p.add_run(f"{', '.join(dimensiones_riesgo_bajog) if dimensiones_riesgo_bajog else 'Ninguna'}\n")


        # Salto de página
        doc.add_page_break()
        doc.add_picture('ist.jpg', width=Inches(1))  # Ajusta el tamaño según sea necesario
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Alinear a la derecha

        for img_path_te3, te3 in img_paths_te3:
            if te3 is not None:

                # Subtítulo con el nombre del TE3
                doc.add_heading(f"RESULTADOS POR ÁREA O GES {te3}", level=2)

                # Obtener el valor de riesgo para el TE3 y CUV específico
                riesgo_te3 = df_resumen[(df_resumen['CUV'] == datos['CUV']) & (df_resumen['TE3'] == te3)]['Riesgo']
                if not riesgo_te3.empty:
                    riesgo_te3 = riesgo_te3.values[0]
                else:
                    riesgo_te3 = "Información no disponible"

                p = doc.add_paragraph()
                p.add_run('Nivel de riesgo: ').bold = True
                p.add_run(f"{riesgo_te3}\n")
                p.style = destacado

                # Insertar la imagen del gráfico para este TE3
                doc.add_picture(img_path_te3, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Obtener dimensiones en riesgo alto y medio
                dimensiones_riesgo_alto = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == datos['CUV']) & (df_porcentajes_niveles['TE3'] == te3) & (df_porcentajes_niveles['Puntaje'] == 2)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_medio = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == datos['CUV']) & (df_porcentajes_niveles['TE3'] == te3) & (df_porcentajes_niveles['Puntaje'] == 1)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_bajo = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == datos['CUV']) & (df_porcentajes_niveles['TE3'] == te3) & (df_porcentajes_niveles['Puntaje'] == -2)
                    ]['Dimensión'].tolist()

                # Dimensiones en riesgo alto y medio
                p = doc.add_paragraph()
                p.add_run('Dimensiones en riesgo alto: ').bold = True
                p.add_run(f"{', '.join(dimensiones_riesgo_alto) if dimensiones_riesgo_alto else 'Ninguna'}\n")
                p.add_run('Dimensiones en riesgo medio: ').bold = True
                p.add_run(f"{', '.join(dimensiones_riesgo_medio) if dimensiones_riesgo_medio else 'Ninguna'}\n")
                p.add_run('Dimensiones en riesgo bajo: ').bold = True
                p.add_run(f"{', '.join(dimensiones_riesgo_bajo) if dimensiones_riesgo_bajo else 'Ninguna'}\n")

                # Salto de página
                doc.add_page_break()
                doc.add_picture('ist.jpg', width=Inches(1))  # Ajusta el tamaño según sea necesario
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Alinear a la derecha

        # Llamar a la función para agregar la tabla de GES por dimensión al documento
        doc.add_heading(f"Medidas propuestas para {datos['Nombre Centro de Trabajo']}", level=2)
        agregar_tabla_ges_por_dimension(doc, df_res_dimTE3, cuv, df_recomendaciones, df_resultados_porcentaje,df_porcentajes_niveles, top_glosas, df_res_com)


        # Guardar el documento
        doc.save(nombre_archivo)
        print(f"Informe guardado como: {nombre_archivo}")

    ########

    import os
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, CUV, df_resumen,
                        df_res_dimTE3):
        datos = df_res_com[df_res_com['CUV'] == CUV]
        estado = summary_df[summary_df['CUV'] == CUV]

        if datos.empty:
            print(f"No se encontró el CUV {CUV} en df_res_com.")
            return
        if estado.empty:
            print(f"No se encontró el CUV {CUV} en summary_df.")
            return

        row = datos.iloc[0]
        estado_riesgo = estado['Riesgo'].values[0]

        # Guardar el gráfico principal de resultados usando df_resultados_porcentaje
        img_path_principal = guardar_grafico_principal(df_resultados_porcentaje, CUV)
        if not img_path_principal:
            return

        # Guardar gráficos para cada TE3 dentro del CUV usando df_porcentajes_niveles
        img_paths_te3 = guardar_graficos_por_te3(df_porcentajes_niveles, CUV)

        # Generar el contenido en el documento Word usando python-docx
        nombre_archivo = os.path.join(output_archivos, f"{row['CUV']}-{row['Nombre Centro de Trabajo']}.docx")
        generar_contenido_word(row, estado_riesgo, img_path_principal, img_paths_te3, df_resumen,
                               df_res_dimTE3, df_resultados_porcentaje, nombre_archivo)

        print(f"Informe generado: {nombre_archivo}")

        # Eliminar los archivos temporales de los gráficos
        os.remove(img_path_principal)
        for img_path, _ in img_paths_te3:
            os.remove(img_path)

    #############

    # Generar informes para cada CUV en summary_df
    for cuv in summary_df['CUV'].unique():
        generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, cuv, df_resumen, df_res_dimTE3)
