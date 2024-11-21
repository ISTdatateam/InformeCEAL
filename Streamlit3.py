# Importar las bibliotecas necesarias
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import unicodedata
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuración inicial de la aplicación
st.set_page_config(page_title="Generador de Informes de Riesgos Psicosociales", layout="wide")
st.title("Generador de Informes de Riesgos Psicosociales")
st.write("""
Esta aplicación permite generar informes técnicos basados en datos de riesgos psicosociales.
Por favor, cargue los archivos necesarios y siga las instrucciones.
""")

# Sección 1: Carga de archivos
st.header("1. Carga de archivos")

# Cargar el archivo Excel con múltiples hojas
uploaded_file_combined = st.file_uploader("Selecciona el archivo 'combined_output.xlsx'", type="xlsx")

# Cargar los archivos de recomendaciones y códigos CIIU
# uploaded_file_rec = st.file_uploader("Selecciona el archivo 'Recomendaciones2.xlsx'", type="xlsx")
# uploaded_file_ciiu = st.file_uploader("Selecciona el archivo 'ciiu.xlsx'", type="xlsx")

# Cargar el archivo 'resultados.xlsx'
# uploaded_file_resultados = st.file_uploader("Selecciona el archivo 'resultados.xlsx'", type="xlsx")

#precargas
uploaded_file_rec = 'Recomendaciones2.xlsx'
uploaded_file_ciiu = 'ciiu.xlsx'
uploaded_file_resultados = 'resultados.xlsx'


# Verificar que todos los archivos se hayan cargado
if (uploaded_file_combined is not None and
        uploaded_file_rec is not None and
        uploaded_file_ciiu is not None and
        uploaded_file_resultados is not None):

    # Leer las hojas necesarias desde 'combined_output.xlsx'
    try:
        df_res_com = pd.read_excel(uploaded_file_combined, sheet_name='basecompleta')
        summary_df = pd.read_excel(uploaded_file_combined, sheet_name='Summary')
        df_resultados_porcentaje = pd.read_excel(uploaded_file_combined, sheet_name='resultado')
        df_porcentajes_niveles = pd.read_excel(uploaded_file_combined, sheet_name='df_porcentajes_niveles')
        df_res_dimTE3 = pd.read_excel(uploaded_file_combined, sheet_name='df_res_dimTE3')
        df_resumen = pd.read_excel(uploaded_file_combined, sheet_name='df_resumen')
        top_glosas = pd.read_excel(uploaded_file_combined, sheet_name='top_glosas')

        st.success("Archivo 'combined_output.xlsx' cargado exitosamente.")
    except Exception as e:
        st.error(f"Error al leer las hojas de 'combined_output.xlsx': {e}")

    # Leer los archivos de recomendaciones y CIIU
    try:
        df_rec = pd.read_excel(uploaded_file_rec, sheet_name='Hoja1')  # Cambia 'Hoja1' si es necesario
        df_ciiu = pd.read_excel(uploaded_file_ciiu, sheet_name='ciiu')  # Cambia 'ciiu' si es necesario

        st.success("Archivos 'Recomendaciones2.xlsx' y 'ciiu.xlsx' cargados exitosamente.")
    except Exception as e:
        st.error(f"Error al leer los archivos de recomendaciones y CIIU: {e}")

    # Leer y procesar 'resultados.xlsx'
    try:
        df_resultados = pd.read_excel(uploaded_file_resultados, sheet_name='Datos', usecols=['CUV', 'Folio'])
        df_res_com_resultados = pd.read_excel(uploaded_file_resultados, sheet_name='Datos')

        # Convertir 'CUV' a numérico (int64), manejando errores
        df_resultados['CUV'] = pd.to_numeric(df_resultados['CUV'], errors='coerce').astype(
            'Int64')  # Usa 'Int64' para permitir NaNs

        st.success("Archivo 'resultados.xlsx' cargado y procesado exitosamente.")
    except Exception as e:
        st.error(f"Error al leer y procesar 'resultados.xlsx': {e}")

    # Procesamiento de df_rec para crear df_reco
    try:
        df_reco = pd.DataFrame(columns=['Dimensión', 'Rubro', 'Recomendación'])

        for index, row in df_rec.iterrows():
            dimension = row['Dimensión']
            for i in range(0, len(row) - 1, 2):
                rubro_col = f'Rubro.{i // 2}' if i // 2 > 0 else 'Rubro'
                recomendacion_col = f'Recomendación.{i // 2}' if i // 2 > 0 else 'Recomendación'

                if rubro_col in row and recomendacion_col in row:
                    rubro = row[rubro_col]
                    recomendacion = row[recomendacion_col]

                    if pd.notna(rubro) and pd.notna(recomendacion):
                        df_reco = pd.concat(
                            [df_reco, pd.DataFrame([[dimension, rubro, recomendacion]], columns=df_reco.columns)],
                            ignore_index=True
                        )
        st.success("Recomendaciones procesadas correctamente.")
    except Exception as e:
        st.error(f"Error al procesar las recomendaciones: {e}")

    # Realizar el merge entre df_ciiu y df_reco
    try:
        df_recomendaciones = pd.merge(df_ciiu, df_reco, left_on='Sección', right_on='Rubro', how='left')
        st.success("Merge entre CIIU y Recomendaciones realizado exitosamente.")
    except Exception as e:
        st.error(f"Error al realizar el merge entre CIIU y Recomendaciones: {e}")

    # Asegurar que las columnas 'CUV' y 'ciiu' sean de tipo str
    try:
        df_res_com['CUV'] = df_res_com['CUV'].astype(str)
        summary_df['CUV'] = summary_df['CUV'].astype(str)
        df_resultados_porcentaje['CUV'] = df_resultados_porcentaje['CUV'].astype(str)
        df_porcentajes_niveles['CUV'] = df_porcentajes_niveles['CUV'].astype(str)
        df_res_dimTE3['CUV'] = df_res_dimTE3['CUV'].astype(str)
        df_resumen['CUV'] = df_resumen['CUV'].astype(str)
        top_glosas['CUV'] = top_glosas['CUV'].astype(str)
        df_recomendaciones['ciiu'] = df_recomendaciones['ciiu'].astype(str)

        st.success("Tipos de datos ajustados correctamente.")
    except Exception as e:
        st.error(f"Error al ajustar los tipos de datos: {e}")

    # Convertir las fechas al tipo datetime y luego al formato deseado
    columna_fecha_inicio = 'Fecha Inicio'  # Ajusta si el nombre es diferente

    if columna_fecha_inicio in df_res_com.columns:
        try:
            df_res_com['Fecha Inicio'] = pd.to_datetime(df_res_com[columna_fecha_inicio], errors='coerce').dt.strftime(
                '%d-%m-%Y')
            st.success("Columna 'Fecha Inicio' procesada correctamente.")
        except Exception as e:
            st.error(f"Error al procesar la columna 'Fecha Inicio': {e}")
    else:
        st.error(
            f"La columna '{columna_fecha_inicio}' no se encontró en el DataFrame 'df_res_com'. Por favor, verifica el nombre de la columna.")
        st.write("Columnas disponibles en 'df_res_com':", df_res_com.columns.tolist())

    # Realizar el merge entre df_res_com y df_resultados
    try:
        df_res_com = pd.merge(df_res_com, df_resultados, on='CUV', how='left')
        st.success("Merge entre 'df_res_com' y 'df_resultados' realizado exitosamente.")
    except Exception as e:
        st.error(f"Error al realizar el merge entre 'df_res_com' y 'df_resultados': {e}")


    # Definir funciones auxiliares
    def normalizar_texto(texto):
        if isinstance(texto, str):
            texto = texto.strip().lower()
            texto = ''.join(
                c for c in unicodedata.normalize('NFD', texto)
                if unicodedata.category(c) != 'Mn'
            )
            return texto
        else:
            return ''


    def obtener_dimm_por_dimension(nombre_dimension):
        nombre_dimension_normalizado = normalizar_texto(nombre_dimension)
        df_dimensiones = df_res_dimTE3[['Dimensión', 'dimm']].drop_duplicates()
        df_dimensiones['dimension_normalizada'] = df_dimensiones['Dimensión'].apply(normalizar_texto)
        resultado = df_dimensiones[df_dimensiones['dimension_normalizada'] == nombre_dimension_normalizado]
        if not resultado.empty:
            return resultado.iloc[0]['dimm']
        else:
            st.warning(f"No se encontró el código 'dimm' para la dimensión '{nombre_dimension}'.")
            return None


    # Definir la función para procesar los datos
    def procesar_datos_tabla(df_porcentajes_niveles, df_recomendaciones, top_glosas, selected_cuv, te3, ciiu_valor):
        # [Definición completa de la función según el paso 2.1]
        # ... [Código de la función procesar_datos_tabla] ...
        # Reemplaza este comentario con la definición completa
        pass  # Remover esta línea y pegar la definición completa


    # Definir la función para mostrar la tabla en la web
    def mostrar_tabla_en_web(tabla_te3):
        if not tabla_te3.empty:
            st.markdown("### Análisis de Dimensiones en Riesgo para este TE3")
            st.table(tabla_te3)  # O usa st.dataframe(tabla_te3) para interactividad
        else:
            st.info("No hay dimensiones en riesgo para este TE3.")


    # Definir la función para agregar la tabla al documento de Word
    def agregar_tabla_word(doc, tabla_te3):
        if tabla_te3.empty:
            doc.add_paragraph("No hay dimensiones en riesgo para este TE3.")
            return

        # Crear la tabla en el documento de Word
        table = doc.add_table(rows=1, cols=len(tabla_te3.columns))
        table.style = 'Table Grid'

        # Agregar encabezados
        hdr_cells = table.rows[0].cells
        for idx, column in enumerate(tabla_te3.columns):
            hdr_cells[idx].text = column

        # Rellenar la tabla con los datos
        for _, row in tabla_te3.iterrows():
            row_cells = table.add_row().cells
            for idx, item in enumerate(row):
                row_cells[idx].text = str(item)

        # Opcional: Ajustar el tamaño de fuente o estilos adicionales
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Ajusta el tamaño de la fuente según sea necesario


    # Sección 2: Selección de CUV para generar el informe
    st.header("2. Seleccionar CUV para generar el informe")

    cuvs_disponibles = summary_df['CUV'].unique()
    selected_cuv = st.selectbox("Selecciona un CUV", cuvs_disponibles)

    # Filtrar los datos para el CUV seleccionado
    datos = df_res_com[df_res_com['CUV'] == selected_cuv]
    estado = summary_df[summary_df['CUV'] == selected_cuv]

    if datos.empty or estado.empty:
        st.error(f"No se encontraron datos para el CUV {selected_cuv}.")
    else:
        # Obtener la primera fila de datos
        datos = datos.iloc[0]
        estado_riesgo = estado['Riesgo'].values[0]

        # Extraer 'ciiu_valor' del DataFrame 'datos'
        if 'CIIU' in datos:
            ciiu_valor = datos['CIIU']
        else:
            ciiu_valor = None
            st.error("La columna 'CIIU' no está presente en los datos.")

        # Sección 3: Generación y visualización de gráficos
        st.header("3. Gráfico principal de resultados")


        def generar_grafico_principal(df, CUV):
            """
            Genera el gráfico principal de resultados.

            Parámetros:
            - df (pd.DataFrame): DataFrame con los resultados porcentuales.
            - CUV (str): CUV seleccionado.

            Retorna:
            - fig (matplotlib.figure.Figure): Figura del gráfico generado.
            """
            df_filtrado = df[df['CUV'] == CUV]
            if df_filtrado.empty:
                st.warning(f"No se encontraron datos para el CUV {CUV}.")
                return None

            try:
                df_pivot = df_filtrado.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[
                           ::-1]
            except Exception as e:
                st.error(f"Error al pivotear los datos: {e}")
                return None

            fig, ax = plt.subplots(figsize=(12, 8))

            niveles = ["Bajo", "Medio", "Alto"]
            colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
            posiciones = np.arange(len(df_pivot.index))
            ancho_barra = 0.2

            for i, nivel in enumerate(niveles):
                if nivel in df_pivot.columns:
                    valores = df_pivot[nivel]
                    ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                            label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
                else:
                    st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot.")

            ax.axvline(50, color="blue", linestyle="--", linewidth=1)
            ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}", pad=20)
            ax.set_xlabel("Porcentaje")
            ax.set_ylabel("Dimensiones")
            ax.set_xlim(0, 100)
            ax.set_yticks(posiciones + ancho_barra)
            ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

            fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 1.05), ncol=3)
            plt.tight_layout()

            return fig


        # Generar y mostrar el gráfico principal
        fig_principal = generar_grafico_principal(df_resultados_porcentaje, selected_cuv)
        if fig_principal:
            st.pyplot(fig_principal)
        else:
            st.warning("No se pudo generar el gráfico principal.")

        # ---- NUEVA SECCIÓN PARA MOSTRAR DIMENSIONES EN RIESGO ----
        st.header("4. Dimensiones en Riesgo")

        # Obtener dimensiones en riesgo
        dimensiones_riesgo_alto = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == 2)
            ]['Dimensión'].tolist()

        dimensiones_riesgo_medio = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == 1)
            ]['Dimensión'].tolist()

        dimensiones_riesgo_bajo = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == -2)
            ]['Dimensión'].tolist()

        # Mostrar dimensiones en riesgo en la web
        st.subheader("Dimensiones en riesgo")
        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("**Alto:**")
            if dimensiones_riesgo_alto:
                st.write(", ".join(dimensiones_riesgo_alto))
            else:
                st.write("Ninguna")

        with col2:
            st.markdown("**Medio:**")
            if dimensiones_riesgo_medio:
                st.write(", ".join(dimensiones_riesgo_medio))
            else:
                st.write("Ninguna")

        with col3:
            st.markdown("**Bajo:**")
            if dimensiones_riesgo_bajo:
                st.write(", ".join(dimensiones_riesgo_bajo))
            else:
                st.write("Ninguna")
        # ------------------------------------------------------------

        # Sección 5: Gráficos por TE3
        st.header("5. Gráficos por TE3")

        figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles, selected_cuv)

        if figs_te3:
            for fig_te3, te3 in figs_te3:
                st.subheader(f"Gráfico para TE3: {te3}")
                st.pyplot(fig_te3)

                dimensiones_riesgo_alto = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == 2)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_medio = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == 1)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_bajo = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == -2)
                    ]['Dimensión'].tolist()

                # Crear tres columnas
                col1, col2, col3 = st.columns(3)

                with col1:
                    st.markdown("**Alto:**")
                    if dimensiones_riesgo_alto:
                        st.write(", ".join(dimensiones_riesgo_alto))
                    else:
                        st.write("Ninguna")

                with col2:
                    st.markdown("**Medio:**")
                    if dimensiones_riesgo_medio:
                        st.write(", ".join(dimensiones_riesgo_medio))
                    else:
                        st.write("Ninguna")

                with col3:
                    st.markdown("**Bajo:**")
                    if dimensiones_riesgo_bajo:
                        st.write(", ".join(dimensiones_riesgo_bajo))
                    else:
                        st.write("Ninguna")







        else:
            st.info("No se generaron gráficos adicionales por TE3.")

        # Opcional: Botón para generar y descargar el informe en Word
        if st.button("Generar Informe Word"):
            # Crear un nuevo documento de Word
            doc = Document()
            doc.add_heading("Informe de Riesgos Psicosociales", 0)

            # ... [Agregar otras secciones al documento] ...

            # Agregar tablas por TE3
            if figs_te3:
                for _, te3 in figs_te3:
                    # Procesar los datos para la tabla
                    tabla_te3 = procesar_datos_tabla(
                        df_porcentajes_niveles=df_porcentajes_niveles,
                        df_recomendaciones=df_recomendaciones,
                        top_glosas=top_glosas,
                        selected_cuv=selected_cuv,
                        te3=te3,
                        ciiu_valor=ciiu_valor
                    )

                    # Agregar la tabla al documento
                    agregar_tabla_word(doc, tabla_te3)

            # Guardar el documento en un objeto BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Descargar el documento
            st.download_button(
                label="Descargar Informe Word",
                data=buffer,
                file_name="Informe_Riesgos_Psicosociales.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
