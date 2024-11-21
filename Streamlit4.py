# app.py

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import pyodbc
import threading
import unicodedata

# ===============================
# Funciones de data_processing.py
# ===============================

def cargar_datos(uploaded_files):
    """
    Carga y retorna los datos desde los archivos subidos.

    Parámetros:
    - uploaded_files (dict): Diccionario con los archivos subidos.

    Retorna:
    - dict: Diccionario con DataFrames cargados.
    """
    try:
        data_dict = {}
        data_dict['df_combined'] = pd.read_excel(uploaded_files['combined'], sheet_name=None)
        data_dict['df_rec'] = pd.read_excel(uploaded_files['recomendaciones'])
        data_dict['df_ciiu'] = pd.read_excel(uploaded_files['ciiu'])
        data_dict['df_res_com'] = pd.read_excel(uploaded_files['resultados'])
        data_dict['summary_df'] = data_dict['df_combined']['Summary']  # Asumiendo que hay una hoja 'Summary'
        data_dict['df_resultados_porcentaje'] = data_dict['df_combined']['ResultadosPorcentaje']
        data_dict['df_porcentajes_niveles'] = data_dict['df_combined']['PorcentajesNiveles']
        data_dict['df_resumen'] = data_dict['df_combined']['Resumen']
        data_dict['df_res_dimTE3'] = data_dict['df_combined']['ResDimTE3']
        data_dict['top_glosas'] = data_dict['df_combined'].get('TopGlosas', pd.DataFrame())
        return data_dict
    except Exception as e:
        st.error(f"Error al cargar los datos: {e}")
        return None

def procesar_recomendaciones(df_rec, df_ciiu):
    """
    Procesa las recomendaciones combinando información de recomendaciones y CIIU.

    Parámetros:
    - df_rec (pd.DataFrame): DataFrame con recomendaciones.
    - df_ciiu (pd.DataFrame): DataFrame con códigos CIIU.

    Retorna:
    - pd.DataFrame: DataFrame procesado con recomendaciones.
    """
    try:
        df = pd.merge(df_rec, df_ciiu, on='CIIU', how='left')
        return df
    except Exception as e:
        st.error(f"Error al procesar recomendaciones: {e}")
        return pd.DataFrame()

def convertir_columnas(data_dict):
    """
    Convierte ciertas columnas a tipo string.

    Parámetros:
    - data_dict (dict): Diccionario con DataFrames.

    Retorna:
    - dict: Diccionario con DataFrames actualizados.
    """
    try:
        for key, df in data_dict.items():
            if isinstance(df, pd.DataFrame):
                df = df.applymap(lambda x: str(x) if isinstance(x, (int, float)) else x)
                data_dict[key] = df
        return data_dict
    except Exception as e:
        st.error(f"Error al convertir columnas: {e}")
        return data_dict

def convertir_fechas(df_res_com):
    """
    Convierte columnas de fechas a formato datetime.

    Parámetros:
    - df_res_com (pd.DataFrame): DataFrame con columnas de fechas.

    Retorna:
    - pd.DataFrame: DataFrame con columnas de fechas convertidas.
    """
    try:
        fecha_cols = ['Fecha Inicio', 'Fecha Fin']
        for col in fecha_cols:
            if col in df_res_com.columns:
                df_res_com[col] = pd.to_datetime(df_res_com[col], errors='coerce')
        return df_res_com
    except Exception as e:
        st.error(f"Error al convertir fechas: {e}")
        return df_res_com

# ===============================
# Funciones de utils.py
# ===============================

def obtener_dimm_por_dimension(nombre_dimension, df_dimensiones):
    nombre_dimension_normalizado = normalizar_texto(nombre_dimension)
    resultado = df_dimensiones[df_dimensiones['dimension_normalizada'] == nombre_dimension_normalizado]
    if not resultado.empty:
        return resultado.iloc[0]['dimm']
    else:
        st.warning(f"No se encontró el código 'dimm' para la dimensión '{nombre_dimension}'.")
        return None

def normalizar_texto(texto):
    if isinstance(texto, str):
        texto = texto.strip().lower()
        texto = ''.join(
            c for c in unicodedata.normalize('NFD', texto)
            if unicodedata.category(c) != 'Mn'
        )
        return texto
    else:
        return ''

# ===============================
# Funciones de graph_generation.py
# ===============================

def generar_grafico_principal(df, CUV):
    df_filtrado = df[df['CUV'] == CUV]
    if df_filtrado.empty:
        st.warning(f"No se encontraron datos para el CUV {CUV}.")
        return None

    try:
        df_pivot = df_filtrado.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[::-1]
    except Exception as e:
        st.error(f"Error al pivotear los datos: {e}")
        return None

    fig, ax = plt.subplots(figsize=(12, 8))

    niveles = ["Bajo", "Medio", "Alto"]
    colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
    posiciones = np.arange(len(df_pivot.index))
    ancho_barra = 0.2

    for i, nivel in enumerate(niveles):
        if nivel in df_pivot.columns:
            valores = df_pivot[nivel]
            ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                    label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
        else:
            st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot.")

    ax.axvline(50, color="blue", linestyle="--", linewidth=1)
    ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}", pad=50)
    ax.set_xlabel("Porcentaje")
    ax.set_ylabel("Dimensiones")
    ax.set_xlim(0, 100)
    ax.set_yticks(posiciones + ancho_barra)
    ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

    fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.6, 0.96), ncol=3)
    plt.subplots_adjust(left=0.3, top=0.85)
    plt.tight_layout()

    return fig

def generar_graficos_por_te3(df, CUV):
    df_cuv = df[df['CUV'] == CUV]
    if df_cuv.empty:
        st.warning(f"No se encontraron datos para el CUV {CUV}.")
        return []

    if 'TE3' not in df_cuv.columns:
        st.warning(f"La columna 'TE3' no existe en el DataFrame para CUV {CUV}.")
        return []

    te3_values = df_cuv['TE3'].unique()

    if len(te3_values) <= 1:
        st.info(f"Solo hay un valor de TE3 para el CUV {CUV}. No se generarán gráficos adicionales.")
        return []

    figs_te3 = []

    for te3 in te3_values:
        df_te3 = df_cuv[df_cuv['TE3'] == te3]
        if df_te3.empty:
            continue
        try:
            df_pivot = df_te3.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[::-1]
        except Exception as e:
            st.error(f"Error al pivotear los datos para TE3 {te3}: {e}")
            continue

        fig, ax = plt.subplots(figsize=(12, 8))

        niveles = ["Bajo", "Medio", "Alto"]
        colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
        posiciones = np.arange(len(df_pivot.index))
        ancho_barra = 0.2

        for i, nivel in enumerate(niveles):
            if nivel in df_pivot.columns:
                valores = df_pivot[nivel]
                ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                        label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
            else:
                st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot para TE3 {te3}.")

        ax.axvline(50, color="blue", linestyle="--", linewidth=1)
        ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}, TE3 {te3}", pad=50)
        ax.set_xlabel("Porcentaje")
        ax.set_ylabel("Dimensiones")
        ax.set_xlim(0, 100)
        ax.set_yticks(posiciones + ancho_barra)
        ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

        fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.96), ncol=3)
        plt.tight_layout()

        figs_te3.append((fig, te3))

    return figs_te3

# ===============================
# Funciones de report_generation.py
# ===============================

def establecer_orientacion_apaisada(doc):
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

def agregar_tabla_ges_por_dimension(doc, dimensiones, cuv, dimensiones_recomendaciones,
                                    dimensiones_resultados_porcentaje, dimensiones_porcentajes_niveles, top_glosas,
                                    datos):
    """
    Agrega una tabla de dimensiones y GES para un CUV específico en el documento de Word.

    Parámetros:
    - doc (Document): El objeto del documento de Word.
    - dimensiones (list): Lista de diccionarios con información de dimensiones.
    - cuv (str): El CUV específico para el que se generará la tabla.
    - dimensiones_recomendaciones (pd.DataFrame): DataFrame con las recomendaciones por dimensión.
    - dimensiones_resultados_porcentaje (pd.DataFrame): DataFrame con las descripciones por dimensión y CUV.
    - dimensiones_porcentajes_niveles (pd.DataFrame): DataFrame con porcentajes y niveles por dimensión y GES.
    - top_glosas (pd.DataFrame): DataFrame con preguntas clave por dimensión y CUV.
    - datos (pd.DataFrame): DataFrame con información adicional, incluyendo CIIU.
    """

    # Crear la tabla en el documento
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    column_widths = [Inches(1), Inches(2), Inches(1), Inches(3), Inches(1), Inches(1)]

    # Configurar el ancho de cada columna
    for col_idx, width in enumerate(column_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width

    # Agregar encabezados
    hdr_cells = table.rows[0].cells
    headers = ['Dimensión en riesgo', 'Preguntas clave', 'Explicación',
               'Medidas propuestas', 'Fecha monitoreo', 'Responsable seguimiento']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Procesar cada dimensión en la lista
    for dim in dimensiones:
        dimension = dim['dimension']
        preguntas = dim['preguntas']
        medidas = dim['medidas']
        riesgo = dim.get('riesgo', '')
        area = dim.get('area', '')

        # Obtener las recomendaciones para esta dimensión y CIIU
        ciiu_valor = datos.loc[datos['CUV'] == cuv, 'CIIU'].iloc[0].split('_')[-1] if not datos.empty else ''
        recomendaciones = dimensiones_recomendaciones[
            (dimensiones_recomendaciones['Dimensión'] == dimension) &
            (dimensiones_recomendaciones['ciiu'] == str(ciiu_valor))
        ]['Recomendación'].tolist()
        medidas_propuestas = '\n'.join([f"- {rec}" for rec in recomendaciones]) if recomendaciones else 'Ninguna'

        # Obtener las preguntas clave desde top_glosas
        filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dimension) & (top_glosas['CUV'] == cuv)]
        preguntas_text = '\n'.join(filtro_glosas['Pregunta'].tolist()) if not filtro_glosas.empty else 'Ninguna'

        # Obtener descripciones relacionadas
        descripcion = dimensiones_resultados_porcentaje[
            (dimensiones_resultados_porcentaje['Dimensión'] == dimension) &
            (dimensiones_resultados_porcentaje['CUV'] == cuv)
        ]['Descripción'].dropna().tolist()
        descripcion_text = descripcion[0] + " para todo el centro de trabajo\n" if descripcion else ""

        # Obtener descripciones adicionales por GES
        descripcion2 = []
        for medida in medidas:
            ges = medida.get('fecha_monitoreo', '').split('(')[-1].replace(')', '') if '(' in medida.get(
                'fecha_monitoreo', '') else ''
            desc = dimensiones_porcentajes_niveles[
                (dimensiones_porcentajes_niveles['Dimensión'] == dimension) &
                (dimensiones_porcentajes_niveles['CUV'] == cuv) &
                (dimensiones_porcentajes_niveles['TE3'] == ges) &
                (dimensiones_porcentajes_niveles['Descripción'].str.strip() != '')
            ]['Descripción'].tolist()
            if desc:
                descripcion2.append(f"{desc[0]} en {ges}")

        descripcion2_text = '\n'.join(descripcion2).replace("[", "").replace("]", "").replace("'", "") if descripcion2 else ""

        # Rellenar las celdas de la tabla
        row_cells = table.add_row().cells
        row_cells[0].text = f"{descripcion_text}\n{descripcion2_text}".strip()
        row_cells[1].text = preguntas_text.strip()
        row_cells[2].text = f"{riesgo} en {area}" if riesgo and area else ""
        row_cells[3].text = medidas_propuestas.strip()
        row_cells[4].text = ''  # Fecha monitoreo (ya incluida en medidas)
        row_cells[5].text = ''  # Responsable seguimiento (ya incluido en medidas)

        # Opcional: Ajustar el tamaño de fuente de las celdas
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

def generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, CUV, df_resumen,
                    df_res_dimTE3, generar_grafico_principal, generar_graficos_por_te3, dimensiones, cuv,
                    df_recomendaciones, top_glosas, datos_df):
    """
    Genera el informe en Word para un CUV específico.

    Parámetros:
    - df_res_com (pd.DataFrame): DataFrame con resultados de la evaluación por CUV.
    - summary_df (pd.DataFrame): DataFrame con resumen de los CUV.
    - df_resultados_porcentaje (pd.DataFrame): DataFrame con resultados porcentuales.
    - df_porcentajes_niveles (pd.DataFrame): DataFrame con porcentajes y niveles por dimensión y GES.
    - CUV (str): Código único de evaluación.
    - df_resumen (pd.DataFrame): DataFrame con resumen de resultados.
    - df_res_dimTE3 (pd.DataFrame): DataFrame con resultados por dimensión y TE3.
    - generar_grafico_principal (function): Función para generar el gráfico principal.
    - generar_graficos_por_te3 (function): Función para generar gráficos por TE3.
    - dimensiones (list): Lista de diccionarios con información de dimensiones.
    - cuv (str): Código único de evaluación (redundante, podría eliminarse).
    - df_recomendaciones (pd.DataFrame): DataFrame con recomendaciones por dimensión.
    - top_glosas (pd.DataFrame): DataFrame con preguntas clave por dimensión y CUV.
    - datos_df (pd.DataFrame): DataFrame con información adicional, incluyendo CIIU.

    Retorna:
    - docx_buffer (BytesIO): Buffer con el documento de Word generado.
    """
    datos = df_res_com[df_res_com['CUV'] == CUV]
    estado = summary_df[summary_df['CUV'] == CUV]

    if datos.empty:
        st.error(f"No se encontró el CUV {CUV} en df_res_com.")
        return None
    if estado.empty:
        st.error(f"No se encontró el CUV {CUV} en summary_df.")
        return None

    row = datos.iloc[0]
    estado_riesgo = estado['Riesgo'].values[0]

    # Generar el gráfico principal
    fig_principal = generar_grafico_principal(df_resultados_porcentaje, CUV)
    if not fig_principal:
        st.warning("No se pudo generar el gráfico principal.")
        return None

    # Generar gráficos por TE3
    figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles, CUV)

    # Obtener las dimensiones y medidas desde el parámetro
    if dimensiones:
        dimensiones = dimensiones
    else:
        dimensiones = []

    # Obtener el DataFrame 'datos' necesario para CIIU
    datos_df = df_res_com[df_res_com['CUV'] == CUV]

    # Generar el contenido en el documento Word usando python-docx
    doc = Document()
    establecer_orientacion_apaisada(doc)

    # Configurar estilos y agregar contenido al documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9)

    # Título principal
    titulo = doc.add_heading('INFORME TÉCNICO', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulos
    subtitulo = doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Información general
    p = doc.add_paragraph()
    p.add_run('Razón Social: ').bold = True
    p.add_run(f"{row['Nombre Empresa']}\n")
    p.add_run('RUT: ').bold = True
    p.add_run(f"{row['RUT Empresa Lugar Geográfico']}\n")
    p.add_run('Nombre del centro de trabajo: ').bold = True
    p.add_run(f"{row['Nombre Centro de Trabajo']}\n")
    p.add_run('CUV: ').bold = True
    p.add_run(f"{row['CUV']}\n")
    p.add_run('CIIU: ').bold = True
    p.add_run(f"{row['CIIU CT'].split('_')[-1]}\n")
    p.add_run('Fecha de activación del cuestionario: ').bold = True
    p.add_run(f"{row['Fecha Inicio']}\n")
    p.add_run('Fecha de cierre del cuestionario: ').bold = True
    p.add_run(f"{row['Fecha Fin']}\n")
    p.add_run('Universo de trabajadores de evaluación: ').bold = True
    p.add_run(f"{row['Nº Trabajadores CT']}\n")
    p.paragraph_format.left_indent = Cm(1)

    # Salto de página
    doc.add_page_break()

    # Título de sección
    doc.add_heading('RESULTADOS GENERALES', level=2)

    # Información de riesgo general
    p = doc.add_paragraph()
    p.add_run('Nivel de riesgo: ').bold = True
    p.add_run(f"{estado_riesgo}\n")
    p.style.font.size = Pt(12)

    # Insertar imagen del gráfico principal
    if fig_principal:
        img_buffer = BytesIO()
        fig_principal.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        img_buffer.close()
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Obtener dimensiones en riesgo
    dimensiones_riesgo_alto = df_resultados_porcentaje[
        (df_resultados_porcentaje['CUV'] == cuv) & (df_resultados_porcentaje['Puntaje'] == 2)
        ]['Dimensión'].tolist()

    dimensiones_riesgo_medio = df_resultados_porcentaje[
        (df_resultados_porcentaje['CUV'] == cuv) & (df_resultados_porcentaje['Puntaje'] == 1)
        ]['Dimensión'].tolist()

    dimensiones_riesgo_bajo = df_resultados_porcentaje[
        (df_resultados_porcentaje['CUV'] == cuv) & (df_resultados_porcentaje['Puntaje'] == -2)
        ]['Dimensión'].tolist()

    # Dimensiones en riesgo
    p = doc.add_paragraph()
    p.add_run('Dimensiones en riesgo alto: ').bold = True
    p.add_run(f"{', '.join(dimensiones_riesgo_alto) if dimensiones_riesgo_alto else 'Ninguna'}\n")
    p.add_run('Dimensiones en riesgo medio: ').bold = True
    p.add_run(f"{', '.join(dimensiones_riesgo_medio) if dimensiones_riesgo_medio else 'Ninguna'}\n")
    p.add_run('Dimensiones en riesgo bajo: ').bold = True
    p.add_run(f"{', '.join(dimensiones_riesgo_bajo) if dimensiones_riesgo_bajo else 'Ninguna'}\n")

    # Salto de página
    doc.add_page_break()

    # Agregar gráficos por TE3
    for fig_te3, te3 in figs_te3:
        doc.add_heading(f"RESULTADOS POR ÁREA O GES {te3}", level=2)

        # Obtener el nivel de riesgo para este TE3
        riesgo_te3 = df_resumen[
            (df_resumen['CUV'] == cuv) &
            (df_resumen['TE3'] == te3)
            ]['Riesgo']
        riesgo_te3 = riesgo_te3.values[0] if not riesgo_te3.empty else "Información no disponible"

        p = doc.add_paragraph()
        p.add_run('Nivel de riesgo: ').bold = True
        p.add_run(f"{riesgo_te3}\n")
        p.style.font.size = Pt(12)

        # Insertar el gráfico
        img_buffer = BytesIO()
        fig_te3.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        img_buffer.close()
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Salto de página
        doc.add_page_break()

    # Agregar tabla de medidas propuestas
    doc.add_heading(f"Medidas propuestas para {row['Nombre Centro de Trabajo']}", level=2)
    agregar_tabla_ges_por_dimension(doc, dimensiones, cuv, df_recomendaciones, df_resultados_porcentaje,
                                    df_porcentajes_niveles, top_glosas, datos_df)

    # Guardar el documento en un BytesIO para descarga
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer

# ===============================
# Configuración de la Página
# ===============================

st.set_page_config(page_title="Generador de Informes de Riesgos Psicosociales", layout="wide")

# ===============================
# Título y Descripción
# ===============================

st.title("Generador de Informes de Riesgos Psicosociales")
st.write("""
Esta aplicación permite generar informes técnicos basados en datos de riesgos psicosociales.
Por favor, cargue los archivos necesarios y siga las instrucciones.
""")

# ===============================
# Inicializar Estado de Sesión
# ===============================

if 'dimensiones' not in st.session_state:
    st.session_state['dimensiones'] = []

# ===============================
# Sección 1: Carga de Archivos
# ===============================

st.header("1. Carga de archivos")

# Cargar el archivo Excel con múltiples hojas
uploaded_file_combined = st.file_uploader("Selecciona el archivo 'combined_output.xlsx'", type="xlsx")

# Cargar los archivos
# uploaded_file_rec = st.file_uploader("Selecciona el archivo 'Recomendaciones2.xlsx'", type="xlsx")
# uploaded_file_ciiu = st.file_uploader("Selecciona el archivo 'ciiu.xlsx'", type="xlsx")
# uploaded_file_resultados = st.file_uploader("Selecciona el archivo 'resultados.xlsx'", type="xlsx")

#precargas
uploaded_file_rec = 'Recomendaciones2.xlsx'
uploaded_file_ciiu = 'ciiu.xlsx'
uploaded_file_resultados = 'resultados.xlsx'

if all([uploaded_file_combined, uploaded_file_rec, uploaded_file_ciiu, uploaded_file_resultados]):
    # Mapear los archivos cargados
    uploaded_files = {
        'combined': uploaded_file_combined,
        'recomendaciones': uploaded_file_rec,
        'ciiu': uploaded_file_ciiu,
        'resultados': uploaded_file_resultados
    }

    # Cargar y procesar los datos
    data_dict = cargar_datos(uploaded_files)
    if data_dict:
        # Convertir columnas a string donde sea necesario
        data_dict = convertir_columnas(data_dict)

        # Procesar recomendaciones
        df_recomendaciones = procesar_recomendaciones(data_dict['df_rec'], data_dict['df_ciiu'])

        # Convertir fechas
        data_dict['df_res_com'] = convertir_fechas(data_dict['df_res_com'])

        # Asegurarse de que 'top_glosas' esté cargado
        top_glosas = data_dict.get('top_glosas', pd.DataFrame())

        # Sección 2: Selección de CUV para generar el informe
        st.header("2. Seleccionar CUV para generar el informe")

        cuvs_disponibles = data_dict['summary_df']['CUV'].unique()
        selected_cuv = st.selectbox("Selecciona un CUV", cuvs_disponibles)

        # Filtrar los datos para el CUV seleccionado
        datos = data_dict['df_res_com'][data_dict['df_res_com']['CUV'] == selected_cuv]
        estado = data_dict['summary_df'][data_dict['summary_df']['CUV'] == selected_cuv]

        if datos.empty or estado.empty:
            st.error(f"No se encontraron datos para el CUV {selected_cuv}.")
        else:
            # Obtener la primera fila de datos
            datos = datos.iloc[0]
            estado_riesgo = estado['Riesgo'].values[0]

            # Sección 3: Generación y visualización de gráficos
            st.header("3. Gráfico principal de resultados")

            # Generar y mostrar el gráfico principal
            fig_principal = generar_grafico_principal(data_dict['df_resultados_porcentaje'], selected_cuv)
            if fig_principal:
                st.pyplot(fig_principal)
            else:
                st.warning("No se pudo generar el gráfico principal.")

            # Sección 4: Generación de gráficos por TE3
            st.header("4. Generación de gráficos por GES")

            # Generar y mostrar los gráficos por TE3
            figs_te3 = generar_graficos_por_te3(data_dict['df_porcentajes_niveles'], selected_cuv)

            if figs_te3:
                for fig_te3, te3 in figs_te3:
                    st.subheader(f"Gráfico para GES: {te3}")
                    st.pyplot(fig_te3)
            else:
                st.info("No se generaron gráficos adicionales por TE3.")

            # Sección 5: Prescripciones de medidas
            st.header("5. Prescripciones de medidas")

            # Obtener dimensiones relacionadas al CUV seleccionado
            dimensiones = data_dict['df_res_dimTE3'][data_dict['df_res_dimTE3']['CUV'] == selected_cuv].to_dict('records')
            st.session_state['dimensiones'] = dimensiones

            # Iterar sobre las dimensiones para mostrar campos de entrada
            for idx, dim in enumerate(st.session_state.dimensiones):
                st.subheader(f"Dimensión en riesgo: **{dim['dimension']}**")

                # Resultados (No editables)
                st.markdown("**Resultados:**")
                st.text(dim.get('resultados', "No disponible"))

                # Preguntas clave (No editables)
                st.markdown("**Preguntas clave:**")
                for pregunta in dim.get('preguntas', []):
                    st.text(f"- {pregunta}")

                # Campo editable para la explicación
                st.markdown("**Explicación:**")
                dim['explicacion'] = st.text_area(
                    f"Explicación para {dim['dimension']}", value=dim.get('explicacion', ""), key=f"explicacion_{idx}"
                )

                # Medidas propuestas (editable y prellenado)
                st.markdown("**Medidas propuestas:**")
                medidas = dim.get('medidas', [])
                for i, medida in enumerate(medidas):
                    medida['medida'] = st.text_area(
                        f"Medida {i + 1} para {dim['dimension']}", value=medida.get('medida', ''), key=f"medida_{idx}_{i}"
                    )
                    medida['fecha_monitoreo'] = st.selectbox(
                        f"Fecha monitoreo para medida {i + 1}",
                        options=["01-03-2024", "01-06-2024", "01-09-2024"],
                        index=["01-03-2024", "01-06-2024", "01-09-2024"].index(medida.get('fecha_monitoreo', "01-03-2024")),
                        key=f"fecha_{idx}_{i}"
                    )
                    medida['responsable'] = st.text_input(
                        f"Responsable seguimiento para medida {i + 1}", value=medida.get('responsable', ''),
                        key=f"responsable_{idx}_{i}"
                    )

                    # Botón para eliminar una medida
                    if st.button(f"Eliminar medida {i + 1} para {dim['dimension']}", key=f"eliminar_medida_{idx}_{i}"):
                        medidas.pop(i)
                        st.session_state.dimensiones[idx]['medidas'] = medidas
                        st.experimental_rerun()

                # Botón para agregar una nueva medida
                if st.button(f"Agregar nueva medida para {dim['dimension']}", key=f"agregar_medida_{idx}"):
                    medidas.append({'medida': '', 'fecha_monitoreo': "01-03-2024", 'responsable': ''})
                    st.session_state.dimensiones[idx]['medidas'] = medidas
                    st.experimental_rerun()

                # Guardar las medidas en el estado de sesión
                dim['medidas'] = medidas

                # Separador entre dimensiones
                st.markdown("---")

            # Sección 6: Generación del informe en Word
            st.header("6. Generación del informe en Word")


            def on_generate_report():
                doc_buffer = generar_informe(
                    df_res_com=data_dict['df_res_com'],
                    summary_df=data_dict['summary_df'],
                    df_resultados_porcentaje=data_dict['df_resultados_porcentaje'],
                    df_porcentajes_niveles=data_dict['df_porcentajes_niveles'],
                    CUV=selected_cuv,
                    df_resumen=data_dict['df_resumen'],
                    df_res_dimTE3=data_dict['df_res_dimTE3'],
                    generar_grafico_principal=generar_grafico_principal,
                    generar_graficos_por_te3=generar_graficos_por_te3,
                    dimensiones=st.session_state.dimensiones,
                    cuv=selected_cuv,
                    df_recomendaciones=df_recomendaciones,
                    top_glosas=top_glosas,
                    datos_df=data_dict['df_res_com']
                )

                if doc_buffer:
                    st.download_button(
                        label="Descargar informe",
                        data=doc_buffer,
                        file_name=f"Informe_{selected_cuv}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("No se pudo generar el informe.")


            if st.button("Generar informe en Word"):
                with st.spinner("Generando el informe, por favor espera..."):
                    on_generate_report()

else:
    st.info("Por favor, carga todos los archivos requeridos para continuar.")
