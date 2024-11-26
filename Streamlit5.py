# Importar las bibliotecas necesarias
import matplotlib.pyplot as plt
import unicodedata
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from openpyxl.worksheet.dimensions import Dimension
import streamlit as st
import pandas as pd
import numpy as np
from datetime import timedelta, datetime


# Configuración inicial de la aplicación
st.set_page_config(page_title="Generador de Informes de Riesgos Psicosociales")
st.title("Generador de Informes de Riesgos Psicosociales")
st.write("""
Esta aplicación permite generar informes técnicos basados en datos de riesgos psicosociales.
Por favor, cargue los archivos necesarios y siga las instrucciones.
""")

# Sección 1: Carga de archivos
#st.header("1. Carga de archivos")

# Cargar el archivo Excel con múltiples hojas
#uploaded_file_combined = st.file_uploader("Selecciona el archivo 'combined_output.xlsx'", type="xlsx")

# Cargar los archivos de recomendaciones y códigos CIIU
# uploaded_file_rec = st.file_uploader("Selecciona el archivo 'Recomendaciones2.xlsx'", type="xlsx")
# uploaded_file_ciiu = st.file_uploader("Selecciona el archivo 'ciiu.xlsx'", type="xlsx")

# Cargar el archivo 'resultados.xlsx'
# uploaded_file_resultados = st.file_uploader("Selecciona el archivo 'resultados.xlsx'", type="xlsx")
# Estamos casi a punto de salir a produccion, esta seria la version test.



#precargas
uploaded_file_combined = "combined_output.xlsx"
uploaded_file_rec = 'Recomendaciones.xlsx'
uploaded_file_ciiu = 'ciiu.xlsx'
uploaded_file_resultados = 'resultados.xlsx'


if (uploaded_file_combined is not None and
        uploaded_file_rec is not None and
        uploaded_file_ciiu is not None and
        uploaded_file_resultados is not None):

    # Leer las hojas necesarias desde 'combined_output.xlsx'
    try:
        combined_df_base_complet3 = pd.read_excel(uploaded_file_combined, sheet_name='basecompleta')
        summary_df = pd.read_excel(uploaded_file_combined, sheet_name='Summary')
        df_resultados_porcentaje = pd.read_excel(uploaded_file_combined, sheet_name='resultado')
        df_porcentajes_niveles = pd.read_excel(uploaded_file_combined, sheet_name='df_porcentajes_niveles')
        df_res_dimTE3 = pd.read_excel(uploaded_file_combined, sheet_name='df_res_dimTE3')
        df_resumen = pd.read_excel(uploaded_file_combined, sheet_name='df_resumen')
        top_glosas = pd.read_excel(uploaded_file_combined, sheet_name='top_glosas')

        #st.success("Archivo 'combined_output.xlsx' cargado exitosamente.")
    except Exception as e:
        st.error(f"Error al leer las hojas de 'combined_output.xlsx': {e}")

    # Leer los archivos de recomendaciones y CIIU
    try:
        df_recomendaciones = pd.read_excel(uploaded_file_rec, sheet_name='df_recomendaciones')  # Cambia 'Hoja1' si es necesario
        df_ciiu = pd.read_excel(uploaded_file_ciiu, sheet_name='ciiu')  # Cambia 'ciiu' si es necesario

        #st.success("Archivos 'Recomendaciones2.xlsx' y 'ciiu.xlsx' cargados exitosamente.")
    except Exception as e:
        st.error(f"Error al leer los archivos de recomendaciones y CIIU: {e}")

    # Leer y procesar 'resultados.xlsx'
    try:
        df_resultados = pd.read_excel(uploaded_file_resultados, sheet_name='Datos', usecols=['CUV', 'Folio'])
        df_res_com = pd.read_excel(uploaded_file_resultados, sheet_name='Datos')

        # Convertir 'CUV' a int64 en df_resultados
        #df_resultados['CUV'] = pd.to_numeric(df_resultados['CUV'], errors='coerce').astype(
        #    'Int64')  # Usa 'Int64' para permitir NaNs

        #st.success("Archivo 'resultados.xlsx' cargado y procesado exitosamente.")
    except Exception as e:
        st.error(f"Error al leer y procesar 'resultados.xlsx': {e}")


    # Asegurar que las columnas 'CUV' y 'ciiu' sean de tipo str
    df_resultados['CUV'] = df_resultados['CUV'].astype(str)
    df_res_com['CUV'] = df_res_com['CUV'].astype(str)
    summary_df['CUV'] = summary_df['CUV'].astype(str)
    df_resultados_porcentaje['CUV'] = df_resultados_porcentaje['CUV'].astype(str)
    df_porcentajes_niveles['CUV'] = df_porcentajes_niveles['CUV'].astype(str)
    df_res_dimTE3['CUV'] = df_res_dimTE3['CUV'].astype(str)
    df_resumen['CUV'] = df_resumen['CUV'].astype(str)
    top_glosas['CUV'] = top_glosas['CUV'].astype(str)
    df_recomendaciones['ciiu'] = df_recomendaciones['ciiu'].astype(str)

    def procesar_columna_fecha(df, columna, formato='%d-%m-%Y'):
        """
        Convierte una columna de fechas en el DataFrame al formato deseado.

        Parameters:
        - df (pd.DataFrame): El DataFrame que contiene la columna.
        - columna (str): El nombre de la columna a procesar.
        - formato (str): El formato al que se desea convertir la fecha.

        Returns:
        - pd.DataFrame: El DataFrame con la columna procesada.
        """
        if columna in df.columns:
            df[columna] = pd.to_datetime(df[columna], errors='coerce').dt.strftime(formato)
            #st.success(f"Columna '{columna}' procesada correctamente.")
        else:
            st.error(
                f"La columna '{columna}' no se encontró en el DataFrame 'df_res_com'. Por favor, verifica el nombre de la columna.")
            st.write("Columnas disponibles en 'df_res_com':", df.columns.tolist())
        return df


    columnas_fecha = ['Fecha Inicio', 'Fecha Fin']

    # Procesar cada columna de fecha
    for columna in columnas_fecha:
        df_res_com = procesar_columna_fecha(df_res_com, columna)





    # Definir funciones auxiliares
    def normalizar_texto(texto):
        if isinstance(texto, str):
            texto = texto.strip().lower()
            texto = ''.join(
                c for c in unicodedata.normalize('NFD', texto)
                if unicodedata.category(c) != 'Mn'
            )
            return texto
        else:
            return ''


    def obtener_dimm_por_dimension(nombre_dimension):
        nombre_dimension_normalizado = normalizar_texto(nombre_dimension)
        # df_dimensiones debería estar definido o cargado desde alguna parte
        # Aquí debes asegurarte de tener el DataFrame 'dimensiones' disponible
        # Por ejemplo, podrías cargarlo desde un archivo o definirlo manualmente
        # Supongamos que está en 'df_res_dimTE3'
        # Ajusta esto según tu estructura de datos
        df_dimensiones = df_res_dimTE3[['Dimensión', 'dimm']].drop_duplicates()
        df_dimensiones['dimension_normalizada'] = df_dimensiones['Dimensión'].apply(normalizar_texto)
        resultado = df_dimensiones[df_dimensiones['dimension_normalizada'] == nombre_dimension_normalizado]
        if not resultado.empty:
            return resultado.iloc[0]['dimm']
        else:
            st.warning(f"No se encontró el código 'dimm' para la dimensión '{nombre_dimension}'.")
            return None


    def agregar_tabla_ges_por_dimension(doc, df, cuv, df_recomendaciones, df_resultados_porcentaje,
                                        df_porcentajes_niveles, top_glosas, datos):
        """
        Agrega una tabla de dimensiones y GES para un CUV específico en el documento de Word.

        Parámetros:
        doc (Document): El objeto del documento de Word.
        df (pd.DataFrame): DataFrame con los datos de dimensiones y GES filtrados.
        cuv (str): El CUV específico para el que se generará la tabla.
        df_recomendaciones (pd.DataFrame): DataFrame con las recomendaciones por dimensión.
        """

        # Filtrar el DataFrame para el CUV específico
        df_revision = df[df['CUV'] == cuv]
        unique_te3 = df_revision['TE3'].dropna().unique()

        if len(unique_te3) < 2:
            noges = 1
            print(f"Solo hay un GES para el CUV {cuv}. No se generarán recomendaciones por GES.")
        else:
            noges = 0

        # Filtrar el DataFrame para el CUV específico y puntajes 1 y 2
        df_cuv = df[(df['CUV'] == cuv) & (df['Puntaje'].isin([1, 2]))]

        if df_cuv.empty:
            st.warning(f"No hay datos con puntaje 1 o 2 para el CUV {cuv}.")
            return

        # Agrupar por 'Dimensión' y combinar los valores únicos de 'TE3' en una lista separada por "; "
        resultado = df_cuv.groupby('Dimensión')['TE3'].unique().reset_index()
        resultado['GES'] = resultado['TE3'].apply(lambda x: '; '.join(map(str, x)))

        # Limpiar el campo 'GES' reemplazando ciertos caracteres
        resultado['GES'] = resultado['GES'].str.replace('|', '_', regex=False) \
            .str.replace(':', '_', regex=False) \
            .str.replace('?', '_', regex=False)

        # Asegúrate de que los valores en 'CIIU' son strings y extraer la parte necesaria
        datos['CIIU'] = datos['CIIU'].apply(lambda x: x.split('_')[-1] if isinstance(x, str) else x)

        # Filtrar el DataFrame para el CUV específico y obtener el valor único de CIIU
        ciiu_valor = datos.loc[datos['CUV'] == cuv, 'CIIU'].copy()

        if len(ciiu_valor) > 0:
            ciiu_unico = ciiu_valor.iloc[0]
            if isinstance(ciiu_unico, str) and ciiu_unico.isdigit():
                ciiu = int(ciiu_unico[:2]) if len(ciiu_unico) > 5 else int(ciiu_unico[:1])

            else:
                print("El valor de CIIU no es numérico.")
                ciiu = None
        else:
            print("CUV no encontrado en la tabla de datos.")
            ciiu = None

        if ciiu is None:
            st.error(f"No se pudo determinar el valor de CIIU para el CUV {cuv}.")
            return

        # Crear la tabla en el documento
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        column_widths = [Inches(0.5), Inches(0.5), Inches(0.5), Inches(7), Inches(0.5), Inches(0.5)]

        # Configurar el ancho de cada columna
        for col_idx, width in enumerate(column_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width

        # Agregar encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dimensión en riesgo'
        hdr_cells[1].text = 'Preguntas clave'
        hdr_cells[2].text = 'Explicación'
        hdr_cells[3].text = 'Medidas propuestas'
        hdr_cells[4].text = 'Fecha monitoreo'
        hdr_cells[5].text = 'Responsable seguimiento'

        # Asegurarse de que 'Descripción' sea string y reemplazar NaN
        df_resultados_porcentaje['Descripción'] = df_resultados_porcentaje['Descripción'].fillna('').astype(str)

        # Rellenar la tabla con los datos de 'Dimensión' y 'GES'
        for _, row in resultado.iterrows():
            dim = row['Dimensión']
            ges = row['GES']

            # Obtener las recomendaciones para esta dimensión
            recomendaciones = df_recomendaciones[
                (df_recomendaciones['Dimensión'] == dim) &
                (df_recomendaciones['ciiu'] == str(ciiu))
                ]['Recomendación'].tolist()
            medidas_propuestas = '\n'.join([f"- {rec}" for rec in recomendaciones]) if recomendaciones else 'N/A'

            # Obtener la descripción relacionada desde df_resultados_porcentaje
            descripcion = df_resultados_porcentaje[
                (df_resultados_porcentaje['Dimensión'] == dim) &
                (df_resultados_porcentaje['CUV'] == cuv)
                ]['Descripción'].values

            # Filtrar solo cadenas no vacías
            descripcion = [desc for desc in descripcion if isinstance(desc, str) and desc.strip() != '']

            descripcion2 = [
                f"{desc} en {ges}"
                for desc in df_porcentajes_niveles[
                    (df_porcentajes_niveles['Dimensión'] == dim) &
                    (df_porcentajes_niveles['CUV'] == cuv) &
                    (df_porcentajes_niveles['TE3'] == ges) &
                    (df_porcentajes_niveles['Descripción'].str.strip() != '')
                    ]['Descripción'].values
            ]

            descripcion2_text = '\n'.join(descripcion2).replace("[", "").replace("]", "").replace("'",
                                                                                                  "") if descripcion2 else ""

            # Construir descripcion_text
            descripcion_text = ""
            if len(descripcion) > 0 and isinstance(descripcion[0], str) and len(descripcion[0]) > 0:
                descripcion_text = descripcion[0] + " para todo el centro de trabajo\n"
            elif len(descripcion) > 1 and isinstance(descripcion[1], str) and len(descripcion[1]) > 0:
                descripcion_text = descripcion[1] + " para todo el centro de trabajo\n"
            elif len(descripcion) > 2 and isinstance(descripcion[2], str) and len(descripcion[2]) > 0:
                descripcion_text = descripcion[2] + " para todo el centro de trabajo\n"
            else:
                descripcion_text = ""

            # Verificar si hay múltiples GES
            if noges == 1:
                descripcion2_text = ""
                print(f"Solo hay un GES para el CUV {cuv}. No se generarán recomendaciones por GES.")

            # Obtener las preguntas clave desde top_glosas
            filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dim) & (top_glosas['CUV'] == cuv)]
            preguntas = filtro_glosas['Pregunta'].tolist()
            preguntas_text = '\n'.join(preguntas) if preguntas else 'N/A'

            # Rellenar las celdas de la tabla
            row_cells = table.add_row().cells
            row_cells[0].text = f"{descripcion_text}{descripcion2_text}".strip()
            row_cells[1].text = preguntas_text.strip()
            row_cells[2].text = ''  # Espacio para 'Explicación'
            row_cells[3].text = medidas_propuestas.strip()
            row_cells[4].text = ''  # Fecha de monitoreo
            row_cells[5].text = ''  # Responsable seguimiento

        # Ajustar el tamaño de fuente de las celdas (opcional)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Ajusta el tamaño de la fuente


    def agregar_tabla_ges_por_dimension_streamlit(df, cuv, df_recomendaciones, df_porcentajes_niveles, top_glosas, df_res_com):
        """
        Prepara una estructura de datos con medidas propuestas por dimensión y retorna una lista de diccionarios.
        """
        # Verificar tipos de entrada
        if not isinstance(df, pd.DataFrame):
            st.error("El parámetro 'df' debe ser un DataFrame.")
            return []
        if not isinstance(df_recomendaciones, pd.DataFrame):
            st.error("El parámetro 'df_recomendaciones' debe ser un DataFrame.")
            return []
        if not isinstance(df_porcentajes_niveles, pd.DataFrame):
            st.error("El parámetro 'df_porcentajes_niveles' debe ser un DataFrame.")
            return []
        if not isinstance(top_glosas, pd.DataFrame):
            st.error("El parámetro 'top_glosas' debe ser un DataFrame.")
            return []
        if not isinstance(df_res_com, pd.DataFrame):
            st.error("El parámetro 'df_res_com' debe ser un DataFrame.")
            return []

        # Filtrar el DataFrame para el CUV y puntajes 1 y 2
        df_cuv = df[(df['CUV'] == cuv) & (df['Puntaje'].isin([1, 2]))]
        if df_cuv.empty:
            st.warning(f"No hay datos con puntaje 1 o 2 para el CUV {cuv}.")
            return []

        # Agrupar por 'Dimensión' y obtener los valores únicos de 'TE3'
        resultado = df_cuv.groupby('Dimensión')['TE3'].unique().reset_index()
        resultado['GES'] = resultado['TE3'].apply(lambda x: '; '.join(map(str, x)))

        # Limpiar el campo 'GES' reemplazando ciertos caracteres
        resultado['GES'] = resultado['GES'].str.replace('|', '_', regex=False) \
            .str.replace(':', '_', regex=False) \
            .str.replace('?', '_', regex=False)

        # Asegúrate de que los valores en 'CIIU' son strings y extraer la parte necesaria
        df_res_com['CIIU'] = df_res_com['CIIU'].apply(lambda x: x.split('_')[-1] if isinstance(x, str) else x)

        # Filtrar el DataFrame para el CUV específico y obtener el valor único de CIIU
        ciiu_valor = df_res_com.loc[df_res_com['CUV'] == cuv, 'CIIU']

        if not ciiu_valor.empty:
            ciiu_unico = ciiu_valor.iloc[0]
            if isinstance(ciiu_unico, str) and ciiu_unico.isdigit():
                ciiu = int(ciiu_unico[:2]) if len(ciiu_unico) > 5 else int(ciiu_unico[:1])

            else:
                st.error("El valor de CIIU no es numérico.")
                return []
        else:
            st.error(f"No se encontró el valor de CIIU para el CUV {cuv}.")
            return []

        df_porcentajes_niveles['Descripción'] = df_porcentajes_niveles.apply(
            lambda row: f"{row['Porcentaje']}% Riesgo {row['Nivel']}, {row['Respuestas']} personas"
            if row['Puntaje'] in [1, 2] else "",
            axis=1
        )

        # Asegurarse de que 'Descripción' sea string y reemplazar NaN
        df_resultados_porcentaje['Descripción'] = df_resultados_porcentaje['Descripción'].fillna('').astype(str)

        # Inicializar una lista para almacenar las dimensiones
        dimensiones = []

        # Rellenar la lista con los datos de 'Dimensión' y 'GES'
        for _, row in resultado.iterrows():
            dim = row['Dimensión']
            ges = row['GES']

            # Obtener las recomendaciones para esta dimensión y CIIU
            recomendaciones = df_recomendaciones[
                (df_recomendaciones['Dimensión'] == dim) &
                (df_recomendaciones['ciiu'] == str(ciiu))
                ]['Recomendación'].tolist()
            medidas_propuestas = recomendaciones if recomendaciones else ['N/A']

            # Obtener las descripciones desde df_porcentajes_niveles
            descripcion = [
                f"{desc} en {ges}"
                for desc in df_porcentajes_niveles[
                    (df_porcentajes_niveles['Dimensión'] == dim) &
                    (df_porcentajes_niveles['CUV'] == cuv) &
                    (df_porcentajes_niveles['TE3'] == ges) &
                    (df_porcentajes_niveles['Descripción'].str.strip() != '')
                    ]['Descripción'].values
            ]

            descripcion_text = '\n'.join(descripcion).replace("[", "").replace("]", "").replace("'",
                                                                                                  "") if descripcion else ""

            # Obtener las preguntas clave desde top_glosas
            filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dim) & (top_glosas['CUV'] == cuv)]
            preguntas = filtro_glosas['Pregunta'].tolist()
            preguntas_text = '\n'.join([f"- {pregunta}" for pregunta in preguntas]) if preguntas else 'N/A'

            # Agregar la dimensión a la lista
            dimensiones.append({
                'GES': ges,
                'Dimensión en riesgo': dim,
                'Descripción riesgo': descripcion_text.strip(),
                'Preguntas clave': preguntas_text.strip(),
                'Interpretación del grupo de discusión': "",
                'Medidas propuestas': medidas_propuestas  # Lista de medidas
            })

        return dimensiones


    def mostrar_datos(datos):
        """
        Muestra los datos de una empresa formateados en Markdown.

        Parameters:
        - datos (dict): Diccionario con los datos de la empresa.
        """
        contenido = f"""
        **Razón Social:** {datos.get('Nombre Empresa', 'N/A')}  
        **RUT:** {datos.get('RUT Empresa Lugar Geográfico', 'N/A')}  
        **Nombre del centro de trabajo:** {datos.get('Nombre Centro de Trabajo', 'N/A')}  
        **CUV:** {datos.get('CUV', 'N/A')}  
        **CIIU:** {datos.get('CIIU CT', 'N/A').split('_')[-1] if 'CIIU CT' in datos else 'N/A'}  
        **Fecha de activación del cuestionario:** {datos.get('Fecha Inicio', 'N/A')}  
        **Fecha de cierre del cuestionario:** {datos.get('Fecha Fin', 'N/A')}  
        **Universo de trabajadores de evaluación:** {datos.get('Nº Trabajadores CT', 'N/A')}  
        """
        st.markdown(contenido)


    # Sección 2: Selección de CUV para generar el informe
    st.header("2. Seleccionar CUV para generar el informe")

    cuvs_disponibles = summary_df['CUV'].unique()
    selected_cuv = st.selectbox("Selecciona un CUV", cuvs_disponibles)

    # Filtrar los datos para el CUV seleccionado
    datos = df_res_com[df_res_com['CUV'] == selected_cuv]
    estado = summary_df[summary_df['CUV'] == selected_cuv]

    fecha_fin = pd.to_datetime(datos.get('Fecha Fin'))

    # Mostrar la información de la empresa
    st.subheader("Información de la Empresa")
    for _, row in datos.iterrows():
        mostrar_datos(row.to_dict())
        st.markdown("---")  # Línea separadora entre empresas


    if datos.empty or estado.empty:
        st.error(f"No se encontraron datos para el CUV {selected_cuv}.")
    else:
        # Obtener la primera fila de datos
        datos = datos.iloc[0]
        estado_riesgo = estado['Riesgo'].values[0]

        # Sección 3: Generación y visualización de gráficos
        st.header("3. Gráfico principal de resultados")


        def generar_grafico_principal(df, CUV):
            df_filtrado = df[df['CUV'] == CUV]
            if df_filtrado.empty:
                st.warning(f"No se encontraron datos para el CUV {CUV}.")
                return None

            try:
                df_pivot = df_filtrado.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[
                           ::-1]
            except Exception as e:
                st.error(f"Error al pivotear los datos: {e}")
                return None

            fig, ax = plt.subplots(figsize=(12, 8))

            niveles = ["Bajo", "Medio", "Alto"]
            colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
            posiciones = np.arange(len(df_pivot.index))
            ancho_barra = 0.2

            for i, nivel in enumerate(niveles):
                if nivel in df_pivot.columns:
                    valores = df_pivot[nivel]
                    ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                            label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
                else:
                    st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot.")

            ax.axvline(50, color="blue", linestyle="--", linewidth=1)
            ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}", pad=50)
            ax.set_xlabel("Porcentaje")
            ax.set_ylabel("Dimensiones")
            ax.set_xlim(0, 100)
            ax.set_yticks(posiciones + ancho_barra)
            ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

            fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.6, 0.96), ncol=3)
            plt.subplots_adjust(left=0.3, top=0.85)
            plt.tight_layout()

            return fig


        # Generar y mostrar el gráfico principal
        fig_principal = generar_grafico_principal(df_resultados_porcentaje, selected_cuv)
        if fig_principal:
            st.pyplot(fig_principal)
        else:
            st.warning("No se pudo generar el gráfico principal.")

        # ---- NUEVA SECCIÓN PARA MOSTRAR DIMENSIONES EN RIESGO ----
        #st.header("4. Dimensiones en Riesgo")

        # Obtener dimensiones en riesgo
        dimensiones_riesgo_alto = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == 2)
                ]['Dimensión'].tolist()

        dimensiones_riesgo_medio = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == 1)
                ]['Dimensión'].tolist()

        dimensiones_riesgo_bajo = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == -2)
                ]['Dimensión'].tolist()

        # Mostrar dimensiones en riesgo en la web
        st.subheader("Dimensiones en riesgo")
        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("**Alto:**")
            if dimensiones_riesgo_alto:
                st.write(", ".join(dimensiones_riesgo_alto))
            else:
                st.write("Ninguna")

        with col2:
            st.markdown("**Medio:**")
            if dimensiones_riesgo_medio:
                st.write(", ".join(dimensiones_riesgo_medio))
            else:
                st.write("Ninguna")

        with col3:
            st.markdown("**Bajo:**")
            if dimensiones_riesgo_bajo:
                st.write(", ".join(dimensiones_riesgo_bajo))
            else:
                st.write("Ninguna")


        # Sección 4: Generación de gráficos por TE3
        st.header("4. Generación de gráficos por GES")


        def generar_graficos_por_te3(df, CUV):
            """
            Genera una lista de figuras de gráficos para cada valor de TE3 dentro del CUV especificado.
            """
            df_cuv = df[df['CUV'] == CUV]
            if df_cuv.empty:
                st.warning(f"No se encontraron datos para el CUV {CUV}.")
                return []

            if 'TE3' not in df_cuv.columns:
                st.warning(f"La columna 'TE3' no existe en el DataFrame para CUV {CUV}.")
                return []

            te3_values = df_cuv['TE3'].unique()

            # Verificar si hay más de un valor de TE3
            if len(te3_values) <= 1:
                st.info(f"Solo hay un GES para el CUV {CUV}. No se generarán gráficos adicionales.")
                return []

            figs_te3 = []

            for te3 in te3_values:
                df_te3 = df_cuv[df_cuv['TE3'] == te3]
                if df_te3.empty:
                    continue
                try:
                    df_pivot = df_te3.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[
                               ::-1]
                except Exception as e:
                    st.error(f"Error al pivotear los datos para TE3 {te3}: {e}")
                    continue

                fig, ax = plt.subplots(figsize=(12, 8))

                niveles = ["Bajo", "Medio", "Alto"]
                colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
                posiciones = np.arange(len(df_pivot.index))
                ancho_barra = 0.2

                for i, nivel in enumerate(niveles):
                    if nivel in df_pivot.columns:
                        valores = df_pivot[nivel]
                        ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                                label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
                    else:
                        st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot para TE3 {te3}.")

                ax.axvline(50, color="blue", linestyle="--", linewidth=1)
                ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}, TE3 {te3}", pad=50)
                ax.set_xlabel("Porcentaje")
                ax.set_ylabel("Dimensiones")
                ax.set_xlim(0, 100)
                ax.set_yticks(posiciones + ancho_barra)
                ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

                fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.96), ncol=3)
                plt.tight_layout()

                figs_te3.append((fig, te3))

            return figs_te3


        # Generar y mostrar los gráficos por TE3
        figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles, selected_cuv)

        if figs_te3:
            for fig_te3, te3 in figs_te3:
                st.subheader(f"Gráfico para GES: {te3}")
                st.pyplot(fig_te3)
                dimensiones_riesgo_alto = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == 2)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_medio = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == 1)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_bajo = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == -2)
                    ]['Dimensión'].tolist()

                # Crear tres columnas
                col1, col2, col3 = st.columns(3)

                with col1:
                    st.markdown("**Alto:**")
                    if dimensiones_riesgo_alto:
                        st.write(", ".join(dimensiones_riesgo_alto))
                    else:
                        st.write("Ninguna")

                with col2:
                    st.markdown("**Medio:**")
                    if dimensiones_riesgo_medio:
                        st.write(", ".join(dimensiones_riesgo_medio))
                    else:
                        st.write("Ninguna")

                with col3:
                    st.markdown("**Bajo:**")
                    if dimensiones_riesgo_bajo:
                        st.write(", ".join(dimensiones_riesgo_bajo))
                    else:
                        st.write("Ninguna")


        else:
            st.info("No se generaron gráficos adicionales por GES.")


        # Sección 5: Prescripciones de medidas
        st.header("5. Prescripciones de medidas")

        dimensiones_te3 = agregar_tabla_ges_por_dimension_streamlit(df_res_dimTE3, selected_cuv, df_recomendaciones,
                                            df_porcentajes_niveles, top_glosas, df_res_com)

        #st.session_state.dimensiones_te3 = dimensiones_te3

        # Función para actualizar la numeración
        def actualizar_numeracion(df):
            df = df.reset_index(drop=True)
            df['N°'] = df.index + 1
            return df


        # Agrupar las dimensiones por GES correctamente desglosando valores combinados
        ges_groups = {}
        for dimension in dimensiones_te3:
            ges_values = dimension['GES'].split(";")  # Dividir los valores de GES
            for ges in ges_values:
                ges = ges.strip()  # Eliminar espacios adicionales
                if ges not in ges_groups:
                    ges_groups[ges] = []
                # Crear una copia de la dimensión para evitar duplicados en diferentes GES
                dimension_copy = dimension.copy()
                dimension_copy['GES'] = ges  # Asignar el GES desglosado
                ges_groups[ges].append(dimension_copy)


        # Definir la función para formatear opciones del selectbox
        def format_option(option):
            nombre, fecha = option
            if pd.notnull(fecha):
                return f"{nombre} - {fecha.strftime('%d-%m-%Y')}"
            else:
                return nombre


        # Crear un diccionario temporal para almacenar interpretaciones en st.session_state si no existe
        if 'interpretaciones_temporales' not in st.session_state:
            st.session_state['interpretaciones_temporales'] = {}

        # Procesar cada GES
        for ges, dimensiones in ges_groups.items():
            st.header(f"GES: {ges}")

            # Procesar cada Dimensión dentro del GES
            for idx, dimension in enumerate(dimensiones, 1):
                st.subheader(f"Dimensión: {dimension['Dimensión en riesgo']}")
                st.write(f"**Descripción del riesgo:** {dimension['Descripción riesgo']}")
                st.write("**Preguntas clave:**")
                st.write(dimension["Preguntas clave"])

                # Gestionar la interpretación del grupo de discusión
                interpretacion_key = f"interpretacion_{ges}_{idx}"

                # Establecer un valor inicial en el diccionario temporal si no existe
                if interpretacion_key not in st.session_state['interpretaciones_temporales']:
                    st.session_state['interpretaciones_temporales'][interpretacion_key] = st.session_state.get(
                        interpretacion_key, "")

                # Mostrar el cuadro de texto para la interpretación
                interpretacion = st.text_area(
                    label="Interpretación del grupo de discusión",
                    value=st.session_state['interpretaciones_temporales'][interpretacion_key],
                    height=150,
                    key=interpretacion_key
                )

                # Actualizar el valor en el diccionario temporal sin modificar el valor en el estado directamente asociado al widget
                st.session_state['interpretaciones_temporales'][interpretacion_key] = interpretacion

                # Gestionar medidas propuestas
                st.write("#### Medidas propuestas")
                session_key = f"measures_{ges}_{idx}"
                if session_key not in st.session_state:
                    medidas_data = [
                        {
                            'N°': i + 1,
                            'GES': ges,
                            'Dimensión': dimension['Dimensión en riesgo'],
                            'Medida': medida,
                            'Fecha monitoreo': '',
                            'Responsable': '',
                            'Activo': True,
                            'Seleccionada': False
                        }
                        for i, medida in enumerate(dimension['Medidas propuestas'])
                    ]
                    st.session_state[session_key] = pd.DataFrame(medidas_data)

                df = st.session_state[session_key]
                medidas_list = [""] + df.loc[df['Activo'], 'Medida'].tolist()  # Añadir opción vacía al inicio
                selected_measure = st.selectbox(
                    "Seleccione una medida para editar o deje vacío para crear una nueva",
                    medidas_list,
                    key=f"select_{ges}_{idx}"
                )

                if selected_measure:  # Si selecciona una medida existente
                    medida_idx = df[df['Medida'] == selected_measure].index[0]
                    st.write("#### Editar medida seleccionada")
                else:  # Si no selecciona nada, permite crear una nueva medida
                    st.write("#### Crear una nueva medida")
                    medida_idx = None

                # Calcular las fechas de corto, mediano y largo plazo
                fecha_fin = pd.to_datetime("2024-11-26")  # Reemplaza con tu lógica para obtener 'Fecha Fin'
                fecha_corto_plazo = fecha_fin + timedelta(days=240)
                fecha_mediano_plazo = fecha_fin + timedelta(days=330)
                fecha_largo_plazo = fecha_fin + timedelta(days=420)

                # Crear una lista con las fechas calculadas
                fechas_opciones = [
                    ("Corto Plazo (180 días)", fecha_corto_plazo),
                    ("Mediano Plazo (270 días)", fecha_mediano_plazo),
                    ("Largo Plazo (360 días)", fecha_largo_plazo),
                    ("Otra fecha", None)  # Opción para seleccionar una fecha personalizada
                ]

                # Crear formulario para editar o crear medida
                with st.form(key=f"form_{ges}_{idx}"):
                    medida = st.text_area(
                        "Descripción de la medida",
                        value=df.at[medida_idx, 'Medida'] if medida_idx is not None else "",
                        key=f"edit_medida_{ges}_{idx}",
                        height=90
                    )

                    # Mostrar la lista desplegable para seleccionar la fecha de monitoreo
                    opcion_seleccionada = st.selectbox(
                        "Selecciona la Fecha de Monitoreo",
                        options=fechas_opciones,
                        format_func=format_option,
                        key=f"select_fecha_{ges}_{idx}"
                    )

                    # Determinar la fecha seleccionada
                    if opcion_seleccionada[1]:
                        # Si se selecciona una de las opciones predefinidas
                        fecha = opcion_seleccionada[1]
                    else:
                        # Si se selecciona 'Otra fecha', mostrar el date_input para elegir manualmente
                        fecha_default = pd.to_datetime(df.at[medida_idx, 'Fecha monitoreo']) if (
                                medida_idx is not None and pd.notna(df.at[medida_idx, 'Fecha monitoreo'])
                        ) else datetime.today()
                        fecha = st.date_input(
                            "Selecciona una Fecha de Monitoreo personalizada",
                            value=fecha_default,
                            key=f"edit_fecha_personalizada_{ges}_{idx}"
                        )

                    # Asegurarse de que 'fecha' sea un objeto datetime
                    if isinstance(fecha, datetime):
                        fecha_formateada = fecha.strftime('%d-%m-%Y')
                    else:
                        # Si 'fecha' es un objeto 'date', convertir a datetime
                        fecha = datetime.combine(fecha, datetime.min.time())
                        fecha_formateada = fecha.strftime('%d-%m-%Y')

                    responsable = st.text_input(
                        "Responsable",
                        value=df.at[medida_idx, 'Responsable'] if medida_idx is not None else "",
                        key=f"edit_responsable_{ges}_{idx}"
                    )

                    # Botón para enviar el formulario
                    submit_button = st.form_submit_button(label="Confirmar selección o crear nueva medida")

                # Procesar la acción del formulario
                if submit_button:
                    if medida_idx is not None:  # Editar medida existente
                        st.session_state[session_key].at[medida_idx, 'Medida'] = medida
                        st.session_state[session_key].at[medida_idx, 'Fecha monitoreo'] = fecha.strftime(
                            '%Y-%m-%d') if fecha else ''
                        st.session_state[session_key].at[medida_idx, 'Responsable'] = responsable
                        st.session_state[session_key].at[medida_idx, 'Seleccionada'] = True
                        st.success("Medida actualizada correctamente")
                    else:  # Crear nueva medida
                        nueva_medida = {
                            "N°": len(st.session_state[session_key]) + 1,
                            "GES": ges,
                            "Dimensión": dimension['Dimensión en riesgo'],
                            "Medida": medida,
                            "Fecha monitoreo": fecha.strftime('%d-%m-%Y') if fecha else '',
                            "Responsable": responsable,
                            "Activo": True,
                            "Seleccionada": True
                        }
                        st.session_state[session_key] = pd.concat(
                            [st.session_state[session_key], pd.DataFrame([nueva_medida])],
                            ignore_index=True
                        )
                        st.success("Nueva medida creada correctamente")

        # Botón para guardar todas las interpretaciones
        if st.button("Guardar todas las interpretaciones"):
            for key, interpretacion in st.session_state['interpretaciones_temporales'].items():
                if key not in st.session_state:
                    st.session_state[
                        key] = interpretacion  # Guardar en `st.session_state` solo si no ha sido instanciado por el widget
            st.success("Todas las interpretaciones se han guardado correctamente")

        # Nueva Sección: Resumen de datos confirmados
        st.header("Resumen de datos confirmados")
        confirmed_measures = []
        interpretaciones_data = []

        for ges, dimensiones in ges_groups.items():
            for idx, dimension in enumerate(dimensiones, 1):
                session_key = f"measures_{ges}_{idx}"
                interpretacion_key = f"interpretacion_{ges}_{idx}"

                # Procesar medidas confirmadas
                if session_key in st.session_state:
                    temp_df = st.session_state[session_key].copy()
                    temp_df = temp_df[temp_df['Seleccionada']]  # Filtrar solo medidas seleccionadas

                    # Agregar información del GES y la Dimensión
                    temp_df['Dimensión'] = dimension["Dimensión en riesgo"]
                    temp_df['GES'] = ges

                    # Añadir la interpretación correspondiente a todas las medidas de la dimensión
                    if interpretacion_key in st.session_state['interpretaciones_temporales']:
                        temp_df['Interpretación'] = st.session_state['interpretaciones_temporales'][interpretacion_key]
                    else:
                        temp_df['Interpretación'] = ""

                    confirmed_measures.append(temp_df)

                # Procesar las interpretaciones de cada dimensión
                if interpretacion_key in st.session_state['interpretaciones_temporales']:
                    interpretacion = st.session_state['interpretaciones_temporales'][interpretacion_key]
                    interpretaciones_data.append({
                        'GES': ges,
                        'Dimensión': dimension["Dimensión en riesgo"],
                        'Interpretación': interpretacion
                    })

        # Mostrar Resumen de Interpretaciones
        if interpretaciones_data:
            interpretaciones_df = pd.DataFrame(interpretaciones_data)
            if not interpretaciones_df.empty:
                st.write("Las interpretaciones ingresadas hasta el momento:")
                st.dataframe(interpretaciones_df[['GES', 'Dimensión', 'Interpretación']])
        else:
            st.info("No hay interpretaciones ingresadas hasta el momento.")


        # Mostrar Resumen de Medidas Confirmadas
        if confirmed_measures:
            summary_df = pd.concat(confirmed_measures, ignore_index=True)
            if not summary_df.empty:
                st.write("Las siguientes medidas han sido confirmadas hasta el momento:")
                st.dataframe(
                    summary_df[['GES', 'Dimensión', 'Medida', 'Fecha monitoreo', 'Responsable', 'Interpretación']])
            else:
                st.info("No hay medidas confirmadas hasta el momento.")
        else:
            st.info("No hay medidas confirmadas hasta el momento.")


        # Exportar como CSV las interpretaciones ingresadas
        if 'interpretaciones_df' in locals() and not interpretaciones_df.empty:
            csv_interpretaciones = interpretaciones_df.to_csv(index=False)
            st.download_button(
                label="Descargar archivo CSV con Interpretaciones",
                data=csv_interpretaciones,
                file_name="interpretaciones_ingresadas.csv",
                mime="text/csv",
            )
            st.success("Datos de interpretaciones guardados correctamente.")

        # Exportar como CSV
        if 'summary_df' in locals() and not summary_df.empty:
            csv = summary_df.to_csv(index=False)
            st.download_button(
                label="Descargar archivo CSV con Medidas Confirmadas",
                data=csv,
                file_name="medidas_seleccionadas.csv",
                mime="text/csv",
            )
            st.success("Datos de medidas guardados correctamente.")
        else:
            st.warning("No se han seleccionado medidas para guardar.")



        # Sección 6: Generación del informe en Word
        st.header("6. Generación del informe en Word")

        def establecer_orientacion_apaisada(doc):
            """
            Configura el documento en orientación horizontal (apaisado).
            """
            section = doc.sections[0]
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)




        def generar_contenido_word(datos, estado_riesgo, fig_principal, figs_te3, df_resumen, df_porcentajes_niveles,
                                   df_resultados_porcentaje):
            """
            Genera el contenido del informe en un objeto Document de python-docx.
            """
            # Crear un nuevo documento
            doc = Document()
            establecer_orientacion_apaisada(doc)
            # section = doc.sections[0]
            # section.page_height = Inches(11)  # 11 pulgadas de alto para Carta
            # section.page_width = Inches(8.5)  # 8.5 pulgadas de ancho para Carta

            # Establecer Calibri como fuente predeterminada para el estilo 'Normal'
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(9)  # Tamaño de fuente opcional; ajusta según prefieras

            # Crear un nuevo estilo llamado 'destacado' con Calibri y tamaño de fuente 12
            destacado = doc.styles.add_style('destacado', 1)  # 1 para párrafos
            destacado_font = destacado.font
            destacado_font.name = 'Calibri'
            destacado_font.size = Pt(12)  # Tamaño de la fuente en puntos

            # Configurar el idioma del documento en español
            lang = doc.styles['Normal'].element
            lang.set(qn('w:lang'), 'es-ES')

            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            # Agregar imagen del logo (ajusta la ruta de la imagen a tu ubicación)
            doc.add_picture('/mount/src/informeceal/IST.jpg', width=Inches(2))  # Ajusta el tamaño según sea necesario
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinear a la derecha
            doc.add_paragraph()

            # Título principal
            titulo = doc.add_heading('INFORME TÉCNICO', level=1)
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Subtítulo
            subtitulo = doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2)
            subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Subtítulo
            subtitulo = doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2)
            subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()

            # Información general
            p = doc.add_paragraph()
            p.add_run('Razón Social: ').bold = True
            p.add_run(f"{datos['Nombre Empresa']}\n")
            p.add_run('RUT: ').bold = True
            p.add_run(f"{datos['RUT Empresa Lugar Geográfico']}\n")
            p.add_run('Nombre del centro de trabajo: ').bold = True
            p.add_run(f"{datos['Nombre Centro de Trabajo']}\n")
            p.add_run('CUV: ').bold = True
            p.add_run(f"{datos['CUV']}\n")
            p.add_run('CIIU: ').bold = True
            p.add_run(f"{datos['CIIU CT'].split('_')[-1]}\n")
            p.add_run('Fecha de activación del cuestionario: ').bold = True
            p.add_run(f"{datos['Fecha Inicio']}\n")
            p.add_run('Fecha de cierre del cuestionario: ').bold = True
            p.add_run(f"{datos['Fecha Fin']}\n")
            p.add_run('Universo de trabajadores de evaluación: ').bold = True
            p.add_run(f"{datos['Nº Trabajadores CT']}\n")
            p.paragraph_format.left_indent = Cm(15)

            # Salto de página
            doc.add_page_break()
            # Agregar imagen del logo en el encabezado (opcional)
            # doc.add_picture('logo.jpg', width=Inches(1))
            # last_paragraph = doc.paragraphs[-1]
            # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Título de sección
            doc.add_heading('RESULTADOS GENERALES', level=2)

            # Información de riesgo general
            p = doc.add_paragraph()
            p.add_run('Nivel de riesgo: ').bold = True
            p.add_run(f"{estado_riesgo}\n")
            p.style.font.size = Pt(12)

            # Insertar imagen del gráfico principal
            if fig_principal:
                img_buffer = BytesIO()
                fig_principal.savefig(img_buffer, format='png')
                img_buffer.seek(0)
                doc.add_picture(img_buffer, width=Inches(6))
                img_buffer.close()
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Obtener dimensiones en riesgo
            dimensiones_riesgo_alto = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (df_resultados_porcentaje['Puntaje'] == 2)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_medio = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (df_resultados_porcentaje['Puntaje'] == 1)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_bajo = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (df_resultados_porcentaje['Puntaje'] == -2)
                ]['Dimensión'].tolist()

            # Dimensiones en riesgo
            p = doc.add_paragraph()
            p.add_run('Dimensiones en riesgo alto: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_alto) if dimensiones_riesgo_alto else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo medio: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_medio) if dimensiones_riesgo_medio else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo bajo: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_bajo) if dimensiones_riesgo_bajo else 'Ninguna'}\n")

            # Salto de página
            doc.add_page_break()

            # Agregar gráficos por TE3
            for fig_te3, te3 in figs_te3:
                doc.add_heading(f"RESULTADOS POR ÁREA O GES {te3}", level=2)

                # Obtener el nivel de riesgo para este TE3
                riesgo_te3 = df_resumen[(df_resumen['CUV'] == datos['CUV']) & (df_resumen['TE3'] == te3)]['Riesgo']
                if not riesgo_te3.empty:
                    riesgo_te3 = riesgo_te3.values[0]
                else:
                    riesgo_te3 = "Información no disponible"

                p = doc.add_paragraph()
                p.add_run('Nivel de riesgo: ').bold = True
                p.add_run(f"{riesgo_te3}\n")
                p.style.font.size = Pt(12)

                # Insertar el gráfico
                img_buffer = BytesIO()
                fig_te3.savefig(img_buffer, format='png')
                img_buffer.seek(0)
                doc.add_picture(img_buffer, width=Inches(6))
                img_buffer.close()
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Obtener dimensiones en riesgo alto y medio
                dimensiones_riesgo_alto = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == datos['CUV']) & (df_porcentajes_niveles['TE3'] == te3) & (
                                df_porcentajes_niveles['Puntaje'] == 2)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_medio = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == datos['CUV']) & (df_porcentajes_niveles['TE3'] == te3) & (
                                df_porcentajes_niveles['Puntaje'] == 1)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_bajo = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == datos['CUV']) & (df_porcentajes_niveles['TE3'] == te3) & (
                                df_porcentajes_niveles['Puntaje'] == -2)
                    ]['Dimensión'].tolist()

                # Dimensiones en riesgo alto y medio
                p = doc.add_paragraph()
                p.add_run('Dimensiones en riesgo alto: ').bold = True
                p.add_run(f"{', '.join(dimensiones_riesgo_alto) if dimensiones_riesgo_alto else 'Ninguna'}\n")
                p.add_run('Dimensiones en riesgo medio: ').bold = True
                p.add_run(f"{', '.join(dimensiones_riesgo_medio) if dimensiones_riesgo_medio else 'Ninguna'}\n")
                p.add_run('Dimensiones en riesgo bajo: ').bold = True
                p.add_run(f"{', '.join(dimensiones_riesgo_bajo) if dimensiones_riesgo_bajo else 'Ninguna'}\n")

                # Salto de página
                doc.add_page_break()

            # Agregar tabla de medidas propuestas
            doc.add_heading(f"Medidas propuestas para {datos['Nombre Centro de Trabajo']}", level=2)
            agregar_tabla_ges_por_dimension(doc, df_res_dimTE3, datos['CUV'], df_recomendaciones,
                                            df_resultados_porcentaje, df_porcentajes_niveles, top_glosas, df_res_com)

            # Retornar el objeto Document
            return doc


    def generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, CUV, df_resumen,
                        df_res_dimTE3):
        """
        Genera el informe en Word para un CUV específico.
        """
        datos = df_res_com[df_res_com['CUV'] == CUV]
        estado = summary_df[summary_df['CUV'] == CUV]

        if datos.empty:
            st.error(f"No se encontró el CUV {CUV} en df_res_com.")
            return None
        if estado.empty:
            st.error(f"No se encontró el CUV {CUV} en summary_df.")
            return None

        row = datos.iloc[0]
        estado_riesgo = estado['Riesgo'].values[0]

        # Generar el gráfico principal
        fig_principal = generar_grafico_principal(df_resultados_porcentaje, CUV)
        if not fig_principal:
            st.warning("No se pudo generar el gráfico principal.")
            return None

        # Generar gráficos por TE3
        figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles, CUV)

        # Generar el contenido en el documento Word usando python-docx
        doc = generar_contenido_word(row, estado_riesgo, fig_principal, figs_te3, df_resumen,
                                     df_porcentajes_niveles, df_resultados_porcentaje)

        # Guardar el documento en un BytesIO para descarga
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer


    # Botón para generar y descargar el informe
    if st.button("Generar informe en Word"):

        if (uploaded_file_combined is not None and
                uploaded_file_rec is not None and
                uploaded_file_ciiu is not None and
                uploaded_file_resultados is not None and
                'df_res_com' in locals() and 'summary_df' in locals()):
            with st.spinner("Generando el informe, por favor espera..."):
                # Generar el documento
                doc_buffer = generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles,
                                             selected_cuv, df_resumen, df_res_dimTE3)

                if doc_buffer:
                    # Botón de descarga
                    st.download_button(
                        label="Descargar informe",
                        data=doc_buffer,
                        file_name=f"Informe_{selected_cuv}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning(
                "Los datos necesarios para generar el informe no están disponibles. Asegúrate de haber cargado todos los archivos requeridos.")





