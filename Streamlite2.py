# Importar las bibliotecas necesarias
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import unicodedata
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Configuración inicial de la aplicación
st.set_page_config(page_title="Generador de Informes de Riesgos Psicosociales", layout="wide")
st.title("Generador de Informes de Riesgos Psicosociales")
st.write("""
Esta aplicación permite generar informes técnicos basados en datos de riesgos psicosociales.
Por favor, cargue los archivos necesarios y siga las instrucciones.
""")

# Sección 1: Carga de archivos
st.header("1. Carga de archivos")

# Cargar el archivo Excel con múltiples hojas
uploaded_file_combined = st.file_uploader("Selecciona el archivo 'combined_output.xlsx'", type="xlsx")

# Cargar los archivos de recomendaciones y códigos CIIU
# uploaded_file_rec = st.file_uploader("Selecciona el archivo 'Recomendaciones2.xlsx'", type="xlsx")
# uploaded_file_ciiu = st.file_uploader("Selecciona el archivo 'ciiu.xlsx'", type="xlsx")

# Cargar el archivo 'resultados.xlsx'
# uploaded_file_resultados = st.file_uploader("Selecciona el archivo 'resultados.xlsx'", type="xlsx")

#precargas
uploaded_file_rec = 'Recomendaciones2.xlsx'
uploaded_file_ciiu = 'ciiu.xlsx'
uploaded_file_resultados = 'resultados.xlsx'


if (uploaded_file_combined is not None and
        uploaded_file_rec is not None and
        uploaded_file_ciiu is not None and
        uploaded_file_resultados is not None):

    # Leer las hojas necesarias desde 'combined_output.xlsx'
    try:
        df_res_com = pd.read_excel(uploaded_file_resultados, sheet_name='Datos')
        combined_df_base_complet3 = pd.read_excel(uploaded_file_combined, sheet_name='basecompleta')
        summary_df = pd.read_excel(uploaded_file_combined, sheet_name='Summary')
        df_resultados_porcentaje = pd.read_excel(uploaded_file_combined, sheet_name='resultado')
        df_porcentajes_niveles = pd.read_excel(uploaded_file_combined, sheet_name='df_porcentajes_niveles')
        df_res_dimTE3 = pd.read_excel(uploaded_file_combined, sheet_name='df_res_dimTE3')
        df_resumen = pd.read_excel(uploaded_file_combined, sheet_name='df_resumen')
        top_glosas = pd.read_excel(uploaded_file_combined, sheet_name='top_glosas')

        st.success("Archivo 'combined_output.xlsx' cargado exitosamente.")
    except Exception as e:
        st.error(f"Error al leer las hojas de 'combined_output.xlsx': {e}")

    # Leer los archivos de recomendaciones y CIIU
    try:
        df_rec = pd.read_excel(uploaded_file_rec, sheet_name='Hoja1')  # Cambia 'Hoja1' si es necesario
        df_ciiu = pd.read_excel(uploaded_file_ciiu, sheet_name='ciiu')  # Cambia 'ciiu' si es necesario

        st.success("Archivos 'Recomendaciones2.xlsx' y 'ciiu.xlsx' cargados exitosamente.")
    except Exception as e:
        st.error(f"Error al leer los archivos de recomendaciones y CIIU: {e}")

    # Leer y procesar 'resultados.xlsx'
    try:
        df_resultados = pd.read_excel(uploaded_file_resultados, sheet_name='Datos', usecols=['CUV', 'Folio'])
        df_res_com_resultados = pd.read_excel(uploaded_file_resultados, sheet_name='Datos')

        # Convertir 'CUV' a int64 en df_resultados
        df_resultados['CUV'] = pd.to_numeric(df_resultados['CUV'], errors='coerce').astype(
            'Int64')  # Usa 'Int64' para permitir NaNs

        st.success("Archivo 'resultados.xlsx' cargado y procesado exitosamente.")
    except Exception as e:
        st.error(f"Error al leer y procesar 'resultados.xlsx': {e}")

    # Procesamiento de df_rec para crear df_reco
    df_reco = pd.DataFrame(columns=['Dimensión', 'Rubro', 'Recomendación'])

    for index, row in df_rec.iterrows():
        dimension = row['Dimensión']
        for i in range(0, len(row) - 1, 2):
            rubro_col = f'Rubro.{i // 2}' if i // 2 > 0 else 'Rubro'
            recomendacion_col = f'Recomendación.{i // 2}' if i // 2 > 0 else 'Recomendación'

            if rubro_col in row and recomendacion_col in row:
                rubro = row[rubro_col]
                recomendacion = row[recomendacion_col]

                if pd.notna(rubro) and pd.notna(recomendacion):
                    df_reco = pd.concat(
                        [df_reco, pd.DataFrame([[dimension, rubro, recomendacion]], columns=df_reco.columns)],
                        ignore_index=True)

    # Realizar el merge entre df_ciiu y df_reco
    df_recomendaciones = pd.merge(df_ciiu, df_reco, left_on='Sección', right_on='Rubro', how='left')

    # Asegurar que las columnas 'CUV' y 'ciiu' sean de tipo str
    df_res_com['CUV'] = df_res_com['CUV'].astype(str)
    summary_df['CUV'] = summary_df['CUV'].astype(str)
    df_resultados_porcentaje['CUV'] = df_resultados_porcentaje['CUV'].astype(str)
    df_porcentajes_niveles['CUV'] = df_porcentajes_niveles['CUV'].astype(str)
    df_res_dimTE3['CUV'] = df_res_dimTE3['CUV'].astype(str)
    df_resumen['CUV'] = df_resumen['CUV'].astype(str)
    top_glosas['CUV'] = top_glosas['CUV'].astype(str)
    df_recomendaciones['ciiu'] = df_recomendaciones['ciiu'].astype(str)

    # Convertir las fechas al tipo datetime y luego al formato deseado
    columna_fecha_inicio = 'Fecha Inicio'  # Ajusta si el nombre es diferente

    if columna_fecha_inicio in df_res_com.columns:
        df_res_com['Fecha Inicio'] = pd.to_datetime(df_res_com[columna_fecha_inicio], errors='coerce').dt.strftime(
            '%d-%m-%Y')
        st.success("Columna 'Fecha Inicio' procesada correctamente.")
    else:
        st.error(
            f"La columna '{columna_fecha_inicio}' no se encontró en el DataFrame 'df_res_com'. Por favor, verifica el nombre de la columna.")
        st.write("Columnas disponibles en 'df_res_com':", df_res_com.columns.tolist())


    # Definir funciones auxiliares
    def normalizar_texto(texto):
        if isinstance(texto, str):
            texto = texto.strip().lower()
            texto = ''.join(
                c for c in unicodedata.normalize('NFD', texto)
                if unicodedata.category(c) != 'Mn'
            )
            return texto
        else:
            return ''


    def obtener_dimm_por_dimension(nombre_dimension):
        nombre_dimension_normalizado = normalizar_texto(nombre_dimension)
        # df_dimensiones debería estar definido o cargado desde alguna parte
        # Aquí debes asegurarte de tener el DataFrame 'dimensiones' disponible
        # Por ejemplo, podrías cargarlo desde un archivo o definirlo manualmente
        # Supongamos que está en 'df_res_dimTE3'
        # Ajusta esto según tu estructura de datos
        df_dimensiones = df_res_dimTE3[['Dimensión', 'dimm']].drop_duplicates()
        df_dimensiones['dimension_normalizada'] = df_dimensiones['Dimensión'].apply(normalizar_texto)
        resultado = df_dimensiones[df_dimensiones['dimension_normalizada'] == nombre_dimension_normalizado]
        if not resultado.empty:
            return resultado.iloc[0]['dimm']
        else:
            st.warning(f"No se encontró el código 'dimm' para la dimensión '{nombre_dimension}'.")
            return None


    def agregar_tabla_ges_por_dimension(doc, df, cuv, df_recomendaciones, df_resultados_porcentaje,
                                        df_porcentajes_niveles, top_glosas, df_res_com):
        """
        Agrega una tabla de medidas propuestas por dimensión en el documento Word.
        """
        # Filtrar el DataFrame para el CUV y puntajes 1 y 2
        df_cuv = df[(df['CUV'] == cuv) & (df['Puntaje'].isin([1, 2]))]
        if df_cuv.empty:
            st.warning(f"No hay datos con puntaje 1 o 2 para el CUV {cuv}.")
            return

        resultado = df_cuv.groupby('Dimensión')['TE3'].unique().reset_index()
        resultado['GES'] = resultado['TE3'].apply(lambda x: '; '.join(x))

        # Crear la tabla en el documento
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Agregar encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dimensión en riesgo'
        hdr_cells[1].text = 'Preguntas clave'
        hdr_cells[2].text = 'Explicación'
        hdr_cells[3].text = 'Medidas propuestas'
        hdr_cells[4].text = 'Fecha monitoreo'
        hdr_cells[5].text = 'Responsable seguimiento'

        # Rellenar la tabla con los datos
        for _, row in resultado.iterrows():
            dim = row['Dimensión']
            ges = row['GES']

            # Obtener las recomendaciones para esta dimensión y CIIU
            ciiu_valor = df_res_com[df_res_com['CUV'] == cuv]['CIIU CT'].iloc[0].split('_')[-1]
            print (ciiu_valor)


            recomendaciones = df_recomendaciones[
                (df_recomendaciones['Dimensión'] == dim) &
                (df_recomendaciones['ciiu'] == ciiu_valor)
                ]['Recomendación'].tolist()
            medidas_propuestas = '\n'.join([f"- {rec}" for rec in recomendaciones])

            # Obtener las preguntas clave desde top_glosas
            filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dim) & (top_glosas['CUV'] == cuv)]
            preguntas = filtro_glosas['Pregunta'].tolist()
            preguntas_text = '\n'.join(preguntas)

            # Obtener la descripción relacionada desde df_resultados_porcentaje
            descripcion = df_resultados_porcentaje[
                (df_resultados_porcentaje['Dimensión'] == dim) &
                (df_resultados_porcentaje['CUV'] == cuv)
                ]['Descripción'].dropna().tolist()

            if descripcion:
                descripcion_text = descripcion[0] + " para todo el centro de trabajo\n"
            else:
                descripcion_text = ""

            descripcion2 = [
                f"{desc} en {ges}"
                for desc in df_porcentajes_niveles[
                    (df_porcentajes_niveles['Dimensión'] == dim) &
                    (df_porcentajes_niveles['CUV'] == cuv) &
                    (df_porcentajes_niveles['TE3'] == ges) &
                    (df_porcentajes_niveles['Descripción'].str.strip() != '')
                    ]['Descripción'].dropna().tolist()
            ]

            descripcion2_text = '\n'.join(descripcion2).replace("[", "").replace("]", "").replace("'", "")

            # Verificar si hay múltiples GES
            te3_values = df_cuv['TE3'].unique()
            noges = 1 if len(te3_values) <= 1 else 0
            if noges == 1:
                descripcion2_text = ""
                st.info(f"Solo hay un GES para el CUV {cuv}. No se generarán recomendaciones por GES.")

            # Rellenar las celdas de la tabla
            row_cells = table.add_row().cells
            row_cells[0].text = f"{descripcion_text}\n{descripcion2_text}".strip()
            row_cells[1].text = preguntas_text.strip()
            row_cells[2].text = ''  # Espacio para 'Explicación'
            row_cells[3].text = medidas_propuestas.strip()
            row_cells[4].text = ''  # Fecha de monitoreo
            row_cells[5].text = ''  # Responsable seguimiento


    # Sección 2: Selección de CUV para generar el informe
    st.header("2. Seleccionar CUV para generar el informe")

    cuvs_disponibles = summary_df['CUV'].unique()
    selected_cuv = st.selectbox("Selecciona un CUV", cuvs_disponibles)

    # Filtrar los datos para el CUV seleccionado
    datos = df_res_com[df_res_com['CUV'] == selected_cuv]
    estado = summary_df[summary_df['CUV'] == selected_cuv]

    if datos.empty or estado.empty:
        st.error(f"No se encontraron datos para el CUV {selected_cuv}.")
    else:
        # Obtener la primera fila de datos
        datos = datos.iloc[0]
        estado_riesgo = estado['Riesgo'].values[0]

        # Sección 3: Generación y visualización de gráficos
        st.header("3. Gráfico principal de resultados")


        def generar_grafico_principal(df, CUV):
            df_filtrado = df[df['CUV'] == CUV]
            if df_filtrado.empty:
                st.warning(f"No se encontraron datos para el CUV {CUV}.")
                return None

            try:
                df_pivot = df_filtrado.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[
                           ::-1]
            except Exception as e:
                st.error(f"Error al pivotear los datos: {e}")
                return None

            fig, ax = plt.subplots(figsize=(12, 8))

            niveles = ["Bajo", "Medio", "Alto"]
            colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
            posiciones = np.arange(len(df_pivot.index))
            ancho_barra = 0.2

            for i, nivel in enumerate(niveles):
                if nivel in df_pivot.columns:
                    valores = df_pivot[nivel]
                    ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                            label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
                else:
                    st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot.")

            ax.axvline(50, color="blue", linestyle="--", linewidth=1)
            ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}", pad=50)
            ax.set_xlabel("Porcentaje")
            ax.set_ylabel("Dimensiones")
            ax.set_xlim(0, 100)
            ax.set_yticks(posiciones + ancho_barra)
            ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

            fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.6, 0.96), ncol=3)
            plt.subplots_adjust(left=0.3, top=0.85)
            plt.tight_layout()

            return fig


        # Generar y mostrar el gráfico principal
        fig_principal = generar_grafico_principal(df_resultados_porcentaje, selected_cuv)
        if fig_principal:
            st.pyplot(fig_principal)
        else:
            st.warning("No se pudo generar el gráfico principal.")

        # ---- NUEVA SECCIÓN PARA MOSTRAR DIMENSIONES EN RIESGO ----
        #st.header("4. Dimensiones en Riesgo")

        # Obtener dimensiones en riesgo
        dimensiones_riesgo_alto = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == 2)
                ]['Dimensión'].tolist()

        dimensiones_riesgo_medio = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == 1)
                ]['Dimensión'].tolist()

        dimensiones_riesgo_bajo = df_resultados_porcentaje[
            (df_resultados_porcentaje['CUV'] == selected_cuv) & (df_resultados_porcentaje['Puntaje'] == -2)
                ]['Dimensión'].tolist()

        # Mostrar dimensiones en riesgo en la web
        st.subheader("Dimensiones en riesgo")
        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("**Alto:**")
            if dimensiones_riesgo_alto:
                st.write(", ".join(dimensiones_riesgo_alto))
            else:
                st.write("Ninguna")

        with col2:
            st.markdown("**Medio:**")
            if dimensiones_riesgo_medio:
                st.write(", ".join(dimensiones_riesgo_medio))
            else:
                st.write("Ninguna")

        with col3:
            st.markdown("**Bajo:**")
            if dimensiones_riesgo_bajo:
                st.write(", ".join(dimensiones_riesgo_bajo))
            else:
                st.write("Ninguna")


        # Sección 4: Generación de gráficos por TE3
        st.header("4. Generación de gráficos por GES")


        def generar_graficos_por_te3(df, CUV):
            """
            Genera una lista de figuras de gráficos para cada valor de TE3 dentro del CUV especificado.
            """
            df_cuv = df[df['CUV'] == CUV]
            if df_cuv.empty:
                st.warning(f"No se encontraron datos para el CUV {CUV}.")
                return []

            if 'TE3' not in df_cuv.columns:
                st.warning(f"La columna 'TE3' no existe en el DataFrame para CUV {CUV}.")
                return []

            te3_values = df_cuv['TE3'].unique()

            # Verificar si hay más de un valor de TE3
            if len(te3_values) <= 1:
                st.info(f"Solo hay un valor de TE3 para el CUV {CUV}. No se generarán gráficos adicionales.")
                return []

            figs_te3 = []

            for te3 in te3_values:
                df_te3 = df_cuv[df_cuv['TE3'] == te3]
                if df_te3.empty:
                    continue
                try:
                    df_pivot = df_te3.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[
                               ::-1]
                except Exception as e:
                    st.error(f"Error al pivotear los datos para TE3 {te3}: {e}")
                    continue

                fig, ax = plt.subplots(figsize=(12, 8))

                niveles = ["Bajo", "Medio", "Alto"]
                colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
                posiciones = np.arange(len(df_pivot.index))
                ancho_barra = 0.2

                for i, nivel in enumerate(niveles):
                    if nivel in df_pivot.columns:
                        valores = df_pivot[nivel]
                        ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                                label=f"Riesgo {nivel.lower()} (%)", color=colores.get(nivel, 'grey'))
                    else:
                        st.warning(f"Nivel '{nivel}' no encontrado en las columnas de pivot para TE3 {te3}.")

                ax.axvline(50, color="blue", linestyle="--", linewidth=1)
                ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {CUV}, TE3 {te3}", pad=50)
                ax.set_xlabel("Porcentaje")
                ax.set_ylabel("Dimensiones")
                ax.set_xlim(0, 100)
                ax.set_yticks(posiciones + ancho_barra)
                ax.set_yticklabels(df_pivot.index, rotation=0, ha='right')

                fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.96), ncol=3)
                plt.tight_layout()

                figs_te3.append((fig, te3))

            return figs_te3


        # Generar y mostrar los gráficos por TE3
        figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles, selected_cuv)

        if figs_te3:
            for fig_te3, te3 in figs_te3:
                st.subheader(f"Gráfico para GES: {te3}")
                st.pyplot(fig_te3)
                dimensiones_riesgo_alto = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == 2)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_medio = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == 1)
                    ]['Dimensión'].tolist()

                dimensiones_riesgo_bajo = df_porcentajes_niveles[
                    (df_porcentajes_niveles['CUV'] == selected_cuv) &
                    (df_porcentajes_niveles['TE3'] == te3) &
                    (df_porcentajes_niveles['Puntaje'] == -2)
                    ]['Dimensión'].tolist()

                # Crear tres columnas
                col1, col2, col3 = st.columns(3)

                with col1:
                    st.markdown("**Alto:**")
                    if dimensiones_riesgo_alto:
                        st.write(", ".join(dimensiones_riesgo_alto))
                    else:
                        st.write("Ninguna")

                with col2:
                    st.markdown("**Medio:**")
                    if dimensiones_riesgo_medio:
                        st.write(", ".join(dimensiones_riesgo_medio))
                    else:
                        st.write("Ninguna")

                with col3:
                    st.markdown("**Bajo:**")
                    if dimensiones_riesgo_bajo:
                        st.write(", ".join(dimensiones_riesgo_bajo))
                    else:
                        st.write("Ninguna")


        else:
            st.info("No se generaron gráficos adicionales por TE3.")


        # Sección 5: Prescripciones de medidas
        st.header("5. Prescripciones de medidas")

        '''tabla_te3 = agregar_tabla_ges_por_dimension(df_porcentajes_niveles, df_recomendaciones, top_glosas, selected_cuv,
                                          te3)

        if not tabla_te3.empty:
            st.markdown("### Análisis de Dimensiones en Riesgo para este TE3")
            st.table(tabla_te3)  # Puedes usar st.dataframe(tabla_te3) si prefieres una tabla interactiva
        else:
            st.info("No hay dimensiones en riesgo para este TE3.")
'''

        # Sección 6: Generación del informe en Word
        st.header("6. Generación del informe en Word")

        def establecer_orientacion_apaisada(doc):
            """
            Configura el documento en orientación horizontal (apaisado).
            """
            section = doc.sections[0]
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)




        def generar_contenido_word(datos, estado_riesgo, fig_principal, figs_te3, df_resumen, df_porcentajes_niveles,
                                   df_resultados_porcentaje):
            """
            Genera el contenido del informe en un objeto Document de python-docx.
            """
            doc = Document()
            establecer_orientacion_apaisada(doc)

            # Configurar estilos y agregar contenido al documento
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(9)

            # Agregar encabezados y secciones
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()

            # Agregar imagen del logo (si tienes una imagen de logo)
            # doc.add_picture('logo.jpg', width=Inches(2))
            # last_paragraph = doc.paragraphs[-1]
            # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Título principal
            titulo = doc.add_heading('INFORME TÉCNICO', level=1)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtítulos
            subtitulo = doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2)
            subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitulo = doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2)
            subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            doc.add_paragraph()

            # Información general
            p = doc.add_paragraph()
            p.add_run('Razón Social: ').bold = True
            p.add_run(f"{datos['Nombre Empresa']}\n")
            p.add_run('RUT: ').bold = True
            p.add_run(f"{datos['RUT Empresa Lugar Geográfico']}\n")
            p.add_run('Nombre del centro de trabajo: ').bold = True
            p.add_run(f"{datos['Nombre Centro de Trabajo']}\n")
            p.add_run('CUV: ').bold = True
            p.add_run(f"{datos['CUV']}\n")
            p.add_run('CIIU: ').bold = True
            p.add_run(f"{datos['CIIU CT'].split('_')[-1]}\n")
            p.add_run('Fecha de activación del cuestionario: ').bold = True
            p.add_run(f"{datos['Fecha Inicio']}\n")
            p.add_run('Fecha de cierre del cuestionario: ').bold = True
            p.add_run(f"{datos['Fecha Fin']}\n")
            p.add_run('Universo de trabajadores de evaluación: ').bold = True
            p.add_run(f"{datos['Nº Trabajadores CT']}\n")
            p.paragraph_format.left_indent = Cm(1)

            # Salto de página
            doc.add_page_break()
            # Agregar imagen del logo en el encabezado (opcional)
            # doc.add_picture('logo.jpg', width=Inches(1))
            # last_paragraph = doc.paragraphs[-1]
            # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Título de sección
            doc.add_heading('RESULTADOS GENERALES', level=2)

            # Información de riesgo general
            p = doc.add_paragraph()
            p.add_run('Nivel de riesgo: ').bold = True
            p.add_run(f"{estado_riesgo}\n")
            p.style.font.size = Pt(12)

            # Insertar imagen del gráfico principal
            if fig_principal:
                img_buffer = BytesIO()
                fig_principal.savefig(img_buffer, format='png')
                img_buffer.seek(0)
                doc.add_picture(img_buffer, width=Inches(6))
                img_buffer.close()
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Obtener dimensiones en riesgo
            dimensiones_riesgo_alto = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (df_resultados_porcentaje['Puntaje'] == 2)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_medio = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (df_resultados_porcentaje['Puntaje'] == 1)
                ]['Dimensión'].tolist()

            dimensiones_riesgo_bajo = df_resultados_porcentaje[
                (df_resultados_porcentaje['CUV'] == datos['CUV']) & (df_resultados_porcentaje['Puntaje'] == -2)
                ]['Dimensión'].tolist()

            # Dimensiones en riesgo
            p = doc.add_paragraph()
            p.add_run('Dimensiones en riesgo alto: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_alto) if dimensiones_riesgo_alto else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo medio: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_medio) if dimensiones_riesgo_medio else 'Ninguna'}\n")
            p.add_run('Dimensiones en riesgo bajo: ').bold = True
            p.add_run(f"{', '.join(dimensiones_riesgo_bajo) if dimensiones_riesgo_bajo else 'Ninguna'}\n")

            # Salto de página
            doc.add_page_break()

            # Agregar gráficos por TE3
            for fig_te3, te3 in figs_te3:
                doc.add_heading(f"RESULTADOS POR ÁREA O GES {te3}", level=2)

                # Obtener el nivel de riesgo para este TE3
                riesgo_te3 = df_resumen[(df_resumen['CUV'] == datos['CUV']) & (df_resumen['TE3'] == te3)]['Riesgo']
                if not riesgo_te3.empty:
                    riesgo_te3 = riesgo_te3.values[0]
                else:
                    riesgo_te3 = "Información no disponible"

                p = doc.add_paragraph()
                p.add_run('Nivel de riesgo: ').bold = True
                p.add_run(f"{riesgo_te3}\n")
                p.style.font.size = Pt(12)

                # Insertar el gráfico
                img_buffer = BytesIO()
                fig_te3.savefig(img_buffer, format='png')
                img_buffer.seek(0)
                doc.add_picture(img_buffer, width=Inches(6))
                img_buffer.close()
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Salto de página
                doc.add_page_break()

            # Agregar tabla de medidas propuestas
            doc.add_heading(f"Medidas propuestas para {datos['Nombre Centro de Trabajo']}", level=2)
            agregar_tabla_ges_por_dimension(doc, df_res_dimTE3, datos['CUV'], df_recomendaciones,
                                            df_resultados_porcentaje, df_porcentajes_niveles, top_glosas, df_res_com)

            # Retornar el objeto Document
            return doc


    def generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, CUV, df_resumen,
                        df_res_dimTE3):
        """
        Genera el informe en Word para un CUV específico.
        """
        datos = df_res_com[df_res_com['CUV'] == CUV]
        estado = summary_df[summary_df['CUV'] == CUV]

        if datos.empty:
            st.error(f"No se encontró el CUV {CUV} en df_res_com.")
            return None
        if estado.empty:
            st.error(f"No se encontró el CUV {CUV} en summary_df.")
            return None

        row = datos.iloc[0]
        estado_riesgo = estado['Riesgo'].values[0]

        # Generar el gráfico principal
        fig_principal = generar_grafico_principal(df_resultados_porcentaje, CUV)
        if not fig_principal:
            st.warning("No se pudo generar el gráfico principal.")
            return None

        # Generar gráficos por TE3
        figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles, CUV)

        # Generar el contenido en el documento Word usando python-docx
        doc = generar_contenido_word(row, estado_riesgo, fig_principal, figs_te3, df_resumen,
                                     df_porcentajes_niveles, df_resultados_porcentaje)

        # Guardar el documento en un BytesIO para descarga
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer


    # Botón para generar y descargar el informe
    if st.button("Generar informe en Word"):

        if (uploaded_file_combined is not None and
                uploaded_file_rec is not None and
                uploaded_file_ciiu is not None and
                uploaded_file_resultados is not None and
                'df_res_com' in locals() and 'summary_df' in locals()):
            with st.spinner("Generando el informe, por favor espera..."):
                # Generar el documento
                doc_buffer = generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles,
                                             selected_cuv, df_resumen, df_res_dimTE3)

                if doc_buffer:
                    # Botón de descarga
                    st.download_button(
                        label="Descargar informe",
                        data=doc_buffer,
                        file_name=f"Informe_{selected_cuv}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning(
                "Los datos necesarios para generar el informe no están disponibles. Asegúrate de haber cargado todos los archivos requeridos.")
