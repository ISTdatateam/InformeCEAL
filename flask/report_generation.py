# report_generation.py

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import os
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def generate_main_chart(df_resultados_porcentaje, cuv):
    """
    Genera y guarda el gráfico principal de resultados para un CUV específico.

    Parámetros:
    df_resultados_porcentaje (pd.DataFrame): DataFrame con los porcentajes por dimensión y nivel.
    cuv (str): Identificador del CUV.

    Retorna:
    str: Ruta del archivo de imagen guardado.
    """
    df_filtrado = df_resultados_porcentaje[df_resultados_porcentaje['CUV'] == cuv]
    if df_filtrado.empty:
        logging.warning(f"No se encontraron datos para el CUV {cuv}.")
        return None

    # Pivoteo de los datos
    df_pivot = df_filtrado.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[::-1]

    # Configuración del gráfico
    fig, ax = plt.subplots(figsize=(12, 8))
    niveles = ["Bajo", "Medio", "Alto"]
    colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
    posiciones = np.arange(len(df_pivot.index))
    ancho_barra = 0.2

    for i, nivel in enumerate(niveles):
        valores = df_pivot[nivel]
        ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                label=f"Riesgo {nivel}", color=colores[nivel])

    ax.axvline(50, color="blue", linestyle="--", linewidth=1)
    ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {cuv}", pad=20)
    ax.set_xlabel("Porcentaje")
    ax.set_ylabel("Dimensiones")
    ax.set_xlim(0, 100)
    ax.set_yticks(posiciones + ancho_barra)
    ax.set_yticklabels(df_pivot.index)

    fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.95), ncol=3)
    plt.tight_layout()

    img_path = f"grafico_resultado_CUV_{cuv}.png"
    plt.savefig(img_path, format="png", bbox_inches="tight")
    plt.close()

    logging.info(f"Gráfico principal guardado en {img_path}")
    return img_path

def generate_te3_charts(df_resultados_porcentaje, cuv):
    """
    Genera y guarda gráficos para cada TE3 dentro del CUV especificado.

    Parámetros:
    df_porcentajes_niveles (pd.DataFrame): DataFrame con los porcentajes por dimensión, nivel y TE3.
    cuv (str): Identificador del CUV.

    Retorna:
    list: Lista de tuplas con la ruta de cada imagen generada y el valor de TE3 correspondiente.
    """

    df_cuv = df_resultados_porcentaje[df_resultados_porcentaje['CUV'] == cuv]

    # Verificar si 'TE3' está en df_cuv
    if 'TE3' not in df_cuv.columns:
        logging.error("La columna 'TE3' no existe en df_cuv.")
        print("Columnas disponibles en df_cuv:", df_cuv.columns.tolist())
        return []

    te3_values = df_cuv['TE3'].unique()
    #img_paths = []
    for te3_value in te3_values:
        df_te3 = df_cuv[df_cuv['TE3'] == te3_value]

    if df_cuv.empty:
        logging.warning(f"No se encontraron datos para el CUV {cuv}.")
        return []

    #te3_values = df_cuv['TE3'].unique()
    img_paths_te3 = []

    for te3 in te3_values:
        df_te3 = df_cuv[df_cuv['TE3'] == te3]
        df_pivot = df_te3.pivot(index="Dimensión", columns="Nivel", values="Porcentaje").fillna(0).iloc[::-1]

        fig, ax = plt.subplots(figsize=(12, 8))
        niveles = ["Bajo", "Medio", "Alto"]
        colores = {"Bajo": "green", "Medio": "orange", "Alto": "red"}
        posiciones = np.arange(len(df_pivot.index))
        ancho_barra = 0.2

        for i, nivel in enumerate(niveles):
            valores = df_pivot[nivel]
            ax.barh(posiciones + i * ancho_barra, valores, height=ancho_barra,
                    label=f"Riesgo {nivel}", color=colores[nivel])

        ax.axvline(50, color="blue", linestyle="--", linewidth=1)
        ax.set_title(f"Porcentaje de trabajadores por nivel de riesgo - CUV {cuv}, TE3 {te3}", pad=20)
        ax.set_xlabel("Porcentaje")
        ax.set_ylabel("Dimensiones")
        ax.set_xlim(0, 100)
        ax.set_yticks(posiciones + ancho_barra)
        ax.set_yticklabels(df_pivot.index)

        fig.legend(title="Nivel de Riesgo", loc="upper center", bbox_to_anchor=(0.5, 0.95), ncol=3)
        plt.tight_layout()

        img_path_te3 = f"grafico_resultado_CUV_{cuv}_TE3_{te3}.png"
        img_path_te3 = img_path_te3.replace('|', '_').replace(':', '_').replace('?', '_')
        plt.savefig(img_path_te3, format="png", bbox_inches="tight")
        plt.close()
        img_paths_te3.append((img_path_te3, te3))

        logging.info(f"Gráfico TE3 {te3} guardado en {img_path_te3}")

    return img_paths_te3

def create_word_report(datos, estado_riesgo, img_path_main, img_paths_te3, output_file):
    """
    Genera un informe en Word con los resultados y gráficos para un CUV específico.

    Parámetros:
    datos (pd.DataFrame): DataFrame con los datos del CUV.
    estado_riesgo (str): Nivel de riesgo del CUV.
    img_path_main (str): Ruta de la imagen principal del gráfico.
    img_paths_te3 (list): Lista de tuplas con rutas de imágenes y valores de TE3.
    output_file (str): Ruta donde se guardará el documento Word generado.
    """
    doc = Document()

    # Establecer orientación apaisada (horizontal)
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Establecer fuente predeterminada
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Título y subtítulo
    doc.add_heading('INFORME TÉCNICO', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Información general
    p = doc.add_paragraph()
    p.add_run('Razón Social: ').bold = True
    p.add_run(f"{datos['Nombre Empresa'].values[0]}\n")
    p.add_run('RUT: ').bold = True
    p.add_run(f"{datos['RUT Empresa Lugar Geográfico'].values[0]}\n")
    p.add_run('Nombre del centro de trabajo: ').bold = True
    p.add_run(f"{datos['Nombre Centro de Trabajo'].values[0]}\n")
    p.add_run('CUV: ').bold = True
    p.add_run(f"{datos['CUV'].values[0]}\n")
    p.add_run('CIIU: ').bold = True
    p.add_run(f"{datos['CIIU CT'].values[0].split('_')[-1]}\n")
    p.add_run('Fecha de activación del cuestionario: ').bold = True
    p.add_run(f"{datos['Fecha Inicio'].values[0]}\n")
    p.add_run('Fecha de cierre del cuestionario: ').bold = True
    p.add_run(f"{datos['Fecha Fin'].values[0]}\n")
    p.add_run('Universo de trabajadores de evaluación: ').bold = True
    p.add_run(f"{datos['Nº Trabajadores CT'].values[0]}\n")
    doc.add_paragraph()

    # Salto de página
    doc.add_page_break()

    # Título de sección
    doc.add_heading('RESULTADOS GENERALES CEAL-SM SUSESO', level=2)

    # Información de riesgo general
    p = doc.add_paragraph()
    p.add_run('Nivel de riesgo: ').bold = True
    p.add_run(f"{estado_riesgo}\n")
    p.style.font.size = Pt(12)

    # Insertar imagen principal
    if img_path_main:
        doc.add_picture(img_path_main, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salto de página
    doc.add_page_break()

    # Insertar gráficos por TE3
    for img_path_te3, te3 in img_paths_te3:
        if te3:
            doc.add_heading(f"RESULTADOS POR ÁREA O GES {te3}", level=2)
            doc.add_picture(img_path_te3, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

    # Guardar el documento
    doc.save(output_file)
    logging.info(f"Informe guardado como: {output_file}")

def generate_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, cuv, df_resumen, df_res_dimTE3):
    """
    Genera el informe completo para un CUV específico.

    Parámetros:
    df_res_com (pd.DataFrame): DataFrame con los datos de resumen.
    summary_df (pd.DataFrame): DataFrame con los puntajes y riesgos totales.
    df_resultados_porcentaje (pd.DataFrame): DataFrame con los porcentajes por dimensión y nivel.
    df_porcentajes_niveles (pd.DataFrame): DataFrame con los porcentajes por dimensión, nivel y TE3.
    cuv (str): Identificador del CUV.
    df_resumen (pd.DataFrame): DataFrame resumen.
    df_res_dimTE3 (pd.DataFrame): DataFrame con datos por dimensión y TE3.
    """
    datos = df_res_com[df_res_com['CUV'] == cuv]
    estado = summary_df[summary_df['CUV'] == cuv]

    if datos.empty:
        logging.warning(f"No se encontró el CUV {cuv} en df_res_com.")
        return
    if estado.empty:
        logging.warning(f"No se encontró el CUV {cuv} en summary_df.")
        return

    estado_riesgo = estado['Riesgo'].values[0]

    # Generar gráfico principal
    img_path_main = generate_main_chart(df_resultados_porcentaje, cuv)
    if not img_path_main:
        return

    # Generar gráficos por TE3
    img_paths_te3 = generate_te3_charts(df_porcentajes_niveles, cuv)

    # Generar informe en Word
    output_file = os.path.join('output_reports', f"{cuv}_informe.docx")
    create_word_report(datos, estado_riesgo, img_path_main, img_paths_te3, output_file)

    # Eliminar archivos temporales de gráficos
    os.remove(img_path_main)
    for img_path, _ in img_paths_te3:
        os.remove(img_path)

    logging.info(f"Informe generado para el CUV {cuv}.")

def agregar_tabla_ges_por_dimension(doc, df, cuv, df_recomendaciones, df_resultados_porcentaje, df_porcentajes_niveles, top_glosas, datos):
    """
    Agrega una tabla de dimensiones y GES para un CUV específico en el documento de Word.

    Parámetros:
    doc (Document): Objeto del documento de Word.
    df (pd.DataFrame): DataFrame con los datos de dimensiones y GES filtrados.
    cuv (str): El CUV específico para el que se generará la tabla.
    df_recomendaciones (pd.DataFrame): DataFrame con las recomendaciones por dimensión.
    df_resultados_porcentaje (pd.DataFrame): DataFrame con los porcentajes de riesgo generales.
    df_porcentajes_niveles (pd.DataFrame): DataFrame con los porcentajes de riesgo por TE3.
    top_glosas (pd.DataFrame): DataFrame con las preguntas clave.
    datos (pd.DataFrame): DataFrame con datos adicionales.

    Retorna:
    None
    """
    # Filtrar el DataFrame para el CUV específico y obtener solo las columnas necesarias
    df_cuv = df[(df['CUV'] == cuv) & (df['Puntaje'].isin([1, 2]))]

    # Agrupar por 'Dimensión' y combinar los valores únicos de 'TE3' en una lista separada por "; "
    resultado = df_cuv.groupby('Dimensión')['TE3'].unique().reset_index()
    resultado['GES'] = resultado['TE3'].apply(lambda x: '; '.join(x))
    resultado['GES'] = resultado['GES'].str.replace('|', '_', regex=False).str.replace(':', '_', regex=False).str.replace('?', '_', regex=False)

    # Obtener el valor de CIIU para el CUV
    datos['CIIU'] = datos['CIIU'].apply(lambda x: x.split('_')[-1] if isinstance(x, str) else x)
    ciiu_valor = datos.loc[datos['CUV'] == cuv, 'CIIU'].copy()
    if len(ciiu_valor) > 0:
        ciiu_unico = ciiu_valor.iloc[0]
        if isinstance(ciiu_unico, str) and ciiu_unico.isdigit():
            ciiu = int(ciiu_unico[:2]) if len(ciiu_unico) > 5 else int(ciiu_unico[:1])
        else:
            ciiu = None
    else:
        ciiu = None

    # Crear la tabla en el documento
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    column_widths = [Inches(1), Inches(1), Inches(1), Inches(4), Inches(1), Inches(1)]

    # Configurar el ancho de cada columna
    for col_idx, width in enumerate(column_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width

    # Agregar encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dimensión en riesgo'
    hdr_cells[1].text = 'Preguntas clave'
    hdr_cells[2].text = 'Explicación'
    hdr_cells[3].text = 'Medidas propuestas'
    hdr_cells[4].text = 'Fecha monitoreo'
    hdr_cells[5].text = 'Responsable seguimiento'

    # Rellenar la tabla con los datos de 'Dimensión' y 'GES'
    for _, row in resultado.iterrows():
        dim = row['Dimensión']
        ges = row['GES']

        # Obtener las recomendaciones para esta dimensión y CIIU
        recomendaciones = df_recomendaciones[
            (df_recomendaciones['Dimensión'] == dim) &
            (df_recomendaciones['ciiu'] == str(ciiu))
        ]['Recomendación'].tolist()
        medidas_propuestas = '\n'.join([f"- {rec}" for rec in recomendaciones])

        # Obtener la descripción relacionada desde df_resultados_porcentaje
        descripcion = df_resultados_porcentaje[
            (df_resultados_porcentaje['Dimensión'] == dim) &
            (df_resultados_porcentaje['CUV'] == cuv)
        ]['Descripción'].values
        descripcion = [desc for desc in descripcion if str(desc).strip() != '']

        # Obtener las preguntas clave desde top_glosas
        preguntas = top_glosas[
            (top_glosas['Dimensión'] == dim) &
            (top_glosas['CUV'] == cuv)
        ]['Pregunta'].values
        preguntas_text = '\n'.join(preguntas)

        # Rellenar las celdas de la tabla
        row_cells = table.add_row().cells
        row_cells[0].text = f"{dim}\n{ges}"
        row_cells[1].text = preguntas_text
        row_cells[2].text = ''  # Espacio para 'Explicación'
        row_cells[3].text = medidas_propuestas  # Medidas propuestas

    # Ajustar el tamaño de fuente de las celdas (opcional)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Ajusta el tamaño de la fuente


def establecer_orientacion_apaisada(doc):
    """
    Configura el documento en orientación apaisada (horizontal).

    Parámetros:
    doc (Document): Objeto del documento de Word.

    Retorna:
    None
    """
    # Acceder al elemento de configuración de la sección
    section = doc.sections[0]
    # Intercambiar los valores de ancho y alto para hacer la página horizontal
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    # Configurar márgenes (opcional)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def generar_contenido_word(datos, estado_riesgo, img_path_principal, img_paths_te3, df_resumen, df_porcentajes_niveles, df_resultados_porcentaje, output_file):
    """
    Genera el contenido del informe en un documento Word, incluyendo la información general,
    gráficos y recomendaciones de intervenciones.

    Parámetros:
    datos (pd.Series): Serie con los datos de un CUV específico.
    estado_riesgo (str): Nivel de riesgo del CUV.
    img_path_principal (str): Ruta de la imagen principal del gráfico.
    img_paths_te3 (list): Lista de tuplas con rutas de imágenes y valores de TE3.
    df_resumen (pd.DataFrame): DataFrame con datos resumidos de riesgo.
    df_porcentajes_niveles (pd.DataFrame): DataFrame con porcentajes por nivel de riesgo.
    df_resultados_porcentaje (pd.DataFrame): DataFrame con resultados por porcentaje.
    output_file (str): Ruta donde se guardará el documento Word generado.
    """
    # Crear un nuevo documento
    doc = Document()
    establecer_orientacion_apaisada(doc)

    # Configurar fuente y estilo
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Configurar idioma del documento en español
    lang = doc.styles['Normal'].element
    lang.set(qn('w:lang'), 'es-ES')

    # Título y subtítulo
    doc.add_heading('INFORME TÉCNICO', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Información general
    p = doc.add_paragraph()
    p.add_run('Razón Social: ').bold = True
    p.add_run(f"{datos['Nombre Empresa']}\n")
    p.add_run('RUT: ').bold = True
    p.add_run(f"{datos['RUT Empresa Lugar Geográfico']}\n")
    p.add_run('Nombre del centro de trabajo: ').bold = True
    p.add_run(f"{datos['Nombre Centro de Trabajo']}\n")
    p.add_run('CUV: ').bold = True
    p.add_run(f"{datos['CUV']}\n")
    p.add_run('CIIU: ').bold = True
    p.add_run(f"{datos['CIIU CT'].split('_')[-1]}\n")
    p.add_run('Fecha de activación del cuestionario: ').bold = True
    p.add_run(f"{datos['Fecha Inicio']}\n")
    p.add_run('Fecha de cierre del cuestionario: ').bold = True
    p.add_run(f"{datos['Fecha Fin']}\n")
    p.add_run('Universo de trabajadores de evaluación: ').bold = True
    p.add_run(f"{datos['Nº Trabajadores CT']}\n")
    doc.add_paragraph()

    # Salto de página
    doc.add_page_break()

    # Título de sección
    doc.add_heading('RESULTADOS GENERALES CEAL-SM SUSESO', level=2)

    # Información de riesgo general
    p = doc.add_paragraph()
    p.add_run('Nivel de riesgo: ').bold = True
    p.add_run(f"{estado_riesgo}\n")
    p.style.font.size = Pt(12)

    # Insertar imagen principal
    if img_path_principal:
        doc.add_picture(img_path_principal, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insertar gráficos por TE3
    for img_path_te3, te3 in img_paths_te3:
        if te3:
            doc.add_heading(f"RESULTADOS POR ÁREA O GES {te3}", level=2)
            doc.add_picture(img_path_te3, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

    # Guardar el documento
    doc.save(output_file)
    logging.info(f"Informe guardado como: {output_file}")

def generar_informes_para_todos_los_cuvs(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, df_resumen, df_res_dimTE3):
    """
    Genera informes para cada CUV en el DataFrame summary_df.

    Parámetros:
    df_res_com (pd.DataFrame): DataFrame con los datos de resumen.
    summary_df (pd.DataFrame): DataFrame con los puntajes y riesgos totales.
    df_resultados_porcentaje (pd.DataFrame): DataFrame con los porcentajes por dimensión y nivel.
    df_porcentajes_niveles (pd.DataFrame): DataFrame con los porcentajes por dimensión, nivel y TE3.
    df_resumen (pd.DataFrame): DataFrame resumen.
    df_res_dimTE3 (pd.DataFrame): DataFrame con datos por dimensión y TE3.
    """
    for cuv in summary_df['CUV'].unique():
        datos = df_res_com[df_res_com['CUV'] == cuv]
        estado = summary_df[summary_df['CUV'] == cuv]

        if datos.empty:
            logging.warning(f"No se encontró el CUV {cuv} en df_res_com.")
            continue
        if estado.empty:
            logging.warning(f"No se encontró el CUV {cuv} en summary_df.")
            continue

        estado_riesgo = estado['Riesgo'].values[0]

        # Generar gráfico principal
        img_path_main = generate_main_chart(df_resultados_porcentaje, cuv)
        if not img_path_main:
            continue

        # Generar gráficos por TE3
        img_paths_te3 = generate_te3_charts(df_porcentajes_niveles, cuv)

        # Generar informe en Word
        output_file = os.path.join('output_reports', f"{cuv}_informe.docx")
        generar_contenido_word(datos.iloc[0], estado_riesgo, img_path_main, img_paths_te3, df_resumen, df_porcentajes_niveles, df_resultados_porcentaje, output_file)

        # Eliminar archivos temporales de gráficos
        os.remove(img_path_main)
        for img_path, _ in img_paths_te3:
            os.remove(img_path)

        logging.info(f"Informe generado para el CUV {cuv}.")
