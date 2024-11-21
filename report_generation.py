# report_generation.py

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import streamlit as st


def establecer_orientacion_apaisada(doc):
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)


def agregar_tabla_ges_por_dimension(doc, dimensiones, cuv, dimensiones_recomendaciones,
                                    dimensiones_resultados_porcentaje, dimensiones_porcentajes_niveles, top_glosas,
                                    datos):
    """
    Agrega una tabla de dimensiones y GES para un CUV específico en el documento de Word.

    Parámetros:
    - doc (Document): El objeto del documento de Word.
    - dimensiones (list): Lista de diccionarios con información de dimensiones.
    - cuv (str): El CUV específico para el que se generará la tabla.
    - dimensiones_recomendaciones (pd.DataFrame): DataFrame con las recomendaciones por dimensión.
    - dimensiones_resultados_porcentaje (pd.DataFrame): DataFrame con las descripciones por dimensión y CUV.
    - dimensiones_porcentajes_niveles (pd.DataFrame): DataFrame con porcentajes y niveles por dimensión y GES.
    - top_glosas (pd.DataFrame): DataFrame con preguntas clave por dimensión y CUV.
    - datos (pd.DataFrame): DataFrame con información adicional, incluyendo CIIU.
    """

    # Crear la tabla en el documento
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    column_widths = [Inches(1), Inches(2), Inches(1), Inches(3), Inches(1), Inches(1)]

    # Configurar el ancho de cada columna
    for col_idx, width in enumerate(column_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width

    # Agregar encabezados
    hdr_cells = table.rows[0].cells
    headers = ['Dimensión en riesgo', 'Preguntas clave', 'Explicación',
               'Medidas propuestas', 'Fecha monitoreo', 'Responsable seguimiento']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Procesar cada dimensión en la lista
    for dim in dimensiones:
        dimension = dim['dimension']
        preguntas = dim['preguntas']
        medidas = dim['medidas']
        riesgo = dim.get('riesgo', '')
        area = dim.get('area', '')

        # Obtener las recomendaciones para esta dimensión y CIIU
        ciiu_valor = datos.loc[datos['CUV'] == cuv, 'CIIU'].iloc[0].split('_')[-1] if not datos.empty else ''
        recomendaciones = dimensiones_recomendaciones[
            (dimensiones_recomendaciones['Dimensión'] == dimension) &
            (dimensiones_recomendaciones['ciiu'] == str(ciiu_valor))
        ]['Recomendación'].tolist()
        medidas_propuestas = '\n'.join([f"- {rec}" for rec in recomendaciones]) if recomendaciones else 'Ninguna'

        # Obtener las preguntas clave desde top_glosas
        filtro_glosas = top_glosas[(top_glosas['Dimensión'] == dimension) & (top_glosas['CUV'] == cuv)]
        preguntas_text = '\n'.join(filtro_glosas['Pregunta'].tolist()) if not filtro_glosas.empty else 'Ninguna'

        # Obtener descripciones relacionadas
        descripcion = dimensiones_resultados_porcentaje[
            (dimensiones_resultados_porcentaje['Dimensión'] == dimension) &
            (dimensiones_resultados_porcentaje['CUV'] == cuv)
        ]['Descripción'].dropna().tolist()
        descripcion_text = descripcion[0] + " para todo el centro de trabajo\n" if descripcion else ""

        # Obtener descripciones adicionales por GES
        descripcion2 = []
        for medida in medidas:
            ges = medida.get('fecha_monitoreo', '').split('(')[-1].replace(')', '') if '(' in medida.get(
                'fecha_monitoreo', '') else ''
            desc = dimensiones_porcentajes_niveles[
                (dimensiones_porcentajes_niveles['Dimensión'] == dimension) &
                (dimensiones_porcentajes_niveles['CUV'] == cuv) &
                (dimensiones_porcentajes_niveles['TE3'] == ges) &
                (dimensiones_porcentajes_niveles['Descripción'].str.strip() != '')
            ]['Descripción'].tolist()
            if desc:
                descripcion2.append(f"{desc[0]} en {ges}")

        descripcion2_text = '\n'.join(descripcion2).replace("[", "").replace("]", "").replace("'",
                                                                                              "") if descripcion2 else ""

        # Rellenar las celdas de la tabla
        row_cells = table.add_row().cells
        row_cells[0].text = f"{descripcion_text}\n{descripcion2_text}".strip()
        row_cells[1].text = preguntas_text.strip()
        row_cells[2].text = f"{riesgo} en {area}" if riesgo and area else ""
        row_cells[3].text = medidas_propuestas.strip()
        row_cells[4].text = ''  # Fecha monitoreo (ya incluida en medidas)
        row_cells[5].text = ''  # Responsable seguimiento (ya incluido en medidas)

        # Opcional: Ajustar el tamaño de fuente de las celdas
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
######


def generar_informe(df_res_com, summary_df, df_resultados_porcentaje, df_porcentajes_niveles, CUV, df_resumen,
                    df_res_dimTE3, generar_grafico_principal, generar_graficos_por_te3, dimensiones, cuv,
                    df_recomendaciones, top_glosas, datos_df):
    """
    Genera el informe en Word para un CUV específico.

    Parámetros:
    - df_res_com (pd.DataFrame): DataFrame con resultados de la evaluación por CUV.
    - summary_df (pd.DataFrame): DataFrame con resumen de los CUV.
    - df_resultados_porcentaje (pd.DataFrame): DataFrame con resultados porcentuales.
    - df_porcentajes_niveles (pd.DataFrame): DataFrame con porcentajes y niveles por dimensión y GES.
    - CUV (str): Código único de evaluación.
    - df_resumen (pd.DataFrame): DataFrame con resumen de resultados.
    - df_res_dimTE3 (pd.DataFrame): DataFrame con resultados por dimensión y TE3.
    - generar_grafico_principal (function): Función para generar el gráfico principal.
    - generar_graficos_por_te3 (function): Función para generar gráficos por TE3.
    - dimensiones (list): Lista de diccionarios con información de dimensiones.
    - cuv (str): Código único de evaluación (redundante, podría eliminarse).
    - df_recomendaciones (pd.DataFrame): DataFrame con recomendaciones por dimensión.
    - top_glosas (pd.DataFrame): DataFrame con preguntas clave por dimensión y CUV.
    - datos_df (pd.DataFrame): DataFrame con información adicional, incluyendo CIIU.

    Retorna:
    - docx_buffer (BytesIO): Buffer con el documento de Word generado.
    """
    datos = df_res_com[df_res_com['CUV'] == CUV]
    estado = summary_df[summary_df['CUV'] == CUV]

    if datos.empty:
        st.error(f"No se encontró el CUV {CUV} en df_res_com.")
        return None
    if estado.empty:
        st.error(f"No se encontró el CUV {CUV} en summary_df.")
        return None

    row = datos.iloc[0]
    estado_riesgo = estado['Riesgo'].values[0]

    # Generar el gráfico principal
    fig_principal = generar_grafico_principal(df_resultados_porcentaje, CUV)
    if not fig_principal:
        st.warning("No se pudo generar el gráfico principal.")
        return None

    # Generar gráficos por TE3
    figs_te3 = generar_graficos_por_te3(df_porcentajes_niveles, CUV)

    # Obtener las dimensiones y medidas desde el parámetro
    if dimensiones:
        dimensiones = dimensiones
    else:
        dimensiones = []

    # Obtener el DataFrame 'datos' necesario para CIIU
    datos_df = df_res_com[df_res_com['CUV'] == CUV]

    # Generar el contenido en el documento Word usando python-docx
    doc = Document()
    establecer_orientacion_apaisada(doc)

    # Configurar estilos y agregar contenido al documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9)

    # Título principal
    titulo = doc.add_heading('INFORME TÉCNICO', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulos
    subtitulo = doc.add_heading('PRESCRIPCIÓN DE MEDIDAS PARA PROTOCOLO DE VIGILANCIA', level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = doc.add_heading('DE RIESGOS PSICOSOCIALES EN EL TRABAJO', level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Información general
    p = doc.add_paragraph()
    p.add_run('Razón Social: ').bold = True
    p.add_run(f"{row['Nombre Empresa']}\n")
    p.add_run('RUT: ').bold = True
    p.add_run(f"{row['RUT Empresa Lugar Geográfico']}\n")
    p.add_run('Nombre del centro de trabajo: ').bold = True
    p.add_run(f"{row['Nombre Centro de Trabajo']}\n")
    p.add_run('CUV: ').bold = True
    p.add_run(f"{row['CUV']}\n")
    p.add_run('CIIU: ').bold = True
    p.add_run(f"{row['CIIU CT'].split('_')[-1]}\n")
    p.add_run('Fecha de activación del cuestionario: ').bold = True
    p.add_run(f"{row['Fecha Inicio']}\n")
    p.add_run('Fecha de cierre del cuestionario: ').bold = True
    p.add_run(f"{row['Fecha Fin']}\n")
    p.add_run('Universo de trabajadores de evaluación: ').bold = True
    p.add_run(f"{row['Nº Trabajadores CT']}\n")
    p.paragraph_format.left_indent = Cm(1)

    # Salto de página
    doc.add_page_break()

    # Título de sección
    doc.add_heading('RESULTADOS GENERALES', level=2)

    # Información de riesgo general
    p = doc.add_paragraph()
    p.add_run('Nivel de riesgo: ').bold = True
    p.add_run(f"{estado_riesgo}\n")
    p.style.font.size = Pt(12)

    # Insertar imagen del gráfico principal
    if fig_principal:
        img_buffer = BytesIO()
        fig_principal.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        img_buffer.close()
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Obtener dimensiones en riesgo
    dimensiones_riesgo_alto = df_resultados_porcentaje[
        (df_resultados_porcentaje['CUV'] == cuv) & (df_resultados_porcentaje['Puntaje'] == 2)
        ]['Dimensión'].tolist()

    dimensiones_riesgo_medio = df_resultados_porcentaje[
        (df_resultados_porcentaje['CUV'] == cuv) & (df_resultados_porcentaje['Puntaje'] == 1)
        ]['Dimensión'].tolist()

    dimensiones_riesgo_bajo = df_resultados_porcentaje[
        (df_resultados_porcentaje['CUV'] == cuv) & (df_resultados_porcentaje['Puntaje'] == -2)
        ]['Dimensión'].tolist()

    # Dimensiones en riesgo
    p = doc.add_paragraph()
    p.add_run('Dimensiones en riesgo alto: ').bold = True
    p.add_run(f"{', '.join(dimensiones_riesgo_alto) if dimensiones_riesgo_alto else 'Ninguna'}\n")
    p.add_run('Dimensiones en riesgo medio: ').bold = True
    p.add_run(f"{', '.join(dimensiones_riesgo_medio) if dimensiones_riesgo_medio else 'Ninguna'}\n")
    p.add_run('Dimensiones en riesgo bajo: ').bold = True
    p.add_run(f"{', '.join(dimensiones_riesgo_bajo) if dimensiones_riesgo_bajo else 'Ninguna'}\n")

    # Salto de página
    doc.add_page_break()

    # Agregar gráficos por TE3
    for fig_te3, te3 in figs_te3:
        doc.add_heading(f"RESULTADOS POR ÁREA O GES {te3}", level=2)

        # Obtener el nivel de riesgo para este TE3
        riesgo_te3 = df_resumen[
            (df_resumen['CUV'] == cuv) &
            (df_resumen['TE3'] == te3)
            ]['Riesgo']
        riesgo_te3 = riesgo_te3.values[0] if not riesgo_te3.empty else "Información no disponible"

        p = doc.add_paragraph()
        p.add_run('Nivel de riesgo: ').bold = True
        p.add_run(f"{riesgo_te3}\n")
        p.style.font.size = Pt(12)

        # Insertar el gráfico
        img_buffer = BytesIO()
        fig_te3.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        img_buffer.close()
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Salto de página
        doc.add_page_break()

    # Agregar tabla de medidas propuestas
    doc.add_heading(f"Medidas propuestas para {row['Nombre Centro de Trabajo']}", level=2)
    agregar_tabla_ges_por_dimension(doc, dimensiones, cuv, df_recomendaciones, df_resultados_porcentaje,
                                    df_porcentajes_niveles, top_glosas, datos_df)

    # Guardar el documento en un BytesIO para descarga
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer

